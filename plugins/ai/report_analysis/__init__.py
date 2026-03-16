"""AI Plugin: Meldungs-Analyse (News Briefing + Export)."""

from plugins import PluginManager
from plugins.ai import AIPlugin

class ReportAnalysisPlugin(AIPlugin):
    plugin_id = "report_analysis"
    meta = {
        "label": "Meldungs-Analyse",
        "icon_svg": '<svg viewBox="0 0 24 24" width="18" height="18" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"><path d="M14 2H6a2 2 0 00-2 2v16a2 2 0 002 2h12a2 2 0 002-2V8z"/><polyline points="14,2 14,8 20,8"/><line x1="16" y1="13" x2="8" y2="13"/><line x1="16" y1="17" x2="8" y2="17"/><polyline points="10,9 9,9 8,9"/></svg>',
        "color": "#a78bfa",
        "description": "KI-gestützte Analyse und Zusammenfassung von Nachrichtenmeldungen.",
        "required_credentials": ["anthropic_api_key"],

    }

    def api_routes(self):
        return [
            {"rule": "/api/events/news-briefing", "endpoint": "api_events_news_briefing",
             "handler": self._handle_briefing, "methods": ["POST"]},
            {"rule": "/api/aipa-export/pdf", "endpoint": "api_aipa_export_pdf",
             "handler": self._handle_aipa_pdf, "methods": ["POST"]},
            {"rule": "/api/aipa-export/docx", "endpoint": "api_aipa_export_docx",
             "handler": self._handle_aipa_docx, "methods": ["POST"]},
            {"rule": "/api/story-export/pdf", "endpoint": "api_story_export_pdf",
             "handler": self._handle_story_pdf, "methods": ["POST"]},
            {"rule": "/api/story-export/docx", "endpoint": "api_story_export_docx",
             "handler": self._handle_story_docx, "methods": ["POST"]},
        ]

    # ── News Briefing ─────────────────────────────────────────────────────

    def _handle_briefing(self):
        import re
        import json as _json
        import logging
        import collections
        from concurrent.futures import ThreadPoolExecutor

        import requests as _req
        from bs4 import BeautifulSoup
        from flask import request, jsonify
        from plugins.ai._llm import call_llm, increment_usage

        log = logging.getLogger(__name__)

        guard = self._guard("report_analysis")
        if not isinstance(guard, tuple) or len(guard) != 2:
            return guard
        user_id, settings = guard

        body = request.get_json(force=True, silent=True) or {}
        articles = body.get("articles", [])
        query = (body.get("query") or "").strip()

        if not articles:
            return jsonify({"error": "Keine Artikel übermittelt."}), 400

        # ── Artikel-Texte parallel scrapen ──────────────────────────────
        scraped = self._scrape_articles(articles, log)

        # ── LLM-Prompt zusammenbauen ────────────────────────────────────
        prompt = self._build_briefing_prompt(scraped, query)

        result = call_llm(prompt, settings, max_tokens=8192)
        if not result["ok"]:
            return jsonify({"error": result["error"]}), result.get("status", 502)

        analysis = result["text"]

        increment_usage(user_id, source="news-briefing",
                        detail=f"Meldungen: {query or '?'} ({len(scraped)} Artikel)")

        # ── Ereignis-JSON extrahieren ───────────────────────────────────
        timeline = self._extract_timeline(analysis, log)

        if timeline:
            # Normalise sources: alte [{domain,url}] → [domain] Format
            for ev in timeline:
                srcs = ev.get("sources", [])
                if srcs and isinstance(srcs[0], dict):
                    ev["sources"] = [s.get("domain", "") for s in srcs if isinstance(s, dict)]
            # Briefing-Text = alles NACH dem letzten JSON-Objekt
            _last_obj_end = self._find_last_obj_end(analysis)
            rest = analysis[_last_obj_end:]
            rest = re.sub(r'^[\s,\]]*', '', rest)
            rest = re.sub(r'<<<\s*/?\s*EVENTS\s*>>>', '', rest)
            rest = re.sub(r'<!--\s*/?\s*TIMELINE\s*-->', '', rest)
            analysis = rest.strip()

        # Restliche Marker aufräumen
        analysis = re.sub(r'<<<\s*/?\s*EVENTS\s*>>>', '', analysis)
        analysis = re.sub(r'<!--\s*/?\s*TIMELINE\s*-->', '', analysis)
        analysis = analysis.strip()

        # ── Word-Cloud: Top-20-Begriffe aus Artikeltexten ───────────────
        top_words = self._extract_top_words(scraped)

        return jsonify({
            "briefing": analysis,
            "scraped_count": len([s for s in scraped if s.get("fulltext")]),
            "timeline": timeline,
            "top_words": top_words,
        })

    @staticmethod
    def _scrape_articles(articles, log):
        import requests as _req
        from bs4 import BeautifulSoup
        from concurrent.futures import ThreadPoolExecutor

        UA = (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        )

        def scrape_one(art):
            import re
            url = art.get("url", "")
            if not url:
                return {**art, "fulltext": ""}
            try:
                r = _req.get(url, headers={"User-Agent": UA}, timeout=12,
                             allow_redirects=True)
                r.raise_for_status()
                soup = BeautifulSoup(r.text, "html.parser")
                for tag in soup(["script", "style", "nav", "footer", "header",
                                 "aside", "form", "iframe", "noscript"]):
                    tag.decompose()
                for el in soup.find_all(attrs={"class": re.compile(
                        r'cookie|consent|gdpr|privacy|banner|overlay|popup|modal|cmp',
                        re.IGNORECASE)}):
                    el.decompose()
                for el in soup.find_all(attrs={"id": re.compile(
                        r'cookie|consent|gdpr|privacy|banner|overlay|cmp',
                        re.IGNORECASE)}):
                    el.decompose()
                article_el = soup.find("article")
                root = article_el if article_el else soup.find("body")
                if root:
                    paragraphs = root.find_all("p")
                    text = "\n".join(p.get_text(strip=True) for p in paragraphs
                                     if len(p.get_text(strip=True)) > 20)
                else:
                    text = soup.get_text(separator="\n", strip=True)
                if len(text) > 3000:
                    text = text[:3000] + " …"
                return {**art, "fulltext": text}
            except Exception as exc:
                log.debug("Scrape fehlgeschlagen für %s: %s", url, exc)
                return {**art, "fulltext": ""}

        with ThreadPoolExecutor(max_workers=6) as pool:
            return list(pool.map(scrape_one, articles[:50]))

    @staticmethod
    def _build_briefing_prompt(scraped, query):
        lines = [
            "Du bist ein erfahrener Nachrichtenanalyst und OSINT-Spezialist.",
            "Analysiere die folgenden Nachrichtenmeldungen und erstelle ein deutschsprachiges Briefing.",
            "",
            "═══════════════════════════════════════════════════════════",
            "DEINE ANTWORT MUSS EXAKT DIESES FORMAT HABEN (zwei Teile):",
            "═══════════════════════════════════════════════════════════",
            "",
            "TEIL 1 – EREIGNIS-JSON (MUSS am Anfang stehen!):",
            "Gib ZUERST eine Zeile mit exakt <<<EVENTS>>> aus,",
            "dann ein JSON-Array mit allen identifizierten Ereignissen,",
            "dann eine Zeile mit exakt <<</EVENTS>>>.",
            "",
            "Beispiel:",
            "<<<EVENTS>>>",
            '[{"dt":"2026-03-05T14:30","tz":"MEZ","title":"Explosion in Chemiefabrik","location":"Ludwigshafen","description":"Schwere Explosion in BASF-Werk","sources":["reuters.com","spiegel.de"]},{"dt":"2026-03-06","title":"Zweites Ereignis","location":"Berlin","description":"...","sources":["tagesschau.de"]}]',
            "<<</EVENTS>>>",
            "",
            "Regeln für das JSON:",
            "- KOMPAKTES JSON auf EINER Zeile pro Objekt – KEIN pretty-printing!",
            "- dt: TATSÄCHLICHER Zeitpunkt des Geschehens, NICHT das Veröffentlichungsdatum!",
            "  Das 'Veröffentlichungsdatum' bei jedem Artikel ist nur wann die Meldung erschien.",
            "  Du MUSST aus dem Artikeltext extrahieren, WANN das Ereignis wirklich stattfand.",
            "  Suche nach: 'am Montag', 'gestern', 'um 14:30 Uhr', konkreten Datumsangaben, etc.",
            "  Format: YYYY-MM-DDTHH:MM (Uhrzeit weglassen wenn unbekannt → nur YYYY-MM-DD)",
            "  Nur wenn KEINE Zeitangabe im Text: Veröffentlichungsdatum als Fallback.",
            "- tz: Zeitzone (UTC, MEZ, EST, etc.) – weglassen wenn unbekannt",
            "- location: Ort/Land des Geschehens",
            "- sources: NUR domain-Name als String-Array, z.B. [\"reuters.com\",\"spiegel.de\"] – KEINE URLs!",
            "  Wenn mehrere Artikel dasselbe Ereignis beschreiben → alle Domains auflisten!",
            "- Chronologisch sortiert (ältestes zuerst)",
            "- NUR gültiges JSON, keine Kommentare",
            "",
            "TEIL 2 – BRIEFING-TEXT (nach dem JSON-Block):",
            "Beginne SOFORT mit einer Zusammenfassung (2–3 Sätze), OHNE Überschrift.",
            "Danach verwende ## und ### für Abschnitte.",
            "Detaillierte chronologische Darstellung mit:",
            "- **fett** für wichtige Namen, Orte, Zeitangaben",
            "- Aufzählungszeichen (- ) für Listen",
            "- Quellenangaben in eckigen Klammern: [domain.com, URL] bei jedem Ereignis",
            "- Mehrere Quellen: [a.com, URL] [b.com, URL]",
            "- Am Ende: Einordnung / Bewertung (1–2 Sätze)",
            "",
        ]

        if query:
            lines += [f"Suchbegriff: {query}", ""]

        lines.append(f"Anzahl Artikel: {len(scraped)}")
        lines.append("")

        for i, art in enumerate(scraped, 1):
            lines.append(f"--- Artikel {i} ---")
            lines.append(f"Titel: {art.get('title', '–')}")
            if art.get("title_original"):
                lines.append(f"Originaltitel: {art['title_original']}")
            lines.append(f"Veröffentlichungsdatum (NICHT Ereignisdatum!): {art.get('seendate', '–')}")
            lines.append(f"Quelle: {art.get('domain', '–')} ({art.get('sourcecountry', '–')})")
            lines.append(f"Sprache: {art.get('language', '–')}")
            lines.append(f"URL: {art.get('url', '–')}")
            fulltext = art.get("fulltext", "").strip()
            if fulltext:
                lines.append(f"Volltext:\n{fulltext}")
            else:
                lines.append("Volltext: [nicht abrufbar]")
            lines.append("")

        return "\n".join(lines)

    @staticmethod
    def _extract_timeline(analysis, log):
        import re
        import json as _json

        log.info("LLM-Antwort Anfang (300z): %s", repr(analysis[:300]))

        timeline = []
        for _m in re.finditer(r'\{', analysis):
            obj_start = _m.start()
            _d, _ins, _j = 0, False, obj_start
            obj_end = -1
            while _j < len(analysis):
                _c = analysis[_j]
                if _ins:
                    if _c == '\\':
                        _j += 2
                        continue
                    elif _c == '"':
                        _ins = False
                else:
                    if _c == '"':
                        _ins = True
                    elif _c == '{':
                        _d += 1
                    elif _c == '}':
                        _d -= 1
                        if _d == 0:
                            obj_end = _j + 1
                            break
                _j += 1
            if obj_end > 0:
                try:
                    obj = _json.loads(analysis[obj_start:obj_end])
                    if isinstance(obj, dict) and 'dt' in obj and 'title' in obj:
                        timeline.append(obj)
                except _json.JSONDecodeError:
                    pass

        if timeline:
            log.info("Events extrahiert: %d Ereignisse", len(timeline))
        else:
            log.warning("Kein Ereignis-Block extrahiert. Anfang: %s", repr(analysis[:200]))

        return timeline

    @staticmethod
    def _find_last_obj_end(analysis):
        """Find the end position of the last JSON object with 'dt' key."""
        import re
        import json as _json

        last_end = 0
        for _m in re.finditer(r'\{', analysis):
            obj_start = _m.start()
            _d, _ins, _j = 0, False, obj_start
            obj_end = -1
            while _j < len(analysis):
                _c = analysis[_j]
                if _ins:
                    if _c == '\\':
                        _j += 2
                        continue
                    elif _c == '"':
                        _ins = False
                else:
                    if _c == '"':
                        _ins = True
                    elif _c == '{':
                        _d += 1
                    elif _c == '}':
                        _d -= 1
                        if _d == 0:
                            obj_end = _j + 1
                            break
                _j += 1
            if obj_end > 0:
                try:
                    obj = _json.loads(analysis[obj_start:obj_end])
                    if isinstance(obj, dict) and 'dt' in obj and 'title' in obj:
                        last_end = max(last_end, obj_end)
                except _json.JSONDecodeError:
                    pass
        return last_end

    @staticmethod
    def _extract_top_words(scraped):
        import re
        import collections

        _STOP = {
            # Deutsch
            "aber","alle","allem","allen","aller","alles","also","andere","anderem",
            "anderen","anderer","anderes","auch","been","beim","bereits","bisher",
            "bitte","dabei","dadurch","dafür","dagegen","daher","dahin","damals",
            "damit","danach","daneben","dann","daran","darauf","daraus","darf",
            "darfst","darin","darum","darunter","davon","davor","dazu","dein",
            "deine","deinem","deinen","deiner","demnach","denen","denn","dennoch",
            "deren","derer","derselbe","derselben","deshalb","dessen","dich","dies",
            "diese","dieselbe","dieselben","diesem","diesen","dieser","dieses",
            "doch","dort","drei","drin","dritte","dritten","dritter","drum","drunter",
            "dass","dürf","dürfe","dürfen","dürft","dürfte","dürften","eben",
            "ebenso","eigenen","eigene","eigentlich","einander","eine","einem",
            "einen","einer","einige","einigem","einigen","einiger","einiges",
            "einmal","erst","erste","erstem","ersten","erster","erstes",
            "etwa","etwas","euch","euer","eure","eurem","euren","eurer",
            "ganz","ganzen","gegen","gehen","genau","gerade","gering","gerne",
            "gibt","ging","grosse","grossen","großen","große","großer","großes",
            "guten","guter","gutes","habe","haben","halt","hatte","hatten",
            "hätt","hätte","hätten","hier","hinter","hoch","häufig",
            "ihre","ihrem","ihren","ihrer","immer","indem","infolge",
            "innen","innerhalb","insgesamt","irgend","jede","jedem","jeden",
            "jeder","jedes","jedoch","jene","jenem","jenen","jener","jenes",
            "jetzt","kann","kannst","kaum","kein","keine","keinem","keinen",
            "keiner","kenn","klar","kommen","konnte","können","könnt","könnte",
            "könnten","lang","lange","langen","längst","laut","lediglich",
            "letzt","letzte","letzten","letzter","letztes","liegt","lässt",
            "machen","macht","manch","manche","manchem","manchen","mancher",
            "manchmal","mehr","mein","meine","meinem","meinen","meiner",
            "mich","mind","mindestens","mirs","mocht","mochte","möchte",
            "mögen","möglich","morgen","morgens","muss","musst","musste",
            "müssen","müsst","müsste","nach","nachdem","nachher","nacher",
            "neben","nehmen","nein","neue","neuem","neuen","neuer","neues",
            "neun","neunte","neunten","neunter","neuntes","nicht","nichts",
            "noch","nochmal","nächste","nächsten","nämlich","nötig",
            "nur","nämlich","oben","obgleich","obwohl","oder","ohne",
            "rund","sagte","schließen","schon","sehr","seid","sein",
            "seine","seinem","seinen","seiner","seit","seitdem","seither",
            "sich","sicher","sicherlich","sind","sogar","solle","sollen",
            "sollte","sollten","solltest","somit","sondern","sonst","sorgt",
            "soviel","sowie","über","überhaupt","übrigens","unser","unsere",
            "unserem","unseren","unserer","unten","unter","viel","viele",
            "vielen","vielleicht","voll","völlig","vom","vor","vorbei",
            "vorher","vorne","vorüber","wahr","wahrscheinlich","wann","warum",
            "weder","wegen","weil","weit","weiter","weitere","weiterem",
            "weiteren","weiterer","weiteres","welch","welche","welchem",
            "welchen","welcher","wenig","wenige","wenigen","weniger","wenigstens",
            "wenn","wer","werde","werden","weshalb","wessen","wieder",
            "will","wir","wird","wirklich","wissen","wohl","wollen","worden",
            "wurde","würde","würden","während","wäre","wären","zehn",
            "zeigt","ziemlich","zwischen","zwar",
            # Englisch
            "about","after","again","also","another","because","been","before",
            "being","between","both","could","does","done","down","each",
            "even","every","from","have","here","into","just","like","made",
            "many","more","most","much","must","never","only","other","over",
            "said","same","should","since","some","still","such","than","that",
            "them","then","there","these","they","this","those","through",
            "under","upon","very","want","well","were","what","when","where",
            "which","while","whom","will","with","within","without","would",
            "your","about","could","would","their","first","people",
            # Web-Boilerplate
            "cookie","cookies","datenschutz","impressum","akzeptieren",
            "newsletter","anmelden","registrieren","webseite","website",
            "klicken","weiterlesen","anzeige","werbung","tracking",
            "facebook","twitter","instagram","google","youtube","whatsapp",
        }
        wc_text = " ".join(
            (art.get("title") or "") + " " + (art.get("fulltext") or "")
            for art in scraped
        ).lower()
        wc_words = re.findall(r'[a-zäöüß]{4,}', wc_text)
        wc_freq = collections.Counter(w for w in wc_words if w not in _STOP)
        return [w for w, _ in wc_freq.most_common(20)]

    # ── APA Export (PDF / DOCX) ───────────────────────────────────────────

    @staticmethod
    def _aipa_build_html(report_md, chat_messages):
        """Build styled HTML from APA report markdown + optional chat."""
        import markdown
        import re
        report_md = re.sub(r'\{\{SNAPSHOT:\d+\}\}', '', report_md)
        report_html = markdown.markdown(report_md, extensions=["tables", "fenced_code"])
        chat_html = ""
        if chat_messages:
            chat_html = '<h2 style="margin-top:36px;padding-top:20px;border-top:2px solid #ccc;">Anschlussfragen</h2>'
            for m in chat_messages:
                role = m.get("role", "")
                text = markdown.markdown(m.get("content", ""), extensions=["tables"])
                if role == "user":
                    chat_html += f'<div style="margin:14px 0 6px;"><strong style="color:#7c5cbf;">Frage:</strong></div>{text}'
                elif role == "assistant":
                    chat_html += f'<div style="margin:6px 0 14px;padding:10px 14px;background:#f4f0ff;border-radius:8px;border-left:3px solid #7c5cbf;">{text}</div>'
        return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
body {{ font-family: 'Segoe UI', Arial, sans-serif; font-size: 13px; line-height: 1.7;
       max-width: 700px; margin: 40px auto; color: #222; }}
h1 {{ font-size: 20px; }} h2 {{ font-size: 16px; color: #444; }} h3 {{ font-size: 14px; }}
table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
th, td {{ border: 1px solid #ccc; padding: 6px 10px; font-size: 12px; }}
th {{ background: #f5f5f5; font-weight: 600; }}
ul, ol {{ margin: 6px 0 6px 22px; }}
strong {{ color: #333; }}
</style></head><body>{report_html}{chat_html}</body></html>"""

    def _handle_aipa_pdf(self):
        import io
        from weasyprint import HTML
        from flask import request, jsonify, send_file
        from flask_login import current_user

        if not current_user.is_authenticated:
            return jsonify({"error": "Nicht eingeloggt."}), 401

        body = request.get_json(force=True, silent=True) or {}
        report_md = (body.get("report") or "").strip()
        chat = body.get("chat") or []
        if not report_md:
            return jsonify({"error": "Kein Bericht vorhanden."}), 400
        html = self._aipa_build_html(report_md, chat)
        buf = io.BytesIO()
        HTML(string=html).write_pdf(buf)
        buf.seek(0)
        from app import audit_log
        audit_log("export_pdf", "report", None, "APA-Bericht als PDF exportiert")
        return send_file(buf, mimetype="application/pdf",
                         as_attachment=True, download_name="APA-Bericht.pdf")

    def _handle_aipa_docx(self):
        import io
        import re
        from docx import Document
        from docx.shared import Pt, RGBColor
        from flask import request, jsonify, send_file
        from flask_login import current_user

        if not current_user.is_authenticated:
            return jsonify({"error": "Nicht eingeloggt."}), 401

        body = request.get_json(force=True, silent=True) or {}
        report_md = (body.get("report") or "").strip()
        chat = body.get("chat") or []
        if not report_md:
            return jsonify({"error": "Kein Bericht vorhanden."}), 400
        report_md = re.sub(r'\{\{SNAPSHOT:\d+\}\}', '', report_md)

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)

        self._add_md_to_docx(doc, report_md)

        if chat:
            doc.add_page_break()
            doc.add_heading("Anschlussfragen", level=1)
            for m in chat:
                role = m.get("role", "")
                content = m.get("content", "")
                if role == "user":
                    p = doc.add_paragraph()
                    run = p.add_run("Frage: ")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x7c, 0x5c, 0xbf)
                    p.add_run(content)
                elif role == "assistant":
                    self._add_md_to_docx(doc, content)
                    doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True, download_name="APA-Bericht.docx")

    # ── Story Export (PDF / DOCX) ─────────────────────────────────────────

    @staticmethod
    def _story_build_html(briefing_md, timeline):
        """Build styled HTML from story briefing + timeline events."""
        import markdown
        import re
        briefing_html = ""
        if briefing_md:
            briefing_md = re.sub(r'\{\{SNAPSHOT:\d+\}\}', '', briefing_md)
            briefing_html = markdown.markdown(briefing_md, extensions=["tables", "fenced_code"])

        tl_html = ""
        if timeline:
            tl_html = '<h2 style="margin-top:30px;padding-top:16px;border-top:2px solid #7c5cbf;">Chronologie</h2>'
            for ev in timeline:
                dt = ev.get("dt", "")
                tz = ev.get("tz", "")
                title = ev.get("title", "–")
                desc = ev.get("description", "")
                location = ev.get("location", "")
                sources = ev.get("sources", [])
                src_list = ", ".join(
                    f'[{s}]' if isinstance(s, str) else f'[{s.get("domain", "?")}]'
                    for s in sources
                ) if sources else ""

                date_str = ""
                if dt:
                    parts = dt.split("T")
                    d = parts[0].split("-")
                    if len(d) == 3:
                        date_str = f"{d[2]}.{d[1]}.{d[0]}"
                    if len(parts) > 1:
                        date_str += f" {parts[1][:5]} Uhr"
                    if tz:
                        date_str += f" {tz}"

                tl_html += f'''<div style="margin:16px 0;padding:12px 16px;border-left:3px solid #7c5cbf;
                    background:#f9f7ff;border-radius:0 6px 6px 0;">
                    <div style="font-size:11px;color:#7c5cbf;font-weight:600;">{date_str}</div>
                    <div style="font-size:14px;font-weight:700;margin:4px 0;">{title}</div>'''
                if desc:
                    tl_html += f'<div style="font-size:12px;margin:4px 0;">{desc}</div>'
                if location:
                    tl_html += f'<div style="font-size:11px;color:#666;">Ort: {location}</div>'
                if src_list:
                    tl_html += f'<div style="font-size:10px;color:#999;margin-top:4px;">Quellen: {src_list}</div>'
                tl_html += '</div>'

        return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
body {{ font-family: 'Segoe UI', Arial, sans-serif; font-size: 13px; line-height: 1.7;
       max-width: 700px; margin: 40px auto; color: #222; }}
h1 {{ font-size: 20px; color: #7c5cbf; }} h2 {{ font-size: 16px; color: #444; }} h3 {{ font-size: 14px; }}
table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
th, td {{ border: 1px solid #ccc; padding: 6px 10px; font-size: 12px; }}
th {{ background: #f5f5f5; font-weight: 600; }}
ul, ol {{ margin: 6px 0 6px 22px; }}
strong {{ color: #333; }}
</style></head><body>
<h1>Meldungsbriefing</h1>
{briefing_html}
{tl_html}
<div style="margin-top:40px;padding-top:12px;border-top:1px solid #ddd;font-size:10px;color:#aaa;">
Erstellt mit VeriTrend.ai</div>
</body></html>"""

    def _handle_story_pdf(self):
        import io
        from weasyprint import HTML
        from flask import request, jsonify, send_file
        from flask_login import current_user

        guard = self._guard("report_analysis")
        if not isinstance(guard, tuple) or len(guard) != 2:
            return guard

        body = request.get_json(force=True, silent=True) or {}
        briefing_md = (body.get("briefing") or "").strip()
        timeline = body.get("timeline") or []
        if not briefing_md and not timeline:
            return jsonify({"error": "Keine Daten vorhanden."}), 400
        html = self._story_build_html(briefing_md, timeline)
        buf = io.BytesIO()
        HTML(string=html).write_pdf(buf)
        buf.seek(0)
        from app import audit_log
        audit_log("export_pdf", "story", None, "Story als PDF exportiert")
        return send_file(buf, mimetype="application/pdf",
                         as_attachment=True, download_name="Meldungsbriefing.pdf")

    def _handle_story_docx(self):
        import io
        import re
        from docx import Document
        from docx.shared import Pt, RGBColor
        from flask import request, jsonify, send_file

        guard = self._guard("report_analysis")
        if not isinstance(guard, tuple) or len(guard) != 2:
            return guard

        body = request.get_json(force=True, silent=True) or {}
        briefing_md = (body.get("briefing") or "").strip()
        timeline = body.get("timeline") or []
        if not briefing_md and not timeline:
            return jsonify({"error": "Keine Daten vorhanden."}), 400
        briefing_md = re.sub(r'\{\{SNAPSHOT:\d+\}\}', '', briefing_md)

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)

        doc.add_heading("Meldungsbriefing", level=0)

        if briefing_md:
            self._add_md_to_docx(doc, briefing_md)

        if timeline:
            doc.add_page_break()
            doc.add_heading("Chronologie", level=1)
            for ev in timeline:
                dt = ev.get("dt", "")
                tz = ev.get("tz", "")
                title = ev.get("title", "–")
                desc = ev.get("description", "")
                location = ev.get("location", "")
                sources = ev.get("sources", [])

                date_str = ""
                if dt:
                    parts = dt.split("T")
                    d = parts[0].split("-")
                    if len(d) == 3:
                        date_str = f"{d[2]}.{d[1]}.{d[0]}"
                    if len(parts) > 1:
                        date_str += f" {parts[1][:5]} Uhr"
                    if tz:
                        date_str += f" {tz}"

                p = doc.add_paragraph()
                if date_str:
                    run = p.add_run(date_str + "  ")
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x7c, 0x5c, 0xbf)
                run = p.add_run(title)
                run.bold = True

                if desc:
                    doc.add_paragraph(desc)
                meta = []
                if location:
                    meta.append(f"Ort: {location}")
                if sources:
                    src_list = ", ".join(
                        f"[{s}]" if isinstance(s, str)
                        else f"[{s.get('domain', '?')}]"
                        for s in sources
                    )
                    meta.append(f"Quellen: {src_list}")
                if meta:
                    p = doc.add_paragraph(" · ".join(meta))
                    p.runs[0].font.size = Pt(9)
                    p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Footer
        p = doc.add_paragraph()
        p.add_run("")
        p = doc.add_paragraph()
        run = p.add_run("Erstellt mit VeriTrend.ai")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        from app import audit_log
        audit_log("export_docx", "story", None, "Story als DOCX exportiert")
        return send_file(buf, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True, download_name="Meldungsbriefing.docx")

    # ── Shared: Markdown → DOCX ──────────────────────────────────────────

    @staticmethod
    def _add_md_to_docx(doc, md_text):
        """Simple markdown-to-docx renderer."""
        import re
        for line in md_text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", stripped):
                doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
            else:
                p = doc.add_paragraph()
                parts = re.split(r"(\*\*.*?\*\*)", stripped)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

PluginManager.register(ReportAnalysisPlugin())
