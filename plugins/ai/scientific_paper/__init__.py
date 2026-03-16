"""AI Plugin: Scientific Paper — KI-generiertes wissenschaftliches Paper als .docx."""

import io
import logging
import re

from plugins import PluginManager
from plugins.ai import AIPlugin

log = logging.getLogger(__name__)

GP = {"": "Web", "news": "News", "images": "Images",
      "youtube": "YouTube", "froogle": "Shopping"}
TF = {"now 1-H": "1h", "now 4-H": "4h", "now 1-d": "24h", "now 7-d": "7 days",
      "today 1-m": "1 month", "today 3-m": "3 months",
      "today 12-m": "12 months", "today 5-y": "5 years"}


class ScientificPaperPlugin(AIPlugin):
    plugin_id = "scientific_paper"
    meta = {
        "label": "Scientific Paper",
        "icon_svg": (
            '<svg viewBox="0 0 24 24" width="18" height="18" fill="none" '
            'stroke="currentColor" stroke-width="1.5" stroke-linecap="round" '
            'stroke-linejoin="round">'
            '<path d="M4 19.5A2.5 2.5 0 016.5 17H20"/>'
            '<path d="M6.5 2H20v20H6.5A2.5 2.5 0 014 19.5v-15A2.5 2.5 0 016.5 2z"/>'
            '<line x1="8" y1="7" x2="16" y2="7"/>'
            '<line x1="8" y1="11" x2="16" y2="11"/>'
            '<line x1="8" y1="15" x2="12" y2="15"/>'
            '</svg>'
        ),
        "color": "#6366f1",
        "description": "KI-generiertes wissenschaftliches Paper als .docx exportieren.",
        "required_credentials": ["anthropic_api_key"],
    }

    def api_routes(self):
        return [{
            "rule": "/api/projects/<int:pid>/export-scientific-paper",
            "endpoint": "api_export_scientific_paper",
            "handler": self._handle,
            "methods": ["GET"],
        }]

    def _handle(self, pid):
        import json as _json
        from copy import deepcopy as _dc
        from flask import abort, jsonify, send_file
        from flask_login import current_user, login_required
        from models import Project, Snapshot, Slide
        from plugins.ai._llm import call_llm, get_ai_settings, increment_usage

        guard = self._guard("scientific_paper")
        if not isinstance(guard, tuple) or len(guard) != 2:
            return guard
        user_id, settings = guard

        proj = Project.query.filter_by(id=pid, user_id=user_id).first()
        if not proj:
            abort(404)

        snaps = Snapshot.query.filter_by(project_id=pid).order_by(
            Snapshot.sort_order, Snapshot.created_at).all()
        slides = Slide.query.filter_by(project_id=pid).order_by(
            Slide.sort_order, Slide.created_at).all()

        increment_usage(user_id, source="scientific_paper",
                        detail=proj.name or f"Projekt #{pid}")

        prompt = self._build_prompt(proj, snaps, slides)
        result = call_llm(prompt, settings, max_tokens=8000)

        if not result["ok"]:
            return jsonify({"error": result["error"]}), result.get("status", 502)

        paper_text = result["text"]

        # Titel und Abbildungstitel aus KI-Antwort extrahieren
        sci_title = proj.name
        fig_titles = {}
        body_text = paper_text.strip()
        if body_text.startswith("TITLE:"):
            nl_idx = body_text.find("\n")
            if nl_idx > 0:
                sci_title = body_text[:nl_idx][6:].strip()
                body_text = body_text[nl_idx:].strip()
        if body_text.startswith("FIGURE_TITLES:"):
            nl_idx = body_text.find("\n")
            if nl_idx > 0:
                ft_line = body_text[:nl_idx][14:].strip()
                body_text = body_text[nl_idx:].strip()
                for i, t in enumerate(ft_line.split("|"), start=1):
                    t = t.strip()
                    if t:
                        fig_titles[i] = t

        # Autor aus Einstellungen
        try:
            author_name = _json.loads(
                settings["resolve"]("pres_template", "{}")).get("author", "").strip()
        except Exception:
            author_name = ""

        buf = self._build_docx(sci_title, body_text, author_name,
                                snaps, fig_titles, _dc)
        safe_name = re.sub(r'[^\w\s\-]', '_', proj.name).strip()
        filename = f"scientific_paper_{safe_name}.docx"
        return send_file(
            buf, as_attachment=True, download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument."
                     "wordprocessingml.document",
        )

    # ── Prompt ─────────────────────────────────────────────────────────────

    @staticmethod
    def _build_prompt(proj, snaps, slides):
        import json as _json

        lines = [
            "You are an expert scientific writer specializing in digital behavioral analytics, "
            "computational social science, and information retrieval research. "
            "Based on the Google Trends research data provided below, write a complete, "
            "formally structured scientific research paper in English, meeting the publication "
            "standards of journals such as PLOS ONE, Nature Scientific Reports, "
            "or Social Science Computer Review.",
            "",
            "Strict output format — follow exactly:",
            "  Line 1:  TITLE: <a concise, formal academic paper title, 10–18 words>",
            "  Line 2:  FIGURE_TITLES: <English title for Fig. 1> | <English title for Fig. 2> | "
            "... (one entry per snapshot figure, in order; omit this line if there are no figures)",
            "  Line 3+: The paper body, beginning immediately with '## Abstract'.",
            "",
            "Rules:",
            "- Use ## for top-level section headings, ### for sub-sections only where warranted.",
            "- Write in formal, impersonal academic style (third-person, passive constructions "
            "  where appropriate, hedged claims, precise quantitative references to the data).",
            "- Reference specific numeric values, trend directions, and correlations from the data.",
            "- When referencing charts or data visualizations, always use 'Figure N' (e.g., Figure 1, Figure 2).",
            "- Do not include any preamble, author notes, commentary, or text outside this format.",
            "- Each section: 200–400 words.",
            "",
            "Required sections in this exact order:",
            "## Abstract",
            "## Keywords",
            "## 1. Introduction",
            "## 2. Data & Methodology",
            "## 3. Results",
            "## 4. Discussion",
            "## 5. Conclusion",
            "",
            "For '## Keywords': provide exactly 5–8 relevant academic keywords as a single",
            "comma-separated line (e.g.: search behavior, digital epidemiology, trend analysis).",
            "",
            f"PROJECT NAME: {proj.name}",
        ]
        if proj.briefing and proj.briefing.strip():
            lines += ["", "## Project Briefing / Research Question",
                       proj.briefing.strip(), ""]

        if snaps:
            lines += ["", f"## Snapshots ({len(snaps)} total)", ""]
            for i, snap in enumerate(snaps, 1):
                chart = _json.loads(snap.chart_json) if snap.chart_json else {}
                markers = _json.loads(snap.markers_json) if snap.markers_json else []
                lines.append(f"### Figure {i}: {snap.title or '(no title)'}")
                if snap.created_at:
                    lines.append(f"Date: {snap.created_at.strftime('%Y-%m-%d')}")
                if snap.comment and snap.comment.strip():
                    lines.append(f"Comment: {snap.comment.strip()}")

                datasets = chart.get("datasets", [])
                kw_meta = chart.get("keywords_meta", [])
                if kw_meta:
                    kws = ", ".join(
                        f"{m.get('keyword', '')} [{m.get('geo', '') or 'Global'}, "
                        f"{GP.get(m.get('gprop', ''), 'Web')}, "
                        f"{TF.get(m.get('timeframe', ''), m.get('timeframe', ''))}]"
                        for m in kw_meta
                    )
                    lines.append(f"Keywords: {kws}")

                for ds in datasets:
                    data = ds.get("data", [])
                    if not data:
                        continue
                    if isinstance(data[0], dict):
                        vals = [p["y"] for p in data if p.get("y") is not None]
                        d0 = str(data[0].get("x", ""))[:10]
                        d1 = str(data[-1].get("x", ""))[:10]
                    else:
                        vals = [v for v in data if v is not None]
                        labels = chart.get("labels", [])
                        d0 = str(labels[0])[:10] if labels else ""
                        d1 = str(labels[-1])[:10] if labels else ""
                    if vals:
                        avg = sum(vals) / len(vals)
                        third = max(1, len(vals) // 3)
                        diff = (sum(vals[-third:]) / third) - (sum(vals[:third]) / third)
                        trend = "rising" if diff > 5 else "falling" if diff < -5 else "stable"
                        lines.append(
                            f"  - {ds.get('label', '?')}: "
                            f"period {d0}–{d1}, "
                            f"min={min(vals):.0f} max={max(vals):.0f} avg={avg:.1f} trend={trend}")

                corrs = chart.get("correlations", [])
                for c in corrs:
                    r = c.get("r")
                    if r is not None:
                        lines.append(
                            f"  - Correlation {c.get('labelA', '?')} ↔ "
                            f"{c.get('labelB', '?')}: r={r:+.3f}")
                for m in markers:
                    lbl = m.get("label", "")
                    cmt = m.get("comment", "")
                    if lbl or cmt:
                        lines.append(f"  - Marker: {lbl}{' – ' + cmt if cmt else ''}")
                lines.append("")

        if slides:
            lines += ["## Additional Slides / Context", ""]
            for sl in slides:
                if sl.slide_type in ("title", "section", "textbild", "chronologie"):
                    lines.append(f"- [{sl.slide_type}] {sl.title or ''}: "
                                 f"{(sl.description or '')[:300]}")
            lines.append("")

        return "\n".join(lines)

    # ── DOCX-Generierung ───────────────────────────────────────────────────

    @staticmethod
    def _build_docx(sci_title, body_text, author_name, snaps, fig_titles, _dc):
        import json as _json

        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.dates as mdates
        from matplotlib.ticker import MaxNLocator

        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # ── Stile ──────────────────────────────────────────────────────────
        _apply_sci_styles(doc, qn, OxmlElement, Pt, RGBColor)

        # Seitenformat A4
        sec = doc.sections[-1]
        sec.page_width = Cm(21)
        sec.page_height = Cm(29.7)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)

        # 2-spaltig
        final_sp = sec._sectPr
        for old in final_sp.findall(qn('w:type')):
            final_sp.remove(old)
        type_el = OxmlElement('w:type')
        type_el.set(qn('w:val'), 'continuous')
        final_sp.insert(0, type_el)
        for old in final_sp.findall(qn('w:cols')):
            final_sp.remove(old)
        cols_el = OxmlElement('w:cols')
        cols_el.set(qn('w:num'), '2')
        cols_el.set(qn('w:space'), '720')
        final_sp.append(cols_el)

        # ── Titel (einspaltig) ─────────────────────────────────────────────
        title_p = doc.add_heading(sci_title, level=0)
        title_p.paragraph_format.page_break_before = False
        for run in title_p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(24)
            run.font.bold = False
            run.font.all_caps = False
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            rPr = run._r.get_or_add_rPr()
            for el in rPr.findall(qn('w:b')):
                rPr.remove(el)
            for el in rPr.findall(qn('w:bCs')):
                rPr.remove(el)
            b_el = OxmlElement('w:b')
            b_el.set(qn('w:val'), '0')
            rPr.append(b_el)
            b_cs = OxmlElement('w:bCs')
            b_cs.set(qn('w:val'), '0')
            rPr.append(b_cs)

        # Autorname
        if author_name:
            from docx.enum.text import WD_ALIGN_PARAGRAPH as _WDA
            ap = doc.add_paragraph()
            ap.alignment = _WDA.CENTER
            ap.paragraph_format.space_before = Pt(2)
            ap.paragraph_format.space_after = Pt(18)
            ap.paragraph_format.first_line_indent = Pt(0)
            ap.paragraph_format.page_break_before = False
            ar = ap.add_run(author_name)
            ar.font.name = 'Times New Roman'
            ar.font.size = Pt(9)
            ar.font.bold = False
            ar.font.italic = True
            ar.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        # Trennabsatz mit inline sectPr (1→2 Spalten)
        sep_p = doc.add_paragraph()
        sep_p.paragraph_format.space_before = Pt(0)
        sep_p.paragraph_format.space_after = Pt(0)
        sep_p.paragraph_format.first_line_indent = Pt(0)
        sep_p.paragraph_format.page_break_before = False
        _no_pbr_pPr = sep_p._p.get_or_add_pPr()
        sec1_pr = OxmlElement('w:sectPr')
        for _tag in (qn('w:pgSz'), qn('w:pgMar')):
            _el = final_sp.find(_tag)
            if _el is not None:
                sec1_pr.append(_dc(_el))
        _cols1 = OxmlElement('w:cols')
        _cols1.set(qn('w:space'), '720')
        sec1_pr.append(_cols1)
        _no_pbr_pPr.append(sec1_pr)

        # ── Markdown → Word ────────────────────────────────────────────────
        _SNAPREF_RE = re.compile(r'(Figure\s+\d+)', re.IGNORECASE)

        def _add_inline(para, text):
            _MD = re.compile(r'(\*\*[^*\n]+\*\*|\*[^*\n]+\*)')
            for md_part in _MD.split(text):
                if md_part.startswith("**") and md_part.endswith("**"):
                    r = para.add_run(md_part[2:-2])
                    r.bold = True
                    r.font.name = 'Times New Roman'
                elif md_part.startswith("*") and md_part.endswith("*"):
                    r = para.add_run(md_part[1:-1])
                    r.italic = True
                    r.font.name = 'Times New Roman'
                else:
                    for ref_part in _SNAPREF_RE.split(md_part):
                        r = para.add_run(ref_part)
                        r.font.name = 'Times New Roman'
                        if _SNAPREF_RE.fullmatch(ref_part):
                            r.bold = True

        _ROMAN = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII',
                   'IX', 'X', 'XI', 'XII', 'XIII', 'XIV', 'XV']
        _roman_ctr = [0]
        _ABC = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
        _abc_ctr = [0]

        def _no_pbr(p):
            pPr = p._p.get_or_add_pPr()
            for el in pPr.findall(qn('w:pageBreakBefore')):
                pPr.remove(el)
            pbr = OxmlElement('w:pageBreakBefore')
            pbr.set(qn('w:val'), '0')
            pPr.append(pbr)

        def _top_heading(text):
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            _abc_ctr[0] = 0
            n = _roman_ctr[0]
            _roman_ctr[0] += 1
            roman = _ROMAN[n] if n < len(_ROMAN) else str(n + 1)
            clean = re.sub(r'^\d+\.\s*', '', text).strip()
            p = doc.add_heading(f"{roman}.  {clean}", level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _no_pbr(p)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.bold = False
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            return p

        def _sub_heading(text):
            n = _abc_ctr[0]
            _abc_ctr[0] += 1
            letter = _ABC[n] if n < len(_ABC) else str(n + 1)
            clean = re.sub(r'^[A-Z]\.\s+', '', text).strip()
            clean = re.sub(r'^\d+\.\d+\s*', '', clean).strip()
            p = doc.add_heading(f"{letter}. {clean}", level=2)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.bold = False
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            pPr = p._p.get_or_add_pPr()
            for el in pPr.findall(qn('w:ind')):
                pPr.remove(el)
            ind = OxmlElement('w:ind')
            ind.set(qn('w:firstLine'), '0')
            pPr.append(ind)
            return p

        def _body_line(line):
            if not line.strip():
                p = doc.add_paragraph()
                _no_pbr(p)
                return
            if re.match(r"^[-*] ", line):
                p = doc.add_paragraph(style="List Bullet")
                _no_pbr(p)
                _add_inline(p, line[2:])
                return
            if re.match(r"^\d+\. ", line):
                p = doc.add_paragraph(style="List Number")
                _no_pbr(p)
                _add_inline(p, re.sub(r"^\d+\. ", "", line))
                return
            p = doc.add_paragraph()
            _no_pbr(p)
            _add_inline(p, line)

        # Sektionen parsen
        sections_order = ["## abstract", "## keywords", "## 1.", "## 2.",
                          "## 3.", "## 4.", "## 5."]
        section_chunks = {}
        current_key = None
        current_lines = []
        current_heading = ""

        for line in body_text.split("\n"):
            lower = line.strip().lower()
            matched = next((k for k in sections_order if lower.startswith(k)), None)
            if matched:
                if current_key:
                    section_chunks[current_key] = (current_heading, current_lines)
                current_key = matched
                current_heading = line
                current_lines = []
            else:
                if current_key:
                    current_lines.append(line)
        if current_key:
            section_chunks[current_key] = (current_heading, current_lines)

        if not section_chunks:
            for line in body_text.split("\n"):
                hm = re.match(r"^(#{1,4}) (.+)", line)
                if hm:
                    md_lvl = len(hm.group(1))
                    if md_lvl == 2:
                        _top_heading(hm.group(2))
                    elif md_lvl == 3:
                        _sub_heading(hm.group(2))
                    else:
                        hp = doc.add_heading(hm.group(2), level=md_lvl - 1)
                        for run in hp.runs:
                            run.font.bold = False
                    continue
                _body_line(line)
        else:
            def _write_section(heading, body_lines, inline_label=False):
                hm = re.match(r"^(#{1,4}) (.+)", heading)
                sec_name = hm.group(2).strip() if hm else heading.strip()

                if inline_label:
                    bt = " ".join(l.strip() for l in body_lines if l.strip())
                    is_kw = sec_name.lower() == "keywords"
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.first_line_indent = Pt(14)
                    _no_pbr(p)
                    r_lbl = p.add_run(f"{sec_name} \u2014 ")
                    r_lbl.bold = True
                    r_lbl.italic = True
                    r_lbl.font.name = 'Times New Roman'
                    r_body = p.add_run(bt)
                    r_body.bold = True
                    r_body.italic = is_kw
                    r_body.font.name = 'Times New Roman'
                    return

                if hm:
                    _top_heading(sec_name)
                for line in body_lines:
                    hm2 = re.match(r"^(#{1,4}) (.+)", line)
                    if hm2:
                        lvl = len(hm2.group(1))
                        if lvl == 3:
                            _sub_heading(hm2.group(2))
                        else:
                            hp = doc.add_heading(hm2.group(2), level=max(1, lvl - 1))
                            for run in hp.runs:
                                run.font.name = 'Times New Roman'
                                run.font.bold = False
                                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
                        continue
                    _body_line(line)

            for key in sections_order:
                if key not in section_chunks:
                    continue
                heading, body = section_chunks[key]
                _write_section(heading, body,
                               inline_label=(key in ("## abstract", "## keywords")))

                # Charts nach Results
                if key == "## 3." and snaps:
                    _top_heading("Figures")
                    for fig_n, snap in enumerate(snaps, 1):
                        chart = _json.loads(snap.chart_json) if snap.chart_json else {}
                        datasets = chart.get("datasets", [])
                        labels = chart.get("labels", [])
                        if not datasets:
                            continue
                        try:
                            w_in = 7 / 2.54
                            fig, ax = plt.subplots(figsize=(w_in, w_in / 2.5))
                            ax.set_facecolor("#ffffff")
                            fig.patch.set_facecolor("#ffffff")
                            for ds in datasets:
                                data = ds.get("data", [])
                                color = ds.get("borderColor", "#1a56db")
                                dlabel = ds.get("label", "")
                                if not data:
                                    continue
                                if isinstance(data[0], dict):
                                    xs = [p.get("x") for p in data]
                                    ys = [p.get("y") for p in data]
                                else:
                                    xs = labels[:len(data)]
                                    ys = data
                                try:
                                    from datetime import datetime as _dt
                                    xs_d = [_dt.fromisoformat(str(x)[:19]) for x in xs]
                                    ax.plot(xs_d, ys, linewidth=1.5, label=dlabel)
                                    ax.xaxis.set_major_formatter(
                                        mdates.DateFormatter("%d.%m.%y"))
                                    ax.xaxis.set_major_locator(
                                        mdates.AutoDateLocator())
                                except Exception:
                                    ax.plot(range(len(ys)), ys, linewidth=1.5, label=dlabel)
                            ax.set_ylim(0, 100)
                            ax.set_ylabel("Search Interest (0\u2013100)")
                            ax.tick_params(labelsize=7)
                            ax.yaxis.set_major_locator(MaxNLocator(integer=True, nbins=5))
                            fig.autofmt_xdate(rotation=35, ha="right")
                            fig.tight_layout(pad=0.5)
                            chart_buf = io.BytesIO()
                            fig.savefig(chart_buf, format="png", dpi=130,
                                        bbox_inches="tight", facecolor="white")
                            plt.close(fig)
                            chart_buf.seek(0)
                            doc.add_picture(chart_buf, width=Cm(7))
                        except Exception as e:
                            log.warning("Chart render failed for scientific paper: %s", e)
                        cap = doc.add_paragraph()
                        cap.paragraph_format.space_before = Pt(2)
                        cap.paragraph_format.space_after = Pt(8)
                        cap.paragraph_format.first_line_indent = Pt(0)
                        cr = cap.add_run(f"Fig. {fig_n}.")
                        cr.bold = True
                        cr.font.name = 'Times New Roman'
                        cr.font.size = Pt(8)
                        cap_title = fig_titles.get(fig_n, (snap.title or "").strip())
                        if cap_title:
                            ct = cap.add_run(f"  {cap_title}")
                            ct.bold = False
                            ct.font.name = 'Times New Roman'
                            ct.font.size = Pt(8)

        # Speichern
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf


# ── Styles-Helper (außerhalb der Klasse, einmal definiert) ─────────────────

def _apply_sci_styles(d, qn, OxmlElement, Pt, RGBColor):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    FONT, BLK = 'Times New Roman', '1A1A1A'

    def _set_font_xml(rPr, font_name):
        for el in rPr.findall(qn('w:rFonts')):
            rPr.remove(el)
        f = OxmlElement('w:rFonts')
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            f.set(qn(attr), font_name)
        rPr.insert(0, f)

    def _fix(name, sz, bold, caps, center=False, sb=0, sa=2, italic=False):
        try:
            sty = d.styles[name]
        except Exception:
            return
        sty.font.size = Pt(sz)
        sty.font.bold = bold
        sty.font.all_caps = caps
        sty.font.underline = False
        sty.font.italic = italic
        pf = sty.paragraph_format
        if center:
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = sty.element.get_or_add_pPr()
        for el in pPr.findall(qn('w:spacing')):
            pPr.remove(el)
        for el in pPr.findall(qn('w:pBdr')):
            pPr.remove(el)
        for el in pPr.findall(qn('w:numPr')):
            pPr.remove(el)
        pf.space_before = Pt(sb)
        pf.space_after = Pt(sa)
        pf.first_line_indent = Pt(0)
        pf.page_break_before = False
        rPr = sty.element.get_or_add_rPr()
        for el in rPr.findall(qn('w:b')):
            rPr.remove(el)
        for el in rPr.findall(qn('w:bCs')):
            rPr.remove(el)
        if not bold:
            b_el = OxmlElement('w:b')
            b_el.set(qn('w:val'), '0')
            rPr.append(b_el)
            b_cs = OxmlElement('w:bCs')
            b_cs.set(qn('w:val'), '0')
            rPr.append(b_cs)
        for el in rPr.findall(qn('w:i')):
            rPr.remove(el)
        for el in rPr.findall(qn('w:iCs')):
            rPr.remove(el)
        if italic:
            rPr.append(OxmlElement('w:i'))
            rPr.append(OxmlElement('w:iCs'))
        else:
            i_el = OxmlElement('w:i')
            i_el.set(qn('w:val'), '0')
            rPr.append(i_el)
            i_cs = OxmlElement('w:iCs')
            i_cs.set(qn('w:val'), '0')
            rPr.append(i_cs)
        _set_font_xml(rPr, FONT)
        for el in rPr.findall(qn('w:color')):
            rPr.remove(el)
        for el in rPr.findall(qn('w:u')):
            rPr.remove(el)
        c = OxmlElement('w:color')
        c.set(qn('w:val'), BLK)
        rPr.append(c)

    _fix('Title', 24, False, False, center=True, sb=0, sa=6)
    _fix('Heading 1', 12, False, True, center=True, sb=10, sa=2)
    _fix('Heading 2', 11, False, False, sb=6, sa=1, italic=True)
    _fix('Heading 3', 10, False, False, sb=4, sa=1)

    try:
        ns = d.styles['Normal']
        ns.font.size = Pt(10)
        rPr = ns.element.get_or_add_rPr()
        _set_font_xml(rPr, FONT)
        for el in rPr.findall(qn('w:color')):
            rPr.remove(el)
        c = OxmlElement('w:color')
        c.set(qn('w:val'), BLK)
        rPr.append(c)
        pPr = ns.element.get_or_add_pPr()
        for el in pPr.findall(qn('w:spacing')):
            pPr.remove(el)
        ns.paragraph_format.space_before = Pt(0)
        ns.paragraph_format.space_after = Pt(0)
        ns.paragraph_format.first_line_indent = Pt(14)
    except Exception:
        pass


PluginManager.register(ScientificPaperPlugin())
