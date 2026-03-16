"""
VeriTrend – Flask App
"""

import logging
import os
from datetime import datetime, timezone, timedelta

from dotenv import load_dotenv
load_dotenv()

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.ticker import MaxNLocator

from flask import Flask, jsonify, render_template, request, abort, send_file, redirect, url_for, Response
from auth import login_required, current_user, login_user, logout_user, superadmin_required, init_auth, MULTI_USER
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.cron import CronTrigger
from sqlalchemy import text

from models import db, login_manager, User, Keyword, TrendData, FetchLog, AppSetting, Event, RelatedQuery, Alert, AlertEvent, RegionInterest, Snapshot, Project, Slide, keyword_projects, WaitlistEntry

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# App-Konfiguration
# ---------------------------------------------------------------------------
BASE_DIR         = os.path.dirname(os.path.abspath(__file__))

app = Flask(__name__)
app.config["SQLALCHEMY_DATABASE_URI"] = (
    os.getenv("DATABASE_URL", f"sqlite:///{BASE_DIR}/trends.db")
)
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["SECRET_KEY"] = os.getenv("SECRET_KEY", "change-me-in-production")
# Templates immer neu laden (kein In-Memory-Cache)
app.config["TEMPLATES_AUTO_RELOAD"] = True
app.jinja_env.auto_reload = True

db.init_app(app)
login_manager.init_app(app)
init_auth(app)

# ── Enterprise-Modul (Benutzerverwaltung) — nur im Multi-User-Modus ─────
if MULTI_USER:
    from enterprise import enterprise_bp
    app.register_blueprint(enterprise_bp)
    # Öffentliche Seiten (Landing, FAQ, About-Us, Waitlist)
    from public_pages import public_bp
    app.register_blueprint(public_bp)

# ── Plugin-System initialisieren ──────────────────────────────────────────
from plugins import PluginManager
PluginManager.discover_all()

# Jinja-Loader um Plugin-Template-Ordner erweitern (gebündelte Plugins)
# Template-Pfade werden als "<plugin_id>/<datei>" aufgelöst,
# z.B. "vessel/_panel.html" → plugins/watchzone/vessel/templates/_panel.html
_plugin_tpl_dirs = PluginManager.template_dirs()
if _plugin_tpl_dirs:
    from jinja2 import ChoiceLoader, FileSystemLoader, PrefixLoader
    app.jinja_loader = ChoiceLoader([
        app.jinja_loader,                       # Standard: templates/
        PrefixLoader({                          # Plugin-Bundles: <id>/<file>
            pid: FileSystemLoader(tpl_dir)
            for pid, tpl_dir in _plugin_tpl_dirs.items()
        }, delimiter="/"),
    ])


# Route fuer statische Dateien gebündelter Plugins
# URL: /plugins/<type>/<id>/static/<filename>
@app.route("/plugins/<ptype>/<pid>/static/<path:filename>")
def plugin_static(ptype, pid, filename):
    import os as _os
    plugin = PluginManager.get(ptype, pid)
    if not plugin:
        abort(404)
    plugin_file = _os.path.abspath(
        __import__(plugin.__class__.__module__, fromlist=["__file__"]).__file__
    )
    static_dir = _os.path.join(_os.path.dirname(plugin_file), "static")
    from flask import send_from_directory
    return send_from_directory(static_dir, filename)


def _register_wz_history_routes():
    """Registriert History-API-Routen aus allen WatchZone-Plugins dynamisch."""
    from models import WatchZone

    def _make_history_view(handler):
        @login_required
        def _view(zid):
            z = WatchZone.query.filter_by(id=zid, user_id=current_user.id).first()
            if not z:
                abort(404)
            return handler(z, request.args, current_user.id)
        return _view

    for pid, plugin in PluginManager.all_of_type("watchzone").items():
        for route_def in plugin.history_routes():
            suffix = route_def["suffix"]
            endpoint = f"wz_{pid}_{suffix.replace('-', '_')}"
            app.add_url_rule(
                f"/api/watchzones/<int:zid>/{suffix}",
                endpoint=endpoint,
                view_func=_make_history_view(route_def["handler"]),
            )

_register_wz_history_routes()


def _register_wz_api_routes():
    """Registriert standalone API-Routen aus allen WatchZone-Plugins dynamisch."""
    for pid, plugin in PluginManager.all_of_type("watchzone").items():
        for route_def in plugin.api_routes():
            rule = route_def["rule"]
            endpoint = f"wz_{pid}_{rule.strip('/').replace('/', '_').replace('<', '').replace('>', '').replace(':', '_')}"
            app.add_url_rule(
                rule,
                endpoint=endpoint,
                view_func=login_required(route_def["handler"]),
                methods=route_def.get("methods", ["GET"]),
            )


_register_wz_api_routes()


def _register_analysis_routes():
    """Registriert API-Routen aus allen Analysis-Plugins dynamisch."""
    for pid, plugin in PluginManager.all_of_type("analysis").items():
        for route_def in plugin.api_routes():
            rule = route_def["rule"]
            endpoint = f"an_{pid}_{rule.strip('/').replace('/', '_').replace('<', '').replace('>', '').replace(':', '_')}"
            app.add_url_rule(
                rule,
                endpoint=endpoint,
                view_func=login_required(route_def["handler"]),
                methods=route_def.get("methods", ["GET"]),
            )


_register_analysis_routes()


def _register_ai_routes():
    """Registriert API-Routen aus allen AI-Plugins dynamisch."""
    for pid, plugin in PluginManager.all_of_type("ai").items():
        for route_def in plugin.api_routes():
            rule = route_def["rule"]
            endpoint = route_def.get("endpoint", f"ai_{pid}_{rule.strip('/').replace('/', '_')}")
            app.add_url_rule(
                rule,
                endpoint=endpoint,
                view_func=login_required(route_def["handler"]),
                methods=route_def.get("methods", ["GET"]),
            )


_register_ai_routes()


@app.context_processor
def inject_multi_user():
    """MULTI_USER-Flag fuer Templates bereitstellen."""
    return {"MULTI_USER": MULTI_USER}


@app.context_processor
def inject_site_logo():
    """Logo aus pres_template-Setting für alle Templates bereitstellen."""
    try:
        uid = current_user.id if current_user.is_authenticated else None
        setting = AppSetting.query.filter_by(key="pres_template", user_id=uid).first()
        if not setting and uid:
            setting = AppSetting.query.filter_by(key="pres_template", user_id=None).first()
        if setting and setting.value:
            import json as _json
            tmpl = _json.loads(setting.value)
            logo = tmpl.get("logo_url", "")
            if logo:
                return {"site_logo_url": logo}
    except Exception:
        pass
    return {"site_logo_url": ""}

@app.context_processor
def inject_current_year():
    from datetime import datetime
    return {"current_year": datetime.now().year}

@app.context_processor
def inject_user_project_count():
    try:
        if current_user.is_authenticated:
            count = Project.query.filter_by(user_id=current_user.id).count()
            return {"user_project_count": count}
    except Exception:
        pass
    return {"user_project_count": 0}


# ---------------------------------------------------------------------------
# Audit Trail
# ---------------------------------------------------------------------------
def _audit_enabled():
    """Prüft ob der Audit Trail aktiviert ist (globale Einstellung)."""
    s = AppSetting.query.filter_by(key="audit_trail_enabled", user_id=None).first()
    return s and s.value == "1"

def audit_log(action, object_type="", object_id=None, detail="",
              content_hash="", project_id=None, user_id=None):
    """Schreibt einen verketteten Audit-Trail-Eintrag (nur wenn Feature aktiviert)."""
    if not _audit_enabled():
        return None
    from models import AuditEntry
    if user_id is None:
        try:
            user_id = current_user.id if current_user.is_authenticated else None
        except Exception:
            user_id = None

    # Vorgänger-Hash für Verkettung
    prev = AuditEntry.query.order_by(AuditEntry.id.desc()).first()
    prev_hash = prev.entry_hash if prev else "GENESIS"

    entry = AuditEntry(
        user_id=user_id,
        project_id=project_id,
        action=action,
        object_type=object_type,
        object_id=object_id,
        detail=detail,
        content_hash=content_hash,
        prev_hash=prev_hash,
    )
    entry.compute_entry_hash()
    db.session.add(entry)
    db.session.commit()
    return entry

@app.context_processor
def inject_audit_trail_enabled():
    try:
        if current_user.is_authenticated:
            return {"audit_trail_enabled": _audit_enabled()}
    except Exception:
        pass
    return {"audit_trail_enabled": False}


def _ai_plugin_enabled(plugin_id):
    """Prüft ob ein AI-Plugin für den aktuellen User aktiviert ist."""
    uid = current_user.id if current_user and current_user.is_authenticated else None
    if not uid:
        return False
    enabled = PluginManager.enabled_for_user("ai", uid)
    return plugin_id in enabled


@app.context_processor
def inject_wz_plugins():
    """Stellt WatchZone-Plugins für Templates bereit (Sidebar, Panels, JS).

    Nur Plugins die sowohl aktiviert als auch verfügbar (Credentials vorhanden) sind.
    """
    plugins = []
    uid = current_user.id if current_user and current_user.is_authenticated else None
    src = PluginManager.enabled_for_user("watchzone", uid) if uid else PluginManager.all_of_type("watchzone")
    for pid, plugin in src.items():
        if not plugin.is_available(uid):
            continue
        plugins.append({
            "plugin_id": pid,
            "meta": plugin.meta,
        })
    return {"wz_plugins": plugins}


@app.context_processor
def inject_ai_plugins():
    """Stellt aktivierte AI-Plugin-IDs für Templates bereit."""
    uid = current_user.id if current_user and current_user.is_authenticated else None
    if uid:
        enabled = PluginManager.enabled_for_user("ai", uid)
    else:
        enabled = PluginManager.all_of_type("ai")
    return {"ai_plugins_enabled": set(enabled.keys())}


@app.context_processor
def inject_analysis_plugins():
    """Stellt Analysis-Plugin-Konfiguration für Templates bereit.

    Liefert analysis_in_lab und analysis_in_popup als Sets von Plugin-IDs,
    sowie analysis_plugins als Liste mit vollständigen Infos.
    """
    uid = current_user.id if current_user and current_user.is_authenticated else None
    enabled = PluginManager.enabled_for_user("analysis", uid) if uid else PluginManager.all_of_type("analysis")
    in_lab = set()
    in_popup = set()
    plugins_list = []
    for pid, plugin in enabled.items():
        show_in = plugin.get_show_in(uid) if uid else list(plugin.meta.get("default_show_in", []))
        if "lab" in show_in:
            in_lab.add(pid)
        if "popup" in show_in:
            in_popup.add(pid)
        plugins_list.append({
            "plugin_id": pid,
            "meta": plugin.meta,
            "show_in": show_in,
        })
    return {
        "analysis_in_lab": in_lab,
        "analysis_in_popup": in_popup,
        "analysis_plugins": plugins_list,
    }


@app.context_processor
def inject_plugin_i18n():
    """Sammelt i18n-Dicts aller Plugins für das Frontend."""
    import json as _json
    merged = {}  # {lang: {key: value}}
    # Seiten-übergreifende UI-Keys (Plugin-Verwaltung etc.)
    _UI_I18N = {
        "de": {
            "plugins_title": "Plugins",
            "plugins_desc": "Verwalte, welche Plugins f\u00fcr dein Konto aktiv sind. Deaktivierte Plugins erscheinen nicht in der Sidebar, im KI-Assistenten oder in der Analyse.",
            "plugins_enabled": "Aktiv", "plugins_disabled": "Inaktiv", "plugins_show_in": "Anzeige:",
        },
        "en": {
            "plugins_title": "Plugins",
            "plugins_desc": "Manage which plugins are active for your account. Disabled plugins will not appear in the sidebar, AI assistant, or analysis.",
            "plugins_enabled": "Enabled", "plugins_disabled": "Disabled", "plugins_show_in": "Show in:",
        },
        "fr": {
            "plugins_title": "Plugins",
            "plugins_desc": "G\u00e9rez les plugins actifs pour votre compte. Les plugins d\u00e9sactiv\u00e9s n\u2019appara\u00eetront pas dans la barre lat\u00e9rale, l\u2019assistant IA ou l\u2019analyse.",
            "plugins_enabled": "Actif", "plugins_disabled": "Inactif", "plugins_show_in": "Afficher dans\u00a0:",
        },
        "es": {
            "plugins_title": "Plugins",
            "plugins_desc": "Administra qu\u00e9 plugins est\u00e1n activos para tu cuenta. Los plugins desactivados no aparecer\u00e1n en la barra lateral, el asistente IA ni el an\u00e1lisis.",
            "plugins_enabled": "Activo", "plugins_disabled": "Inactivo", "plugins_show_in": "Mostrar en:",
        },
    }
    for lang, keys in _UI_I18N.items():
        merged.setdefault(lang, {}).update(keys)
    # Core-i18n der Plugin-Typen (z.B. WZ-Subnav, Global-Panel)
    for ptype, handler_cls in PluginManager._type_handlers.items():
        core = getattr(handler_cls, "_core_i18n", None)
        if core:
            for lang, keys in core.items():
                merged.setdefault(lang, {}).update(keys)
    for ptype in PluginManager._plugins:
        for pid, plugin in PluginManager.all_of_type(ptype).items():
            i18n = (plugin.meta or {}).get("i18n", {})
            for lang, keys in i18n.items():
                merged.setdefault(lang, {}).update(keys)
    return {"plugin_i18n_json": _json.dumps(merged, ensure_ascii=False)}


# ---------------------------------------------------------------------------
# Scheduler
# ---------------------------------------------------------------------------
scheduler = BackgroundScheduler(daemon=True)


def scheduled_fetch():
    from fetcher import run_fetch
    log.info("Scheduler: starte täglichen Trends-Abruf")
    try:
        result = run_fetch(app)
        log.info("Scheduler-Ergebnis: %s", result)
    except Exception as e:
        log.error("Scheduler-Fehler: %s", e)


# Täglich um 06:00 Uhr (Serverzeit)
FETCH_HOUR = int(os.getenv("FETCH_HOUR", "6"))
FETCH_MINUTE = int(os.getenv("FETCH_MINUTE", "0"))

scheduler.add_job(
    scheduled_fetch,
    CronTrigger(hour=FETCH_HOUR, minute=FETCH_MINUTE),
    id="daily_fetch",
    replace_existing=True,
)



# Auth-Routen + superadmin_required → enterprise.py / auth.py

# Community-Modus: /login leitet direkt auf /start weiter
if not MULTI_USER:
    @app.route("/login")
    def login():
        return redirect("/start")


# ---------------------------------------------------------------------------
# API-Richtlinien – pro API definierbar durch Superadmin
# ---------------------------------------------------------------------------

# Mapping: AppSetting-Schlüssel → API-Gruppenname
# (Plugin-spezifische Keys wie bluesky, telegram, copernicus, censys, aishub
#  werden über das Plugin-Credential-System verwaltet, nicht hier.)
_KEY_TO_API_GROUP = {
    "serpapi_key":          "serpapi",
    "newsapi_key":          "newsapi",
    "anthropic_api_key":    "anthropic",
    "openai_api_key":       "openai",
    "gemini_api_key":       "gemini",
    "mistral_api_key":      "mistral",
}

# Alle bekannten API-Gruppen (für Richtlinien-Verwaltung)
_ALL_API_GROUPS = [
    "serpapi", "newsapi", "anthropic", "openai", "gemini", "mistral",
]

_VALID_API_POLICIES = {"own_key", "admin_key", "disabled"}


def _get_api_policy(group_name):
    """Liest die globale Richtlinie für eine API-Gruppe: own_key | admin_key | disabled.
    Standard ist 'admin_key' (Nutzer verwenden den Admin-Schlüssel).
    """
    obj = AppSetting.query.filter_by(key=f"api_policy_{group_name}", user_id=None).first()
    val = obj.value if obj and obj.value else ""
    return val if val in _VALID_API_POLICIES else "admin_key"


def _get_user_api_perm(user_id, group_name):
    """Gibt die effektive API-Berechtigung für einen konkreten User zurück.
    Prüft zuerst die nutzer-spezifische Einstellung, dann die globale Policy.
    """
    if user_id:
        obj = AppSetting.query.filter_by(
            key=f"user_api_perm_{group_name}", user_id=user_id
        ).first()
        if obj and obj.value in _VALID_API_POLICIES:
            return obj.value
    return _get_api_policy(group_name)


# _save_user_api_perms, _load_user_api_perms, _load_plugin_access,
# _user_uses_own_llm_key → enterprise.py
# Fallback für Community-Modus (Single-User, kein enterprise.py)
if '_user_uses_own_llm_key' not in dir():
    def _user_uses_own_llm_key(user_id):
        return False


def _resolve_api_key(key, default="", user_id=None, is_admin=False):
    """Einheitliche Schlüssel-Auflösung mit per-User-Policy-Check.

    Priorität:
      1. Per-User-Richtlinie (user_api_perm_{group}) – fällt auf globale Policy zurück
      2. Richtlinien:
         - admin_key : immer den globalen Admin-Schlüssel verwenden
         - own_key   : nutzer-eigenen Schlüssel, kein Fallback auf Admin
         - disabled  : API nicht erlaubt → default zurückgeben
    """
    if is_admin:
        obj = AppSetting.query.filter_by(key=key, user_id=None).first()
        return obj.value if obj and obj.value else default

    group = _KEY_TO_API_GROUP.get(key)
    if group:
        policy = _get_user_api_perm(user_id, group) if user_id else _get_api_policy(group)
        if policy == "disabled":
            return default
        if policy == "own_key":
            if user_id:
                obj = AppSetting.query.filter_by(key=key, user_id=user_id).first()
                if obj and obj.value:
                    return obj.value
            return default  # own_key aber kein eigener Schlüssel gesetzt
        # admin_key: globalen Schlüssel nutzen
        obj = AppSetting.query.filter_by(key=key, user_id=None).first()
        return obj.value if obj and obj.value else default

    # Schlüssel ohne Richtlinie (z.B. ai_provider, ai_model): bisheriges Verhalten
    if user_id:
        obj = AppSetting.query.filter_by(key=key, user_id=user_id).first()
        if obj and obj.value:
            return obj.value
    obj = AppSetting.query.filter_by(key=key, user_id=None).first()
    return obj.value if obj and obj.value else default


# Benutzerverwaltung (User-CRUD, /api/me) → enterprise.py

# ---------------------------------------------------------------------------
# Seiten-Routen
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    if current_user.is_authenticated:
        return redirect("/start")
    if MULTI_USER:
        return render_template("public_pages/landing.html")
    return redirect("/login")


@app.route("/keywords")
@login_required
def keywords_page():
    return render_template("keywords.html")


@app.route("/analysis")
@login_required
def analysis():
    import time as _time
    # Sammle analysis_provider-Daten aus enabled Plugins
    wz_analysis_types = set()
    enabled = PluginManager.enabled_for_user("watchzone", current_user.id)
    for pid, plugin in enabled.items():
        if not plugin.is_available(current_user.id):
            continue
        prov = plugin.analysis_provider()
        if prov:
            wz_analysis_types.update(prov.get("data_types", []))
    # Modal-Templates aus enabled Analysis-Plugins sammeln
    analysis_modal_templates = []
    an_enabled = PluginManager.enabled_for_user("analysis", current_user.id)
    for pid, plugin in an_enabled.items():
        tpl = (plugin.meta or {}).get("modal_template", "")
        if tpl:
            analysis_modal_templates.append(tpl)
    resp = app.make_response(render_template(
        "analysis.html",
        cache_ts=int(_time.time()),
        wz_analysis_types=wz_analysis_types,
        analysis_modal_templates=analysis_modal_templates,
    ))
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp


@app.route("/events")
@login_required
def events_page():
    return render_template("events.html")


@app.route("/events/find")
@login_required
def events_find_page():
    return render_template("events_find.html")


@app.route("/events/watchzones")
@login_required
def events_watchzones_page():
    return render_template("events_watchzones.html")


@app.route("/audit")
@login_required
def audit_page():
    if not _audit_enabled():
        return redirect("/admin")
    return render_template("audit.html")


@app.route("/statistik")
@login_required
def statistik_page():
    return render_template("statistik.html")


@app.route("/admin")
@login_required
def admin_page():
    return render_template("admin.html")

@app.route("/help")
@login_required
def help_page():
    lang = request.args.get("lang", "").strip().lower()
    if lang not in ("de", "en", "fr", "es"):
        lang = "de"
        gl = AppSetting.query.filter_by(key="ui_language", user_id=None).first()
        if gl and gl.value and gl.value in ("de", "en", "fr", "es"):
            lang = gl.value
        if current_user.is_authenticated:
            ul = AppSetting.query.filter_by(key="ui_language", user_id=current_user.id).first()
            if ul and ul.value and ul.value in ("de", "en", "fr", "es"):
                lang = ul.value
    return render_template(f"help_{lang}.html")


@app.route("/plugins")
@login_required
def plugins_page():
    return render_template("plugins.html")


@app.route("/api/plugins", methods=["GET"])
@login_required
def api_plugins_list():
    """Alle Plugins aller Typen mit enabled/disabled-Status pro User."""
    uid = current_user.id
    result = {}
    for ptype in PluginManager._plugins:
        items = []
        for pid, plugin in PluginManager.all_of_type(ptype).items():
            key = f"plugin_disabled_{ptype}_{pid}"
            row = AppSetting.query.filter_by(key=key, user_id=uid).first()
            disabled = row and row.value == "1"
            # Admin-Zugriffskontrolle
            acc_row = AppSetting.query.filter_by(
                key=f"plugin_access_{ptype}_{pid}", user_id=uid
            ).first()
            blocked = acc_row and acc_row.value == "blocked"
            meta = plugin.meta or {}
            available = plugin.is_available(uid)
            item = {
                "plugin_id": pid,
                "plugin_type": ptype,
                "enabled": (not disabled) and available and (not blocked),
                "blocked": blocked,
                "available": available,
                "meta": meta,
            }
            # API-Gruppen (fuer users.html Plugin-Grid)
            req_creds = meta.get("required_credentials", [])
            api_groups = list({_KEY_TO_API_GROUP[c] for c in req_creds if c in _KEY_TO_API_GROUP})
            if api_groups:
                item["api_groups"] = api_groups
            # Credential-Status + Policy (filled/empty, nie Werte)
            if req_creds:
                cred_status = {}
                cred_policy = {}
                all_admin = True
                for ckey in req_creds:
                    r = AppSetting.query.filter_by(key=ckey, user_id=uid).first()
                    cred_status[ckey] = bool(r and r.value)
                    grp = _KEY_TO_API_GROUP.get(ckey)
                    pol = _get_user_api_perm(uid, grp) if grp else "own_key"
                    cred_policy[ckey] = pol
                    if pol != "admin_key":
                        all_admin = False
                item["credential_status"] = cred_status
                item["credential_policy"] = cred_policy
                item["creds_admin"] = all_admin  # True = alles ueber Admin-Keys
            if ptype == "analysis" and hasattr(plugin, "get_show_in"):
                show_in = plugin.get_show_in(uid)
                item["show_in"] = show_in
            items.append(item)
        result[ptype] = items
    # AI-Settings (Provider/Modell) – einmal fuer alle AI-Plugins
    if "ai" in result:
        _ai_prov_row = AppSetting.query.filter_by(key="ai_provider", user_id=uid).first()
        _ai_mod_row  = AppSetting.query.filter_by(key="ai_model", user_id=uid).first()
        if not _ai_prov_row:
            _ai_prov_row = AppSetting.query.filter_by(key="ai_provider", user_id=None).first()
        if not _ai_mod_row:
            _ai_mod_row = AppSetting.query.filter_by(key="ai_model", user_id=None).first()
        result["_ai_settings"] = {
            "ai_provider": (_ai_prov_row.value if _ai_prov_row and _ai_prov_row.value else "anthropic"),
            "ai_model": (_ai_mod_row.value if _ai_mod_row and _ai_mod_row.value else "claude-haiku-4-5-20251001"),
        }
    result["_uses_own_llm"] = _user_uses_own_llm_key(uid)
    return jsonify(result)


@app.route("/api/plugins/<ptype>/<pid>/toggle", methods=["PUT"])
@login_required
def api_plugin_toggle(ptype, pid):
    """Aktiviert/deaktiviert ein Plugin für den aktuellen User."""
    plugin = PluginManager.get(ptype, pid)
    if not plugin:
        return jsonify({"error": "Plugin not found"}), 404
    key = f"plugin_disabled_{ptype}_{pid}"
    uid = current_user.id
    # Admin-Blockade pruefen
    acc = AppSetting.query.filter_by(key=f"plugin_access_{ptype}_{pid}", user_id=uid).first()
    if acc and acc.value == "blocked":
        return jsonify({"enabled": False, "error": "blocked"}), 403
    row = AppSetting.query.filter_by(key=key, user_id=uid).first()
    if row and row.value == "1":
        # Currently disabled → enable – nur wenn verfuegbar
        if not plugin.is_available(uid):
            return jsonify({"enabled": False, "error": "credentials_missing"}), 400
        db.session.delete(row)
        db.session.commit()
        return jsonify({"enabled": True})
    else:
        # Currently enabled → disable
        if not row:
            row = AppSetting(key=key, user_id=uid, value="1")
            db.session.add(row)
        else:
            row.value = "1"
        db.session.commit()
        return jsonify({"enabled": False})


@app.route("/api/plugins/<ptype>/<pid>/show-in", methods=["PUT"])
@login_required
def api_plugin_show_in(ptype, pid):
    """Setzt wo ein Analysis-Plugin angezeigt wird (lab/popup/both/none)."""
    plugin = PluginManager.get(ptype, pid)
    if not plugin:
        return jsonify({"error": "Plugin not found"}), 404
    if ptype != "analysis":
        return jsonify({"error": "show_in nur für Analysis-Plugins"}), 400
    data = request.get_json(force=True) or {}
    value = data.get("show_in", "both")
    if value not in ("lab", "popup", "both", "none"):
        return jsonify({"error": "Ungültiger Wert"}), 400
    key = f"plugin_show_in_analysis_{pid}"
    uid = current_user.id
    row = AppSetting.query.filter_by(key=key, user_id=uid).first()
    if not row:
        row = AppSetting(key=key, user_id=uid, value=value)
        db.session.add(row)
    else:
        row.value = value
    db.session.commit()
    return jsonify({"show_in": value})


@app.route("/api/plugins/<ptype>/<pid>/credentials", methods=["PUT"])
@login_required
def api_plugin_credentials(ptype, pid):
    """Speichert per-User Credentials fuer ein Plugin."""
    plugin = PluginManager.get(ptype, pid)
    if not plugin:
        return jsonify({"error": "Plugin not found"}), 404
    allowed = set((plugin.meta or {}).get("required_credentials", []))
    if not allowed:
        return jsonify({"error": "Plugin benoetigt keine Credentials"}), 400
    data = request.get_json(force=True) or {}
    creds = data.get("credentials", {})
    uid = current_user.id
    status = {}
    for ckey in allowed:
        if ckey not in creds:
            # Nicht mitgeschickt → unveraendert lassen
            r = AppSetting.query.filter_by(key=ckey, user_id=uid).first()
            status[ckey] = bool(r and r.value)
            continue
        val = (creds[ckey] or "").strip()
        row = AppSetting.query.filter_by(key=ckey, user_id=uid).first()
        if val:
            if not row:
                row = AppSetting(key=ckey, user_id=uid, value=val)
                db.session.add(row)
            else:
                row.value = val
            status[ckey] = True
        else:
            # Leer → loeschen
            if row:
                db.session.delete(row)
            status[ckey] = False
    db.session.commit()
    return jsonify({"credential_status": status})


@app.route("/api/plugins/ai-settings", methods=["PUT"])
@login_required
def api_plugin_ai_settings():
    """Speichert AI-Provider und -Modell fuer den aktuellen User."""
    data = request.get_json(force=True) or {}
    provider = data.get("ai_provider", "").strip()
    model = data.get("ai_model", "").strip()
    if provider not in ("anthropic", "openai", "gemini", "mistral"):
        return jsonify({"error": "Ungueltiger Provider"}), 400
    uid = current_user.id
    for k, v in [("ai_provider", provider), ("ai_model", model)]:
        row = AppSetting.query.filter_by(key=k, user_id=uid).first()
        if not row:
            row = AppSetting(key=k, user_id=uid, value=v)
            db.session.add(row)
        else:
            row.value = v
    db.session.commit()
    return jsonify({"ai_provider": provider, "ai_model": model})


    # /faq, /about-us, /api/waitlist → public_pages.py (nur MULTI_USER)


@app.route("/about")
@login_required
def about_page():
    return render_template("about.html")


@app.route("/alerts")
@login_required
def alerts_page():
    return render_template("alerts.html")


@app.route("/projects")
@login_required
def projects_page():
    return render_template("projects.html")


@app.route("/start")
@app.route("/dashboard")  # Rückwärtskompatibilität
@login_required
def dashboard_page():
    proj_count = Project.query.filter_by(user_id=current_user.id).count()
    max_proj = current_user.max_projects or 0
    return render_template("start.html", proj_count=proj_count, max_proj=max_proj)


# ---------------------------------------------------------------------------
# API – Admin-Einstellungen
# ---------------------------------------------------------------------------

import json as _json
from fetcher import DEFAULT_WORKFLOW as _DEFAULT_WORKFLOW


def _increment_llm_usage(user_id, source="", detail=""):
    """Inkrementiert den monatlichen LLM-Call-Zähler und loggt den Aufruf."""
    from models import LlmLog
    month_key = f"llm_calls_{datetime.now().strftime('%Y-%m')}"
    setting = AppSetting.query.filter_by(key=month_key, user_id=user_id).first()
    if setting:
        setting.value = str(int(setting.value or 0) + 1)
    else:
        db.session.add(AppSetting(key=month_key, value="1", user_id=user_id))
    db.session.add(LlmLog(user_id=user_id, source=source, detail=detail[:255]))
    db.session.commit()


def _get_llm_usage(user_id):
    """Gibt (verwendet, limit) für den aktuellen Monat zurück.
    Eigener LLM-Key → limit=0 (unbegrenzt)."""
    month_key = f"llm_calls_{datetime.now().strftime('%Y-%m')}"
    setting = AppSetting.query.filter_by(key=month_key, user_id=user_id).first()
    used = int(setting.value or 0) if setting else 0
    if _user_uses_own_llm_key(user_id):
        return used, 0  # 0 = unbegrenzt
    user = db.session.get(User, user_id)
    limit = user.max_llm_calls if user else 0
    return used, limit


def _check_llm_quota(user_id):
    """Prüft ob das LLM-Kontingent erschöpft ist. Gibt (ok, used, limit) zurück."""
    used, limit = _get_llm_usage(user_id)
    if limit > 0 and used >= limit:
        return False, used, limit
    return True, used, limit


def _get_query_language(user=None):
    """Ermittelt die Abfragesprache (hl) für Google-Trends-Aufrufe."""
    hl = "auto"
    ql_global = AppSetting.query.filter_by(key="query_language", user_id=None).first()
    if ql_global and ql_global.value:
        hl = ql_global.value
    if user and hasattr(user, "id"):
        ql_user = AppSetting.query.filter_by(key="query_language", user_id=user.id).first()
        if ql_user and ql_user.value:
            hl = ql_user.value
    # "auto" → Accept-Language Header auswerten, Fallback "de"
    if hl == "auto":
        from flask import request as _req
        accept = _req.accept_languages.best_match(
            ["de", "en", "fr", "es", "it", "pt", "nl", "pl", "tr", "ru", "ja", "ko", "zh", "ar"],
            default="de",
        )
        hl = accept
    return hl


def _get_user_workflow(user=None):
    """Lädt den globalen Workflow und filtert nach erlaubten Backends des Users."""
    w_setting = AppSetting.query.filter_by(key="fetch_workflow", user_id=None).first()
    try:
        wf = _json.loads(w_setting.value) if w_setting and w_setting.value else _DEFAULT_WORKFLOW
    except Exception:
        wf = _DEFAULT_WORKFLOW
    if user and not user.is_superadmin and user.can_custom_workflow:
        allowed = user.get_allowed_backends()
        if allowed:
            wf = [e for e in wf if e.get("backend") in allowed]
    return wf


@app.route("/api/admin/settings", methods=["GET"])
@login_required
def api_admin_settings_get():
    """Gibt alle Admin-Einstellungen zurück (DB mit Env-Fallback)."""
    def _get(key, default=""):
        obj = AppSetting.query.filter_by(key=key, user_id=None).first()
        return obj.value if obj and obj.value else default

    raw_workflow = _get("fetch_workflow", "")
    try:
        workflow = _json.loads(raw_workflow) if raw_workflow else _DEFAULT_WORKFLOW
    except (ValueError, _json.JSONDecodeError):
        workflow = _DEFAULT_WORKFLOW

    fetch_hour   = int(_get("fetch_hour",   str(FETCH_HOUR)))
    fetch_minute = int(_get("fetch_minute", str(FETCH_MINUTE)))

    is_sa = current_user.is_superadmin
    result = {"fetch_hour": fetch_hour, "fetch_minute": fetch_minute}

    if is_sa or current_user.can_custom_workflow:
        if not is_sa:
            # Nur erlaubte Backends liefern
            allowed = current_user.get_allowed_backends()
            workflow = [e for e in workflow if e.get("backend") in allowed]
        result["fetch_workflow"] = workflow

    if is_sa or current_user.can_use_own_apis:
        _uid = None if is_sa else current_user.id

        def _get_user(key, default=""):
            """Settings mit Policy-Check: Superadmin bekommt immer den globalen Wert.
            Non-SA-Nutzer bekommen nur den Wert, wenn die Richtlinie es erlaubt."""
            if is_sa:
                return _get(key, default)
            group = _KEY_TO_API_GROUP.get(key)
            if group:
                policy = _get_api_policy(group)
                if policy == "disabled":
                    return ""
                if policy == "own_key":
                    obj = AppSetting.query.filter_by(key=key, user_id=_uid).first()
                    return obj.value if obj and obj.value else ""
                # admin_key: Wert wird angezeigt, aber nicht editierbar
                return _get(key, default)
            # Schlüssel ohne Gruppe (ai_provider etc.)
            obj = AppSetting.query.filter_by(key=key, user_id=_uid).first()
            if obj and obj.value:
                return obj.value
            return _get(key, default)

        result["serpapi_key"]       = _get_user("serpapi_key",       os.getenv("SERPAPI_KEY", ""))
        result["cookies_file"]      = _get_user("cookies_file",      os.getenv("TRENDS_COOKIES_FILE", ""))
        result["ai_provider"]       = _get_user("ai_provider",       "anthropic")
        result["ai_model"]          = _get_user("ai_model",          "claude-haiku-4-5-20251001")
        result["anthropic_api_key"] = _get_user("anthropic_api_key", os.getenv("ANTHROPIC_API_KEY", ""))
        result["openai_api_key"]    = _get_user("openai_api_key",    "")
        result["gemini_api_key"]    = _get_user("gemini_api_key",    "")
        result["mistral_api_key"]   = _get_user("mistral_api_key",   "")
        result["newsapi_key"]       = _get_user("newsapi_key",       os.getenv("NEWSAPI_KEY", ""))
        result["bluesky_handle"]    = _get_user("bluesky_handle",    os.getenv("BLUESKY_HANDLE", ""))
        result["bluesky_app_password"] = _get_user("bluesky_app_password", os.getenv("BLUESKY_APP_PASSWORD", ""))
        result["telegram_api_id"]      = _get_user("telegram_api_id",      os.getenv("TELEGRAM_API_ID", ""))
        result["telegram_api_hash"]    = _get_user("telegram_api_hash",    os.getenv("TELEGRAM_API_HASH", ""))
        result["copernicus_email"]    = _get_user("copernicus_email",    os.getenv("COPERNICUS_EMAIL", ""))
        result["copernicus_password"] = _get_user("copernicus_password", os.getenv("COPERNICUS_PASSWORD", ""))
        result["censys_api_id"]       = _get_user("censys_api_id",       os.getenv("CENSYS_API_ID", ""))
        result["censys_api_secret"]   = _get_user("censys_api_secret",   os.getenv("CENSYS_API_SECRET", ""))

    # API-Richtlinien – nur für Superadmin (zum Bearbeiten), non-SA bekommt Übersicht
    if is_sa:
        result["api_policies"] = {g: _get_api_policy(g) for g in _ALL_API_GROUPS}
    else:
        # Non-SA: nur mitteilen, welche Gruppen als own_key konfiguriert sind
        result["api_policies"] = {
            g: _get_api_policy(g) for g in _ALL_API_GROUPS
            if _get_api_policy(g) in ("own_key", "admin_key", "disabled")
        }

    # E-Mail-Einstellungen nur für Superadmin
    if is_sa:
        result["smtp_host"] = _get("smtp_host", "")
        result["smtp_port"] = int(_get("smtp_port", "587"))
        result["smtp_user"] = _get("smtp_user", "")
        result["smtp_pass"] = _get("smtp_pass", "")
        result["smtp_from"] = _get("smtp_from", "")
        result["smtp_tls"]  = _get("smtp_tls",  "1")

    # Audit Trail (global)
    result["audit_trail_enabled"] = _get("audit_trail_enabled", "0")

    # Datums-/Zeitformat (global)
    result["date_format"] = _get("date_format",  "DD.MM.YY")
    result["time_format"] = _get("time_format",  "HH:mm")
    result["trends_tz"]   = _get("trends_tz",    "60")
    result["ui_language"]    = _get("ui_language",  "de")
    result["query_language"] = _get("query_language", "auto")
    result["accent1"]        = _get("accent1", "")
    result["accent2"]        = _get("accent2", "")

    return jsonify(result)


@app.route("/api/admin/settings", methods=["POST"])
@login_required
def api_admin_settings_set():
    """Speichert Admin-Einstellungen und passt ggf. den Scheduler-Job an."""
    data = request.get_json(force=True) or {}

    def _save(key, value):
        obj = AppSetting.query.filter_by(key=key, user_id=None).first()
        if obj:
            obj.value = value
        else:
            db.session.add(AppSetting(key=key, value=value, user_id=None))

    is_sa = current_user.is_superadmin

    if "fetch_workflow" in data:
        if not is_sa and not current_user.can_custom_workflow:
            return jsonify({"error": "Keine Berechtigung für Workflow-Änderungen"}), 403
        if not is_sa:
            allowed = current_user.get_allowed_backends()
            for entry in data["fetch_workflow"]:
                if entry.get("backend") not in allowed:
                    return jsonify({"error": f"Backend '{entry.get('backend')}' nicht erlaubt"}), 403
        _save("fetch_workflow", _json.dumps(data["fetch_workflow"]))

    api_keys = ["serpapi_key", "cookies_file", "ai_provider", "ai_model",
                 "anthropic_api_key", "openai_api_key", "gemini_api_key",
                 "mistral_api_key", "newsapi_key",
                 "bluesky_handle", "bluesky_app_password",
                 "telegram_api_id", "telegram_api_hash",
                 "copernicus_email", "copernicus_password",
                 "censys_api_id", "censys_api_secret"]
    has_api_fields = any(k in data for k in api_keys)
    if has_api_fields and not is_sa and not current_user.can_use_own_apis:
        return jsonify({"error": "Keine Berechtigung für API-Änderungen"}), 403

    for key in api_keys:
        if key in data:
            if is_sa:
                _save(key, str(data[key]))
            else:
                # Non-SA mit can_use_own_apis: nur wenn Richtlinie "own_key" erlaubt
                group = _KEY_TO_API_GROUP.get(key)
                if group and _get_api_policy(group) != "own_key":
                    continue  # Nicht erlaubt für diesen Nutzer
                obj = AppSetting.query.filter_by(key=key, user_id=current_user.id).first()
                if obj:
                    obj.value = str(data[key])
                else:
                    db.session.add(AppSetting(key=key, value=str(data[key]),
                                              user_id=current_user.id))

    # API-Richtlinien – nur Superadmin darf setzen
    if "api_policies" in data:
        if not is_sa:
            return jsonify({"error": "Keine Berechtigung für API-Richtlinien"}), 403
        for group_name, policy_value in data["api_policies"].items():
            if group_name in _ALL_API_GROUPS and policy_value in _VALID_API_POLICIES:
                _save(f"api_policy_{group_name}", policy_value)

    # Datums-/Zeitformat – nur Superadmin (systemweit)
    _VALID_DATE_FORMATS = {"DD.MM.YY", "DD.MM.YYYY", "MM/DD/YYYY", "YYYY-MM-DD", "DD/MM/YYYY"}
    _VALID_TIME_FORMATS = {"HH:mm", "HH:mm:ss", "hh:mm A"}
    if "date_format" in data and is_sa:
        if data["date_format"] in _VALID_DATE_FORMATS:
            _save("date_format", data["date_format"])
    if "time_format" in data and is_sa:
        if data["time_format"] in _VALID_TIME_FORMATS:
            _save("time_format", data["time_format"])
    if "trends_tz" in data and is_sa:
        try:
            tz_val = int(data["trends_tz"])
            if -720 <= tz_val <= 840:
                _save("trends_tz", str(tz_val))
        except (ValueError, TypeError):
            pass
    if "ui_language" in data and is_sa:
        if data["ui_language"] in {"de", "en", "fr", "es"}:
            _save("ui_language", data["ui_language"])
    if "query_language" in data and is_sa:
        allowed_ql = {"auto", "de", "en", "fr", "es", "it", "pt", "nl", "pl", "tr", "ru", "ja", "ko", "zh", "ar"}
        if data["query_language"] in allowed_ql:
            _save("query_language", data["query_language"])

    # Akzentfarben – nur Superadmin (systemweit)
    import re as _re_mod
    _HEX_RE = _re_mod.compile(r'^#[0-9a-fA-F]{6}$')
    for akey in ("accent1", "accent2"):
        if akey in data and is_sa:
            if _HEX_RE.match(str(data[akey])):
                _save(akey, str(data[akey]))

    # Audit Trail – alle Benutzer dürfen aktivieren/deaktivieren
    if "audit_trail_enabled" in data:
        _save("audit_trail_enabled", "1" if data["audit_trail_enabled"] else "0")

    # SMTP – nur Superadmin
    smtp_keys = ["smtp_host", "smtp_port", "smtp_user", "smtp_pass", "smtp_from", "smtp_tls"]
    has_smtp = any(k in data for k in smtp_keys)
    if has_smtp and not is_sa:
        return jsonify({"error": "Keine Berechtigung für E-Mail-Einstellungen"}), 403
    for key in smtp_keys:
        if key in data:
            _save(key, str(data[key]))

    reschedule = False
    if "fetch_hour" in data or "fetch_minute" in data:
        hour   = int(data.get("fetch_hour",   FETCH_HOUR))
        minute = int(data.get("fetch_minute", FETCH_MINUTE))
        _save("fetch_hour",   str(hour))
        _save("fetch_minute", str(minute))
        reschedule = True

    db.session.commit()

    if reschedule:
        scheduler.reschedule_job(
            "daily_fetch",
            trigger=CronTrigger(hour=hour, minute=minute),
        )
        log.info("Scheduler auf %02d:%02d umgestellt", hour, minute)

    return jsonify({"ok": True})


@app.route("/api/admin/test-email", methods=["POST"])
@login_required
@superadmin_required
def api_admin_test_email():
    """Sendet eine Testmail mit den gespeicherten SMTP-Einstellungen."""
    import smtplib
    from email.mime.text import MIMEText

    def _get(key, default=""):
        obj = AppSetting.query.filter_by(key=key, user_id=None).first()
        return obj.value if obj and obj.value else default

    host = _get("smtp_host")
    port = int(_get("smtp_port", "587"))
    user = _get("smtp_user")
    pw   = _get("smtp_pass")
    frm  = _get("smtp_from")
    tls  = _get("smtp_tls", "1") != "0"

    if not host or not user:
        return jsonify({"error": "SMTP-Server und Benutzername müssen konfiguriert sein"}), 400

    recipient = frm or user
    msg = MIMEText("Dies ist eine Testmail von VeriTrend.\n\nDie SMTP-Konfiguration funktioniert.", "plain", "utf-8")
    msg["Subject"] = "VeriTrend – SMTP-Test"
    msg["From"]    = frm or user
    msg["To"]      = recipient

    try:
        if port == 465:
            srv = smtplib.SMTP_SSL(host, port, timeout=15)
        else:
            srv = smtplib.SMTP(host, port, timeout=15)
            if tls:
                srv.starttls()
        srv.login(user, pw)
        srv.sendmail(frm or user, [recipient], msg.as_string())
        srv.quit()
        return jsonify({"ok": True, "message": f"Testmail an {recipient} gesendet!"})
    except Exception as e:
        return jsonify({"error": f"SMTP-Fehler: {e}"}), 500


# ---------------------------------------------------------------------------
# API – Keywords
# ---------------------------------------------------------------------------

@app.route("/api/keywords", methods=["GET"])
@login_required
def api_keywords_list():
    project_id = request.args.get("project_id", type=int)
    uid = current_user.id
    if project_id:
        keywords = (Keyword.query
                    .filter_by(user_id=uid)
                    .join(keyword_projects, Keyword.id == keyword_projects.c.keyword_id)
                    .filter(keyword_projects.c.project_id == project_id)
                    .order_by(Keyword.created_at.desc()).all())
    else:
        keywords = Keyword.query.filter_by(user_id=uid).order_by(Keyword.created_at.desc()).all()
    counts = dict(
        db.session.query(TrendData.keyword_id, db.func.count(TrendData.id))
        .group_by(TrendData.keyword_id)
        .all()
    )
    result = []
    for kw in keywords:
        d = kw.to_dict()
        d["data_count"] = counts.get(kw.id, 0)
        result.append(d)
    return jsonify(result)


@app.route("/api/keywords", methods=["POST"])
@login_required
def api_keywords_create():
    data = request.get_json(force=True) or {}
    keyword_text = (data.get("keyword") or "").strip()
    if not keyword_text:
        abort(400, "keyword darf nicht leer sein")
    if current_user.max_keywords and current_user.max_keywords > 0:
        count = Keyword.query.filter_by(user_id=current_user.id).count()
        if count >= current_user.max_keywords:
            return jsonify({"error": f"Keyword-Limit erreicht ({current_user.max_keywords})"}), 403

    timeframe_val = data.get("timeframe", "")
    geo_val       = (data.get("geo") if data.get("geo") is not None else "DE").upper()
    gprop_val     = (data.get("gprop") or "").strip()
    if Keyword.query.filter_by(keyword=keyword_text, timeframe=timeframe_val,
                               geo=geo_val, gprop=gprop_val, user_id=current_user.id).first():
        abort(409, "Keyword mit diesen Suchparametern existiert bereits")

    # Sprache erkennen und ggf. deutsche Übersetzung speichern
    detected_lang = ""
    translation_de = ""
    try:
        from translator import detect_language, get_german_translation
        detected_lang = detect_language(keyword_text)
        if detected_lang not in ("de", "unknown") and detected_lang:
            t = get_german_translation(keyword_text)
            if t and t.lower() != keyword_text.lower():
                translation_de = t
    except Exception as e:
        log.warning("Spracherkennung fehlgeschlagen: %s", e)

    kw = Keyword(
        keyword=keyword_text,
        geo=geo_val,
        note=data.get("note", ""),
        active=data.get("active", True),
        detected_lang=detected_lang,
        translation_de=translation_de,
        timeframe=data.get("timeframe", ""),
        gprop=data.get("gprop", ""),
        user_id=current_user.id,
    )
    db.session.add(kw)
    # Projekt-Zuweisung (Projekt-Modus)
    project_id = data.get("project_id")
    if project_id:
        proj = Project.query.filter_by(id=int(project_id), user_id=current_user.id).first()
        if proj:
            kw.kw_projects.append(proj)
    db.session.commit()
    return jsonify(kw.to_dict()), 201


@app.route("/api/keywords/<int:kid>", methods=["PUT"])
@login_required
def api_keywords_update(kid):
    kw = Keyword.query.filter_by(id=kid, user_id=current_user.id).first()
    if not kw:
        abort(404)
    data = request.get_json(force=True) or {}

    if "keyword" in data:
        kw.keyword = data["keyword"].strip()
    if "geo" in data:
        kw.geo = data["geo"].upper()
    if "note" in data:
        kw.note = data["note"]
    if "active" in data:
        kw.active = bool(data["active"])
    if "timeframe" in data:
        kw.timeframe = data["timeframe"] or ""
    if "gprop" in data:
        kw.gprop = data["gprop"] or ""

    db.session.commit()
    return jsonify(kw.to_dict())


@app.route("/api/keywords/<int:kid>", methods=["DELETE"])
@login_required
def api_keywords_delete(kid):
    kw = Keyword.query.filter_by(id=kid, user_id=current_user.id).first()
    if not kw:
        abort(404)
    db.session.delete(kw)
    db.session.commit()
    return jsonify({"ok": True})


@app.route("/api/keywords/<int:kid>/projects/<int:pid>", methods=["POST"])
@login_required
def api_keyword_project_assign(kid, pid):
    """Keyword einem Projekt zuweisen."""
    kw   = Keyword.query.filter_by(id=kid, user_id=current_user.id).first() or abort(404)
    proj = Project.query.filter_by(id=pid, user_id=current_user.id).first() or abort(404)
    if proj not in kw.kw_projects:
        kw.kw_projects.append(proj)
        db.session.commit()
    return jsonify({"ok": True, "project_ids": [p.id for p in kw.kw_projects]})


@app.route("/api/keywords/<int:kid>/projects/<int:pid>", methods=["DELETE"])
@login_required
def api_keyword_project_unassign(kid, pid):
    """Keyword aus einem Projekt entfernen."""
    kw   = Keyword.query.filter_by(id=kid, user_id=current_user.id).first() or abort(404)
    proj = Project.query.filter_by(id=pid, user_id=current_user.id).first() or abort(404)
    if proj in kw.kw_projects:
        kw.kw_projects.remove(proj)
        db.session.commit()
    return jsonify({"ok": True, "project_ids": [p.id for p in kw.kw_projects]})


# ---------------------------------------------------------------------------
# API – Related Queries
# ---------------------------------------------------------------------------

@app.route("/api/keywords/<int:kid>/related-queries", methods=["GET"])
@login_required
def api_related_queries(kid):
    """
    Gibt die neueste Charge verwandter Suchanfragen für ein Keyword zurück.
    Optional: ?history=1 liefert alle gespeicherten Batches.
    """
    kw = Keyword.query.filter_by(id=kid, user_id=current_user.id).first()
    if not kw:
        abort(404)

    history = request.args.get("history") == "1"

    if history:
        rows = (
            db.session.query(RelatedQuery)
            .filter_by(keyword_id=kid)
            .order_by(RelatedQuery.fetched_at.desc(), RelatedQuery.rank)
            .all()
        )
        return jsonify([r.to_dict() | {"query_type": r.query_type} for r in rows])

    # Neuesten Batch ermitteln: höchstes fetched_at
    latest = (
        db.session.query(RelatedQuery)
        .filter_by(keyword_id=kid)
        .order_by(RelatedQuery.fetched_at.desc())
        .first()
    )
    if not latest:
        return jsonify({"rising": [], "top": [], "fetched_at": None})

    latest_ts = latest.fetched_at

    rows = (
        db.session.query(RelatedQuery)
        .filter_by(keyword_id=kid)
        .filter(RelatedQuery.fetched_at == latest_ts)
        .order_by(RelatedQuery.query_type, RelatedQuery.rank)
        .all()
    )

    result = {"rising": [], "top": [], "fetched_at": latest_ts.isoformat()}
    for r in rows:
        result[r.query_type].append(r.to_dict())
    return jsonify(result)


# ---------------------------------------------------------------------------
# API – Regionsinteresse
# ---------------------------------------------------------------------------

REGION_CACHE_HOURS = 24

@app.route("/api/keywords/<int:kid>/region-interest", methods=["GET"])
@login_required
def api_region_interest(kid):
    """
    Gibt Interesse-nach-Region für ein Keyword zurück.
    Query-Parameter:
      resolution – "REGION" (Standard) oder "CITY"
      force      – "1" → Cache ignorieren, live abrufen
      history    – "1" → alle gespeicherten Batches (nach run_tag gruppiert) zurückgeben
    """
    from datetime import timedelta

    kw = Keyword.query.filter_by(id=kid, user_id=current_user.id).first()
    if not kw:
        abort(404)

    resolution = (request.args.get("resolution") or "REGION").upper()
    if resolution not in ("REGION", "CITY", "COUNTRY"):
        abort(400, "resolution muss REGION, CITY oder COUNTRY sein")
    # Weltweite Keywords brauchen COUNTRY statt REGION/CITY
    if kw.geo == "":
        resolution = "COUNTRY"
    force   = request.args.get("force")   == "1"
    history = request.args.get("history") == "1"

    # Alle gespeicherten Batches zurückgeben (für Zeitverlauf-Diagramm)
    if history:
        from sqlalchemy import distinct as sa_distinct
        batch_keys = (
            db.session.query(RegionInterest.run_tag, RegionInterest.fetched_at)
            .filter_by(keyword_id=kid, resolution=resolution)
            .distinct()
            .order_by(RegionInterest.fetched_at.asc())
            .all()
        )
        batches = []
        for rt, fts in batch_keys:
            rows = (
                db.session.query(RegionInterest)
                .filter_by(keyword_id=kid, resolution=resolution,
                           run_tag=rt, fetched_at=fts)
                .order_by(RegionInterest.value.desc())
                .all()
            )
            batches.append({
                "run_tag":    rt or "",
                "fetched_at": fts.isoformat(),
                "data":       [r.to_dict() for r in rows],
            })
        return jsonify({"resolution": resolution, "batches": batches})

    latest = (
        db.session.query(RegionInterest)
        .filter_by(keyword_id=kid, resolution=resolution)
        .order_by(RegionInterest.fetched_at.desc())
        .first()
    )

    now = datetime.now(timezone.utc)
    cache_fresh = (
        latest is not None
        and not force
        and (now - latest.fetched_at.replace(tzinfo=timezone.utc))
            < timedelta(hours=REGION_CACHE_HOURS)
    )

    if cache_fresh:
        rows = (
            db.session.query(RegionInterest)
            .filter_by(keyword_id=kid, resolution=resolution, fetched_at=latest.fetched_at)
            .order_by(RegionInterest.value.desc())
            .all()
        )
        return jsonify({
            "resolution": resolution,
            "fetched_at": latest.fetched_at.isoformat(),
            "from_cache": True,
            "data": [r.to_dict() for r in rows],
        })

    # Live-Abruf via pytrends
    try:
        from fetcher import _fetch_region_pytrends
        _tz_global = AppSetting.query.filter_by(key="trends_tz", user_id=None).first()
        _tz_user   = AppSetting.query.filter_by(key="trends_tz", user_id=current_user.id).first()
        _tz_offset = 60
        if _tz_global and _tz_global.value:
            try: _tz_offset = int(_tz_global.value)
            except ValueError: pass
        if _tz_user and _tz_user.value:
            try: _tz_offset = int(_tz_user.value)
            except ValueError: pass
        _hl = _get_query_language(current_user)
        items = _fetch_region_pytrends(
            kw.keyword,
            kw.geo,
            kw.timeframe or "today 12-m",
            kw.gprop or "",
            resolution,
            tz_offset=_tz_offset,
            hl=_hl,
        )
    except Exception as e:
        log.warning("region-interest Abruf fehlgeschlagen für kw_id=%d: %s", kid, e)
        if latest:
            rows = (
                db.session.query(RegionInterest)
                .filter_by(keyword_id=kid, resolution=resolution, fetched_at=latest.fetched_at)
                .order_by(RegionInterest.value.desc())
                .all()
            )
            return jsonify({
                "resolution": resolution,
                "fetched_at": latest.fetched_at.isoformat(),
                "from_cache": True,
                "stale": True,
                "error": str(e),
                "data": [r.to_dict() for r in rows],
            })
        return jsonify({
            "resolution": resolution,
            "fetched_at": None,
            "from_cache": False,
            "data": [],
            "error": str(e),
        })

    batch_ts = now
    try:
        for item in items:
            db.session.add(RegionInterest(
                keyword_id=kid,
                resolution=resolution,
                geo_name=item["geo_name"],
                geo_code=item.get("geo_code", ""),
                value=item["value"],
                fetched_at=batch_ts,
                run_tag="",
            ))
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        log.error("region-interest DB-Fehler: %s", e)

    return jsonify({
        "resolution": resolution,
        "fetched_at": batch_ts.isoformat(),
        "from_cache": False,
        "data": [
            {
                "geo_name":   it["geo_name"],
                "geo_code":   it.get("geo_code", ""),
                "value":      it["value"],
                "fetched_at": batch_ts.isoformat(),
                "run_tag":    "",
            }
            for it in items
        ],
    })


# ---------------------------------------------------------------------------
# API – Trends-Daten
# ---------------------------------------------------------------------------

@app.route("/api/trends")
@login_required
def api_trends():
    """
    Query-Parameter:
      ids     – kommagetrennte Keyword-IDs (erforderlich)
      from    – Startdatum/Datetime YYYY-MM-DD oder YYYY-MM-DD HH:MM (optional)
      to      – Enddatum/Datetime YYYY-MM-DD oder YYYY-MM-DD HH:MM (optional)
      run_tag – Abruf-Reihe; leer = Standard-Tagesabruf (optional)
    """
    ids_param = request.args.get("ids", "")
    if not ids_param:
        abort(400, "ids Parameter fehlt")

    try:
        ids = [int(i) for i in ids_param.split(",") if i.strip()]
    except ValueError:
        abort(400, "ids muss kommagetrennte Integer sein")

    run_tag   = request.args.get("run_tag", "")
    date_from = request.args.get("from")
    date_to   = request.args.get("to")

    # Nur eigene Keywords zulassen
    own_ids = {k.id for k in Keyword.query.filter(Keyword.id.in_(ids), Keyword.user_id == current_user.id).all()}

    result = {}
    for kid in ids:
        if kid not in own_ids:
            continue
        kw = Keyword.query.filter_by(id=kid, user_id=current_user.id).first()
        if not kw:
            continue

        q = TrendData.query.filter_by(keyword_id=kid, run_tag=run_tag)
        if date_from:
            q = q.filter(TrendData.date >= date_from)
        if date_to:
            q = q.filter(TrendData.date <= date_to)
        q = q.order_by(TrendData.date)

        result[str(kid)] = {
            "keyword": kw.keyword,
            "geo":     kw.geo,
            "run_tag": run_tag,
            "data":    [td.to_dict() for td in q.all()],
        }

    return jsonify(result)


@app.route("/api/wiki-views")
@login_required
def api_wiki_views():
    """Wikipedia-Pageviews für Artikel abrufen.
    Query-Parameter:
      articles – kommagetrennte Artikelnamen (erforderlich)
      lang     – Sprachcode, Standard 'de'
      days     – Anzahl Tage, Standard 365, max 730
    """
    import requests as _req_wiki
    from datetime import timedelta

    articles_param = request.args.get("articles", "")
    if not articles_param:
        abort(400, "articles Parameter fehlt")
    articles = [a.strip() for a in articles_param.split(",") if a.strip()][:5]
    lang = request.args.get("lang", "de")
    days = min(int(request.args.get("days", 365)), 730)

    end_dt = datetime.now()
    start_dt = end_dt - timedelta(days=days)
    start_str = start_dt.strftime("%Y%m%d")
    end_str = end_dt.strftime("%Y%m%d")

    UA = "VeriTrend.ai/1.0 (forensic trend analysis; contact@veritrend.ai)"
    result = []

    def _resolve_wiki_title(term, wlang):
        """Sucht den exakten Wikipedia-Artikeltitel über die Search-API."""
        try:
            search_url = (
                f"https://{wlang}.wikipedia.org/w/api.php?"
                f"action=query&list=search&srsearch={_req_wiki.utils.quote(term)}"
                f"&srlimit=1&format=json"
            )
            sr = _req_wiki.get(search_url, headers={"User-Agent": UA}, timeout=10)
            sr.raise_for_status()
            hits = sr.json().get("query", {}).get("search", [])
            if hits:
                return hits[0]["title"]
        except Exception:
            pass
        return None

    for article in articles:
        # Immer zuerst den korrekten Titel über die Search-API auflösen
        # (Wikimedia Pageviews API ist case-sensitive: "iran" ≠ "Iran")
        resolved = _resolve_wiki_title(article, lang)
        article_clean = (resolved or article).replace(" ", "_")
        url = (
            f"https://wikimedia.org/api/rest_v1/metrics/pageviews/per-article/"
            f"{lang}.wikipedia/all-access/all-agents/"
            f"{_req_wiki.utils.quote(article_clean, safe='')}/daily/{start_str}/{end_str}"
        )
        try:
            r = _req_wiki.get(url, headers={"User-Agent": UA}, timeout=15)
            if r.status_code == 404:
                result.append({"article": article, "error": "Artikel nicht gefunden"})
                continue
            r.raise_for_status()
            data = r.json()
            items = data.get("items", [])
            series = []
            for it in items:
                ts = it.get("timestamp", "")[:8]
                if len(ts) == 8:
                    series.append({
                        "date": f"{ts[:4]}-{ts[4:6]}-{ts[6:8]}",
                        "views": it.get("views", 0),
                    })
            wiki_title = article_clean.replace("_", " ")
            result.append({"article": article, "wiki_title": wiki_title, "lang": lang, "series": series})
        except Exception as exc:
            result.append({"article": article, "error": str(exc)[:120]})

    return jsonify(result)


@app.route("/api/wiki-edits")
@login_required
def api_wiki_edits():
    """Wikipedia-Artikelbearbeitungen (Revisions) pro Tag abrufen.
    Query-Parameter:
      articles – kommagetrennte Artikelnamen (erforderlich, max 5)
      lang     – Sprachcode, Standard 'de'
      days     – Anzahl Tage, Standard 365, max 730
    """
    import requests as _req_wiki
    from collections import Counter

    articles_param = request.args.get("articles", "")
    if not articles_param:
        abort(400, "articles Parameter fehlt")
    articles = [a.strip() for a in articles_param.split(",") if a.strip()][:5]
    lang = request.args.get("lang", "de")
    days = min(int(request.args.get("days", 365)), 730)

    end_dt = datetime.now(timezone.utc)
    start_dt = end_dt - timedelta(days=days)
    UA = "VeriTrend.ai/1.0 (forensic trend analysis; contact@veritrend.ai)"
    result = []

    def _resolve_wiki_title(term, wlang):
        try:
            search_url = (
                f"https://{wlang}.wikipedia.org/w/api.php?"
                f"action=query&list=search&srsearch={_req_wiki.utils.quote(term)}"
                f"&srlimit=1&format=json"
            )
            sr = _req_wiki.get(search_url, headers={"User-Agent": UA}, timeout=10)
            sr.raise_for_status()
            hits = sr.json().get("query", {}).get("search", [])
            if hits:
                return hits[0]["title"]
        except Exception:
            pass
        return None

    for article in articles:
        resolved = _resolve_wiki_title(article, lang)
        wiki_title = resolved or article

        # Revisions per MediaWiki API abrufen (paginiert)
        edits_per_day = Counter()
        rvcontinue = None
        total_revisions = 0
        rv_start = end_dt.strftime("%Y-%m-%dT%H:%M:%SZ")
        rv_end = start_dt.strftime("%Y-%m-%dT%H:%M:%SZ")

        try:
            for _ in range(50):  # Max 50 API-Seiten (à 500 Revisions = 25.000 Edits)
                params = {
                    "action": "query",
                    "prop": "revisions",
                    "titles": wiki_title,
                    "rvprop": "timestamp|size|user",
                    "rvlimit": "500",
                    "rvstart": rv_start,     # neueste zuerst
                    "rvend": rv_end,
                    "format": "json",
                }
                if rvcontinue:
                    params["rvcontinue"] = rvcontinue

                api_url = f"https://{lang}.wikipedia.org/w/api.php"
                r = _req_wiki.get(api_url, params=params,
                                  headers={"User-Agent": UA}, timeout=15)
                r.raise_for_status()
                data = r.json()

                pages = data.get("query", {}).get("pages", {})
                for page_id, page in pages.items():
                    if page_id == "-1":
                        break
                    for rev in page.get("revisions", []):
                        ts = rev.get("timestamp", "")[:10]  # "2026-03-09"
                        if ts:
                            edits_per_day[ts] += 1
                            total_revisions += 1

                # Pagination
                cont = data.get("continue", {})
                rvcontinue = cont.get("rvcontinue")
                if not rvcontinue:
                    break

            # In sortierte Serie umwandeln
            series = [{"date": d, "edits": edits_per_day[d]}
                      for d in sorted(edits_per_day.keys())]

            result.append({
                "article": article,
                "wiki_title": wiki_title,
                "lang": lang,
                "total_edits": total_revisions,
                "series": series,
            })
        except Exception as exc:
            result.append({"article": article, "error": str(exc)[:120]})

    return jsonify(result)


@app.route("/api/wiki-edit-details")
@login_required
def api_wiki_edit_details():
    """Wikipedia-Revisions für einen bestimmten Artikel an einem bestimmten Tag.
    Gibt User/IP, Zeitstempel und Kommentar zurück.
    Query-Parameter:
      article – Artikelname (erforderlich)
      date    – Datum im Format YYYY-MM-DD (erforderlich)
      lang    – Sprachcode, Standard 'de'
    """
    import requests as _req_wiki
    import re as _re_ip

    article = request.args.get("article", "").strip()
    date_str = request.args.get("date", "").strip()
    lang = request.args.get("lang", "de")

    if not article or not date_str:
        return jsonify(ok=False, error="article und date Parameter erforderlich."), 400

    UA = "VeriTrend.ai/1.0 (forensic trend analysis; contact@veritrend.ai)"
    api_url = f"https://{lang}.wikipedia.org/w/api.php"

    # Tagesanfang und -ende als ISO-Timestamps
    rv_start = f"{date_str}T23:59:59Z"
    rv_end = f"{date_str}T00:00:00Z"

    revisions = []
    rvcontinue = None

    try:
        for _ in range(10):  # Max 10 Seiten (à 500 = 5.000 Revisions pro Tag)
            params = {
                "action": "query",
                "prop": "revisions",
                "titles": article,
                "rvprop": "timestamp|user|comment|size|ids",
                "rvlimit": "500",
                "rvstart": rv_start,
                "rvend": rv_end,
                "format": "json",
            }
            if rvcontinue:
                params["rvcontinue"] = rvcontinue

            r = _req_wiki.get(api_url, params=params,
                              headers={"User-Agent": UA}, timeout=15)
            r.raise_for_status()
            data = r.json()

            pages = data.get("query", {}).get("pages", {})
            for page_id, page in pages.items():
                if page_id == "-1":
                    return jsonify(ok=False, error=f"Artikel '{article}' nicht gefunden."), 404
                for rev in page.get("revisions", []):
                    user = rev.get("user", "")
                    # Prüfe ob der Username eine IP-Adresse ist (IPv4 oder IPv6)
                    is_ip = bool(_re_ip.match(
                        r"^(\d{1,3}\.){3}\d{1,3}$|^[0-9a-fA-F:]+$", user
                    )) if user else False
                    revisions.append({
                        "revid": rev.get("revid"),
                        "timestamp": rev.get("timestamp", ""),
                        "user": user,
                        "is_ip": is_ip,
                        "comment": rev.get("comment", ""),
                        "size": rev.get("size"),
                    })

            cont = data.get("continue", {})
            rvcontinue = cont.get("rvcontinue")
            if not rvcontinue:
                break

    except Exception as exc:
        return jsonify(ok=False, error=str(exc)[:200]), 500

    # ── Registrierte Benutzer: Reputation via MediaWiki API ──
    reg_users = list({r["user"] for r in revisions if not r["is_ip"] and r["user"]})
    user_info_map = {}
    if reg_users:
        from datetime import datetime as _dt_rep
        # MediaWiki list=users – max 50 pro Request
        for batch_start in range(0, len(reg_users), 50):
            batch = reg_users[batch_start:batch_start + 50]
            try:
                ui_resp = _req_wiki.get(api_url, params={
                    "action": "query",
                    "list": "users",
                    "ususers": "|".join(batch),
                    "usprop": "editcount|registration|groups|blockinfo",
                    "format": "json",
                }, headers={"User-Agent": UA}, timeout=15)
                if ui_resp.ok:
                    for u in ui_resp.json().get("query", {}).get("users", []):
                        name = u.get("name", "")
                        editcount = u.get("editcount", 0)
                        reg_date = u.get("registration", "")
                        groups = [g for g in u.get("groups", []) if g not in ("*", "user", "autoconfirmed")]
                        blocked = "blockid" in u

                        # Account-Alter in Tagen
                        age_days = None
                        if reg_date:
                            try:
                                rd = _dt_rep.strptime(reg_date, "%Y-%m-%dT%H:%M:%SZ")
                                age_days = (_dt_rep.utcnow() - rd).days
                            except Exception:
                                pass

                        # Reputations-Score (0-100)
                        score = 0
                        if editcount >= 100000: score += 40
                        elif editcount >= 10000: score += 35
                        elif editcount >= 1000: score += 28
                        elif editcount >= 100: score += 18
                        elif editcount >= 10: score += 8
                        elif editcount >= 1: score += 3

                        if age_days is not None:
                            if age_days >= 3650: score += 30
                            elif age_days >= 1825: score += 25
                            elif age_days >= 365: score += 18
                            elif age_days >= 90: score += 10
                            elif age_days >= 30: score += 5

                        if "sysop" in groups or "bureaucrat" in groups: score += 20
                        elif "reviewer" in groups or "editor" in groups: score += 10
                        elif "patroller" in groups or "rollbacker" in groups: score += 8
                        if blocked: score = max(0, score - 30)

                        score = min(100, score)

                        # Label
                        if score >= 70: rep_label = "sehr hoch"
                        elif score >= 50: rep_label = "hoch"
                        elif score >= 30: rep_label = "mittel"
                        elif score >= 15: rep_label = "niedrig"
                        else: rep_label = "sehr niedrig"

                        user_info_map[name] = {
                            "editcount": editcount,
                            "registration": reg_date,
                            "age_days": age_days,
                            "groups": groups,
                            "blocked": blocked,
                            "reputation_score": score,
                            "reputation_label": rep_label,
                        }
            except Exception:
                pass

    # ── IP-Adressen mit GeoIP auflösen (ip-api.com, kostenlos, max 45/min) ──
    ip_users = list({r["user"] for r in revisions if r["is_ip"]})
    geo_map = {}
    if ip_users:
        for batch_start in range(0, len(ip_users), 100):
            batch = ip_users[batch_start:batch_start + 100]
            try:
                geo_resp = _req_wiki.post(
                    "http://ip-api.com/batch?fields=query,status,country,regionName,city,org,isp",
                    json=[{"query": ip} for ip in batch],
                    timeout=10
                )
                if geo_resp.ok:
                    for entry in geo_resp.json():
                        if entry.get("status") == "success":
                            geo_map[entry["query"]] = {
                                "country": entry.get("country", ""),
                                "region": entry.get("regionName", ""),
                                "city": entry.get("city", ""),
                                "org": entry.get("org", ""),
                                "isp": entry.get("isp", ""),
                            }
            except Exception:
                pass

    # Daten an Revisions anfügen
    for rev in revisions:
        if rev["is_ip"] and rev["user"] in geo_map:
            rev["geo"] = geo_map[rev["user"]]
        elif not rev["is_ip"] and rev["user"] in user_info_map:
            rev["user_info"] = user_info_map[rev["user"]]

    # Zusammenfassung: IP-Herkunft aggregieren
    from collections import Counter
    ip_summary = Counter()
    for rev in revisions:
        if rev.get("geo"):
            loc = rev["geo"]
            key = f"{loc['city']}, {loc['region']}, {loc['country']}" if loc['city'] else loc['country']
            ip_summary[key] += 1

    return jsonify(
        ok=True,
        article=article,
        date=date_str,
        total=len(revisions),
        anonymous=sum(1 for r in revisions if r["is_ip"]),
        registered=sum(1 for r in revisions if not r["is_ip"]),
        ip_origins=sorted(ip_summary.items(), key=lambda x: -x[1]),
        user_profiles=user_info_map,
        revisions=revisions,
    )


@app.route("/api/yahoo-search")
@login_required
def api_yahoo_search():
    """Ticker-Suche via Yahoo Finance Autocomplete."""
    import requests as req_lib
    q = request.args.get("q", "").strip()
    if not q or len(q) < 1:
        return jsonify([])
    try:
        r = req_lib.get(
            "https://query2.finance.yahoo.com/v1/finance/search",
            params={"q": q, "quotesCount": 8, "newsCount": 0, "listsCount": 0},
            headers={"User-Agent": "Mozilla/5.0"},
            timeout=5,
        )
        r.raise_for_status()
        data = r.json()
        results = []
        for item in data.get("quotes", []):
            symbol = item.get("symbol", "")
            name = item.get("shortname") or item.get("longname") or ""
            exchange = item.get("exchDisp") or item.get("exchange") or ""
            qtype = item.get("quoteType", "")
            if not symbol:
                continue
            results.append({
                "symbol": symbol,
                "name": name,
                "exchange": exchange,
                "type": qtype,
            })
        return jsonify(results)
    except Exception as exc:
        log.warning("Yahoo Search fehlgeschlagen: %s", exc)
        return jsonify([])


@app.route("/api/yahoo-finance")
@login_required
def api_yahoo_finance():
    """Liefert historische Kursdaten und Volumen für einen oder mehrere Ticker."""
    import yfinance as yf
    tickers = request.args.get("tickers", "").strip()
    days = min(730, max(7, int(request.args.get("days", 365))))
    labels_str = request.args.get("labels", "")

    if not tickers:
        return jsonify([])

    label_list = [l.strip() for l in labels_str.split(",") if l.strip()] if labels_str else []
    ticker_list = [t.strip().upper() for t in tickers.split(",") if t.strip()]

    from datetime import datetime as _dt, timedelta as _td
    end = _dt.now()
    start = end - _td(days=days)

    results = []
    for symbol in ticker_list[:5]:  # max 5 Ticker
        try:
            tk = yf.Ticker(symbol)
            hist = tk.history(start=start.strftime("%Y-%m-%d"),
                              end=end.strftime("%Y-%m-%d"))
            if hist.empty:
                results.append({"ticker": symbol, "error": "Keine Daten"})
                continue

            info = tk.info or {}
            name = info.get("shortName") or info.get("longName") or symbol
            currency = info.get("currency", "")

            price_series = []
            volume_series = []
            for idx, row in hist.iterrows():
                d = idx.strftime("%Y-%m-%d")
                price_series.append({"date": d, "close": round(float(row["Close"]), 2)})
                volume_series.append({"date": d, "volume": int(row["Volume"])})

            # Bei weekly labels: auf Labels aggregieren (letzter Kurs der Woche)
            if label_list:
                from bisect import bisect_right
                label_dates = sorted(label_list)
                price_agg = {}
                vol_agg = {}
                for p in price_series:
                    i = bisect_right(label_dates, p["date"]) - 1
                    if 0 <= i < len(label_dates):
                        price_agg[label_dates[i]] = p["close"]  # letzter Kurs gewinnt
                for v in volume_series:
                    i = bisect_right(label_dates, v["date"]) - 1
                    if 0 <= i < len(label_dates):
                        vol_agg[label_dates[i]] = vol_agg.get(label_dates[i], 0) + v["volume"]
                price_series = [{"date": d, "close": price_agg[d]} for d in label_dates if d in price_agg]
                volume_series = [{"date": d, "volume": vol_agg[d]} for d in label_dates if d in vol_agg]

            results.append({
                "ticker": symbol,
                "name": name,
                "currency": currency,
                "price": price_series,
                "volume": volume_series,
            })
        except Exception as e:
            results.append({"ticker": symbol, "error": str(e)})

    return jsonify(results)


@app.route("/api/ext-source-availability")
@login_required
def api_ext_source_availability():
    """Prüft welche externen Quellen konfiguriert sind."""
    uid = current_user.id
    from transport import _get_credential
    has_bluesky = bool(_get_credential("bluesky_handle", "BLUESKY_HANDLE", uid)) and \
                  bool(_get_credential("bluesky_app_password", "BLUESKY_APP_PASSWORD", uid))
    has_telegram = bool(_get_credential("telegram_api_id", "TELEGRAM_API_ID", uid)) and \
                   bool(_get_credential("telegram_api_hash", "TELEGRAM_API_HASH", uid))
    has_copernicus = bool(_get_credential("copernicus_email", "COPERNICUS_EMAIL", uid)) and \
                     bool(_get_credential("copernicus_password", "COPERNICUS_PASSWORD", uid))
    return jsonify(bluesky=has_bluesky, telegram=has_telegram, copernicus=has_copernicus)



@app.route("/api/gdelt-timeline")
@login_required
def api_gdelt_timeline():
    """GDELT DOC 2.0 – tägliche Medien-Artikelanzahl für Suchbegriffe.
    Query-Parameter:
      terms – kommagetrennte Suchbegriffe (erforderlich)
      days  – Zeitraum in Tagen (Standard 180, max 365)
    """
    import requests as _req_gdelt
    import csv
    import io

    terms_param = request.args.get("terms", "")
    if not terms_param:
        abort(400, "terms Parameter fehlt")
    terms = [t.strip() for t in terms_param.split(",") if t.strip()][:5]
    days = min(int(request.args.get("days", 180)), 365)

    # GDELT nutzt Zeitspannen-Strings
    if days <= 30:
        timespan_str = f"{days}d"
    else:
        months = max(1, days // 30)
        timespan_str = f"{months}m"

    result = []
    for term in terms:
        url = (
            f"https://api.gdeltproject.org/api/v2/doc/doc"
            f"?query={_req_gdelt.utils.quote(term)}"
            f"&mode=timelinevolraw&timespan={timespan_str}&format=csv&TIMERES=day"
        )
        try:
            r = _req_gdelt.get(url, timeout=20)
            r.raise_for_status()
            text = r.text.strip()
            if not text:
                result.append({"term": term, "series": []})
                continue
            # BOM entfernen falls vorhanden
            if text.startswith("\ufeff"):
                text = text[1:]
            reader = csv.reader(io.StringIO(text))
            series = []
            header_skipped = False
            for row in reader:
                if len(row) < 3:
                    continue
                # Header überspringen
                if not header_skipped:
                    header_skipped = True
                    if row[0].strip().lower() == "date":
                        continue
                # Format: Date, Series, Value – nur "Article Count" Zeilen
                series_type = row[1].strip()
                if series_type != "Article Count":
                    continue
                date_str = row[0].strip()[:10]
                try:
                    count = int(float(row[2].strip()))
                except (ValueError, IndexError):
                    continue
                if len(date_str) == 10 and date_str[4] == "-":
                    series.append({"date": date_str, "count": count})
            result.append({"term": term, "series": series})
        except Exception as exc:
            result.append({"term": term, "error": str(exc)[:120]})

    return jsonify(result)


@app.route("/api/trends/runs")
@login_required
def api_trends_runs():
    """
    Liefert alle vorhandenen Abruf-Reihen (run_tags) für die angegebenen Keywords.
    Query: ?ids=1,2,...
    """
    ids_param = request.args.get("ids", "")
    if not ids_param:
        abort(400, "ids Parameter fehlt")
    try:
        ids = [int(i) for i in ids_param.split(",") if i.strip()]
    except ValueError:
        abort(400, "ids muss kommagetrennte Integer sein")

    from sqlalchemy import func as sqlfunc
    rows = (
        db.session.query(
            TrendData.run_tag,
            sqlfunc.max(TrendData.fetched_at).label("latest_fetch"),
            sqlfunc.count(TrendData.id).label("count"),
            sqlfunc.min(TrendData.date).label("min_date"),
            sqlfunc.max(TrendData.date).label("max_date"),
        )
        .filter(TrendData.keyword_id.in_(ids))
        .group_by(TrendData.run_tag)
        .order_by(sqlfunc.max(TrendData.fetched_at).desc())
        .all()
    )
    return jsonify([
        {
            "run_tag":    r.run_tag,
            "label":      r.run_tag if r.run_tag else (
                r.latest_fetch.strftime("%d.%m.%y %H:%M") if r.latest_fetch else "Standard"
            ),
            "fetched_at": r.latest_fetch.isoformat() if r.latest_fetch else None,
            "count":      r.count,
            "min_date":   r.min_date.isoformat() if r.min_date else None,
            "max_date":   r.max_date.isoformat() if r.max_date else None,
        }
        for r in rows
    ])


@app.route("/api/runs/rename", methods=["PATCH"])
@login_required
def api_runs_rename():
    """
    Benennt eine Abruf-Reihe um, indem run_tag in allen TrendData- und
    RegionInterest-Einträgen aktualisiert wird.
    Body: { old_tag: "...", new_label: "..." }
    """
    body    = request.get_json() or {}
    old_tag = body.get("old_tag", "")
    new_tag = (body.get("new_label") or "").strip()
    if not new_tag:
        abort(400, "new_label darf nicht leer sein")
    if old_tag == new_tag:
        return jsonify({"ok": True, "run_tag": new_tag})
    # Prüfen ob neuer Name bereits belegt ist
    exists = TrendData.query.filter_by(run_tag=new_tag).first()
    if exists:
        abort(409, f"Bezeichnung '{new_tag}' existiert bereits")
    TrendData.query.filter_by(run_tag=old_tag).update({"run_tag": new_tag})
    RegionInterest.query.filter_by(run_tag=old_tag).update({"run_tag": new_tag})
    db.session.commit()
    return jsonify({"ok": True, "run_tag": new_tag})


@app.route("/api/trends/run", methods=["DELETE"])
@login_required
def api_delete_run():
    """
    Löscht alle TrendData- und RegionInterest-Einträge mit dem angegebenen run_tag.
    Query: ?run_tag=<tag>  (Leerstring = Standard-Reihe)
    """
    run_tag = request.args.get("run_tag", None)
    if run_tag is None:
        abort(400, "run_tag Parameter fehlt")
    count_td = TrendData.query.filter_by(run_tag=run_tag).delete()
    count_ri = RegionInterest.query.filter_by(run_tag=run_tag).delete()
    db.session.commit()
    return jsonify({"ok": True, "deleted_trend": count_td, "deleted_region": count_ri})


@app.route("/api/runs/preview-data")
@login_required
def api_runs_preview_data():
    """
    Gibt Rohdatenpunkte für einen run_tag und eine Liste von Keyword-IDs zurück.
    Query: ?ids=1,2&run_tag=TAG
    Response: [{keyword_id, keyword, data:[{date, value}]}]
    """
    ids_param = request.args.get("ids", "")
    run_tag   = request.args.get("run_tag", "")
    if not ids_param:
        return jsonify([])
    ids = [int(i) for i in ids_param.split(",") if i.strip()]

    rows = (TrendData.query
            .filter(TrendData.keyword_id.in_(ids), TrendData.run_tag == run_tag)
            .order_by(TrendData.keyword_id, TrendData.date)
            .all())

    from collections import defaultdict
    groups = defaultdict(list)
    for row in rows:
        d = row.date
        if hasattr(d, "hour") and (d.hour != 0 or d.minute != 0 or d.second != 0):
            date_str = d.strftime("%Y-%m-%dT%H:%M")
        else:
            date_str = d.strftime("%Y-%m-%d") if hasattr(d, "strftime") else str(d)[:10]
        groups[row.keyword_id].append({"date": date_str, "value": row.value})

    from models import Keyword
    kw_map = {k.id: k.keyword for k in Keyword.query.filter(Keyword.id.in_(ids)).all()}
    return jsonify([
        {"keyword_id": kw_id, "keyword": kw_map.get(kw_id, str(kw_id)), "data": pts}
        for kw_id, pts in groups.items()
    ])


@app.route("/api/runs/merge", methods=["POST"])
@login_required
def api_runs_merge():
    """
    Führt mehrere Abruf-Reihen zu einer neuen zusammen.
    Body: {
        keyword_ids: [1, 2, ...],
        segments:    [{run_tag: "...", scale: 1.0}, ...],
        new_label:   "Bezeichnung der neuen Reihe"
    }
    """
    from collections import defaultdict
    from datetime import datetime, timezone

    body        = request.get_json(force=True, silent=True) or {}
    keyword_ids = [int(i) for i in body.get("keyword_ids", [])]
    segments    = body.get("segments", [])   # [{run_tag, scale}]
    new_label   = (body.get("new_label") or "").strip()

    if not keyword_ids or len(segments) < 2:
        return jsonify({"error": "keyword_ids und mindestens 2 Segmente erforderlich"}), 400
    if not new_label:
        return jsonify({"error": "Bezeichnung (new_label) fehlt"}), 400

    # Sicherstellen, dass der run_tag eindeutig ist
    new_run_tag = new_label
    if TrendData.query.filter_by(run_tag=new_run_tag).first():
        new_run_tag = f"{new_label} ({datetime.now(timezone.utc).strftime('%H%M%S')})"

    now = datetime.now(timezone.utc)
    total_inserted = 0

    for kw_id in keyword_ids:
        # Pro Datum: skalierte Werte aller Segmente sammeln, dann mitteln
        by_date: dict = {}   # date → [scaled_value, ...]

        for seg in segments:
            rt    = seg.get("run_tag", "")
            scale = max(0.01, min(10.0, float(seg.get("scale", 1.0))))
            rows  = TrendData.query.filter_by(keyword_id=kw_id, run_tag=rt).all()
            for row in rows:
                scaled = int(min(100, max(0, round(row.value * scale))))
                if row.date not in by_date:
                    by_date[row.date] = []
                by_date[row.date].append(scaled)

        # Neue TrendData-Zeilen anlegen (Mittelwert bei Überlappungen)
        for date_key, vals in by_date.items():
            avg_val = round(sum(vals) / len(vals))
            td = TrendData(
                keyword_id=kw_id,
                date=date_key,
                value=avg_val,
                run_tag=new_run_tag,
                fetched_at=now,
            )
            db.session.add(td)
            total_inserted += 1

    db.session.commit()
    return jsonify({"ok": True, "run_tag": new_run_tag, "inserted": total_inserted})


# ---------------------------------------------------------------------------
# API – Einstellungen
# ---------------------------------------------------------------------------

@app.route("/api/system-settings", methods=["GET"])
def api_system_settings():
    """Datums-/Zeitformat: per-User-Präferenz wenn eingeloggt, sonst Systemstandard."""
    def _get_global(key, default=""):
        obj = AppSetting.query.filter_by(key=key, user_id=None).first()
        return obj.value if obj and obj.value else default
    date_fmt   = _get_global("date_format",  "DD.MM.YY")
    time_fmt   = _get_global("time_format",  "HH:mm")
    trends_tz  = _get_global("trends_tz",    "60")
    ui_lang    = _get_global("ui_language",  "de")
    query_lang = _get_global("query_language", "auto")
    accent1    = _get_global("accent1", "")
    accent2    = _get_global("accent2", "")
    if current_user.is_authenticated:
        obj_d  = AppSetting.query.filter_by(key="date_format",    user_id=current_user.id).first()
        obj_t  = AppSetting.query.filter_by(key="time_format",    user_id=current_user.id).first()
        obj_tz = AppSetting.query.filter_by(key="trends_tz",      user_id=current_user.id).first()
        obj_l  = AppSetting.query.filter_by(key="ui_language",    user_id=current_user.id).first()
        obj_ql = AppSetting.query.filter_by(key="query_language", user_id=current_user.id).first()
        if obj_d  and obj_d.value:  date_fmt   = obj_d.value
        if obj_t  and obj_t.value:  time_fmt   = obj_t.value
        if obj_tz and obj_tz.value: trends_tz  = obj_tz.value
        if obj_l  and obj_l.value:  ui_lang    = obj_l.value
        if obj_ql and obj_ql.value: query_lang = obj_ql.value
        obj_a1 = AppSetting.query.filter_by(key="accent1", user_id=current_user.id).first()
        obj_a2 = AppSetting.query.filter_by(key="accent2", user_id=current_user.id).first()
        if obj_a1 and obj_a1.value: accent1 = obj_a1.value
        if obj_a2 and obj_a2.value: accent2 = obj_a2.value
    return jsonify({"date_format": date_fmt, "time_format": time_fmt,
                    "trends_tz": trends_tz, "ui_language": ui_lang,
                    "query_language": query_lang,
                    "accent1": accent1, "accent2": accent2})


@app.route("/api/settings", methods=["GET"])
@login_required
def api_settings_get():
    key = request.args.get("key")
    if key:
        obj = AppSetting.query.filter_by(key=key, user_id=current_user.id).first()
        return jsonify({"key": key, "value": obj.value if obj else ""})
    settings = AppSetting.query.filter_by(user_id=current_user.id).all()
    return jsonify({s.key: s.value for s in settings})


@app.route("/api/settings", methods=["POST"])
@login_required
def api_settings_set():
    data = request.get_json(force=True) or {}
    key = (data.get("key") or "").strip()
    value = data.get("value", "")
    if not key:
        abort(400, "key fehlt")
    obj = AppSetting.query.filter_by(key=key, user_id=current_user.id).first()
    if obj:
        obj.value = value
    else:
        db.session.add(AppSetting(key=key, value=value, user_id=current_user.id))
    db.session.commit()
    return jsonify({"key": key, "value": value})


# ---------------------------------------------------------------------------
# API – Ereignisse
# ---------------------------------------------------------------------------

def _parse_event_dt(s: str) -> datetime:
    """Parst YYYY-MM-DD oder YYYY-MM-DDTHH:MM zu datetime."""
    s = s.strip()
    if "T" in s:
        return datetime.strptime(s[:16], "%Y-%m-%dT%H:%M")
    return datetime.strptime(s[:10], "%Y-%m-%d")


@app.route("/api/events", methods=["GET"])
@login_required
def api_events_list():
    project_id = request.args.get("project_id", type=int)
    q = Event.query.filter_by(user_id=current_user.id).order_by(Event.start_dt)
    if project_id:
        q = q.filter(Event.project_id == project_id)
    return jsonify([e.to_dict() for e in q.all()])


@app.route("/api/events", methods=["POST"])
@login_required
def api_events_create():
    data = request.get_json(force=True) or {}
    title = (data.get("title") or "").strip()
    if not title:
        abort(400, "title fehlt")

    start_str = (data.get("start_dt") or "").strip()
    if not start_str:
        abort(400, "start_dt fehlt")
    try:
        start_dt = _parse_event_dt(start_str)
    except ValueError:
        abort(400, "Ungültiges Start-Datum")

    event_type = data.get("event_type", "point")
    end_dt = None
    end_str = (data.get("end_dt") or "").strip()
    if event_type == "range" and end_str:
        try:
            end_dt = _parse_event_dt(end_str)
        except ValueError:
            abort(400, "Ungültiges End-Datum")

    pid = data.get("project_id")
    evt = Event(
        title=title,
        description=data.get("description", ""),
        event_type=event_type,
        start_dt=start_dt,
        end_dt=end_dt,
        color=data.get("color", "#f75f4f"),
        project_id=int(pid) if pid else None,
        user_id=current_user.id,
    )
    db.session.add(evt)
    db.session.commit()
    return jsonify(evt.to_dict()), 201


@app.route("/api/events/<int:eid>", methods=["PUT"])
@login_required
def api_events_update(eid):
    evt = Event.query.filter_by(id=eid, user_id=current_user.id).first()
    if not evt:
        abort(404)
    data = request.get_json(force=True) or {}

    if "title" in data:
        evt.title = (data["title"] or "").strip()
    if "description" in data:
        evt.description = data["description"]
    if "event_type" in data:
        evt.event_type = data["event_type"]
    if "color" in data:
        evt.color = data["color"]
    if "start_dt" in data and data["start_dt"]:
        try:
            evt.start_dt = _parse_event_dt(data["start_dt"])
        except ValueError:
            abort(400, "Ungültiges Start-Datum")
    if "end_dt" in data:
        if data["end_dt"]:
            try:
                evt.end_dt = _parse_event_dt(data["end_dt"])
            except ValueError:
                abort(400, "Ungültiges End-Datum")
        else:
            evt.end_dt = None
    if "project_id" in data:
        evt.project_id = int(data["project_id"]) if data["project_id"] else None

    db.session.commit()
    return jsonify(evt.to_dict())


@app.route("/api/events/<int:eid>", methods=["DELETE"])
@login_required
def api_events_delete(eid):
    evt = Event.query.filter_by(id=eid, user_id=current_user.id).first()
    if not evt:
        abort(404)
    db.session.delete(evt)
    db.session.commit()
    return jsonify({"ok": True})


# ── Watch Zones CRUD ──────────────────────────────────────────────────────

@app.route("/api/watchzones", methods=["GET"])
@login_required
def api_watchzones_list():
    from models import WatchZone
    zones = WatchZone.query.filter_by(user_id=current_user.id).order_by(WatchZone.created_at.desc()).all()
    return jsonify([z.to_dict() for z in zones])


@app.route("/api/resolve-domain-location", methods=["POST"])
@login_required
def api_resolve_domain_location():
    """Löst eine Domain per DNS auf und gibt den Server-Standort zurück."""
    d = request.get_json(force=True)
    domain = (d.get("domain") or "").strip()
    if not domain:
        return jsonify({"error": "Kein Domain angegeben"}), 400
    from transport import resolve_domain_location
    result = resolve_domain_location(domain)
    if not result:
        return jsonify({"error": f"Server-Standort für '{domain}' nicht ermittelbar"}), 404
    return jsonify(result)


@app.route("/api/watchzones", methods=["POST"])
@login_required
def api_watchzones_create():
    import json as _j
    from models import WatchZone
    d = request.get_json(force=True)
    name = (d.get("name") or "").strip()
    zone_type = d.get("zone_type", "")
    _valid_zone_types = ("global",) + tuple(PluginManager.all_of_type("watchzone").keys())
    if zone_type not in _valid_zone_types:
        return jsonify({"error": "Ung\u00fcltiger zone_type"}), 400
    geometry = d.get("geometry", {})
    config = d.get("config", {})
    project_id = d.get("project_id")
    z = WatchZone(
        name=name or f"Zone {zone_type}",
        zone_type=zone_type,
        geometry=_j.dumps(geometry),
        config=_j.dumps(config),
        active=d.get("active", True),
        project_id=project_id,
        user_id=current_user.id,
    )
    db.session.add(z)
    db.session.commit()
    return jsonify(z.to_dict()), 201


@app.route("/api/watchzones/<int:zid>", methods=["PUT"])
@login_required
def api_watchzones_update(zid):
    import json as _j
    from models import WatchZone
    z = WatchZone.query.filter_by(id=zid, user_id=current_user.id).first()
    if not z:
        abort(404)
    d = request.get_json(force=True)
    if "name" in d:
        z.name = (d["name"] or "").strip()
    if "geometry" in d:
        z.geometry = _j.dumps(d["geometry"])
    if "config" in d:
        z.config = _j.dumps(d["config"])
    if "active" in d:
        z.active = bool(d["active"])
    if "project_id" in d:
        z.project_id = d["project_id"]
    db.session.commit()
    return jsonify(z.to_dict())


@app.route("/api/watchzones/<int:zid>", methods=["DELETE"])
@login_required
def api_watchzones_delete(zid):
    from models import WatchZone
    z = WatchZone.query.filter_by(id=zid, user_id=current_user.id).first()
    if not z:
        abort(404)
    db.session.delete(z)
    db.session.commit()
    return jsonify({"ok": True})


@app.route("/api/watchzones/<int:zid>/live")
@login_required
def api_watchzones_live(zid):
    """Live-Daten für eine Watch Zone abrufen — delegiert an Plugin-Registry."""
    import json as _j
    from models import WatchZone
    z = WatchZone.query.filter_by(id=zid, user_id=current_user.id).first()
    if not z:
        abort(404)
    geo = _j.loads(z.geometry) if z.geometry else {}
    zone_type = request.args.get("as_type") or z.zone_type
    if zone_type == "global":
        return jsonify({"error": "Bitte as_type angeben für globale Zonen"}), 400

    plugin = PluginManager.get("watchzone", zone_type)
    if not plugin:
        return jsonify({"error": f"Unbekannter Zone-Typ: {zone_type}"}), 400

    config = _j.loads(z.config) if z.config else {}
    # Query-Parameter an Config weiterreichen (z.B. from/to für Wayback CDX)
    for _qk in ("from", "to"):
        if request.args.get(_qk):
            config[_qk] = request.args[_qk]
    bbox = _geojson_to_bbox(geo)

    try:
        result = plugin.live_handler(z, config, geo, bbox, current_user.id)
        if "error" in result:
            return jsonify(result), 400
        return jsonify(result)
    except Exception as e:
        log.warning("WatchZone Live-Daten Fehler (Zone %d): %s", zid, e)
        return jsonify({"error": str(e)}), 502


# ── WZ-History-Routen werden dynamisch aus Plugins registriert ───────────
# (siehe _register_wz_history_routes() weiter unten, aufgerufen nach Routendefinition)



@app.route("/api/watchzones/<int:zid>/snapshot-diff")
@login_required
def api_snapshot_diff(zid):
    """Liefert einen Wayback-Snapshot mit optionaler Diff-Markierung gegenüber dem Vorgänger."""
    import json as _j
    from models import WatchZone
    z = WatchZone.query.filter_by(id=zid, user_id=current_user.id).first()
    if not z:
        abort(404)
    config = _j.loads(z.config) if z.config else {}
    url = config.get("url", "")
    if not url:
        return jsonify({"error": "Keine URL konfiguriert"}), 400
    ts2 = request.args.get("ts", "").strip()
    ts1 = request.args.get("ts1", "").strip() or None
    if not ts2:
        return jsonify({"error": "Parameter 'ts' erforderlich"}), 400
    try:
        from plugins.watchzone.website._transport import fetch_wayback_diff_html
        result = fetch_wayback_diff_html(url, ts2, ts1)
        return jsonify(result)
    except Exception as e:
        log.warning("Snapshot-Diff Fehler (Zone %d): %s", zid, e)
        return jsonify({"error": str(e)}), 502



@app.route("/api/ip/forensics")
@login_required
def api_ip_forensics():
    """WHOIS + BGP-Prefix-Daten für eine IP via RIPE Stat."""
    import requests as _req
    ip = request.args.get("ip", "").strip()
    if not ip:
        return jsonify({"error": "ip required"}), 400

    result = {"ip": ip, "whois": {}, "bgp": {}}

    # ── WHOIS via RIPE Stat ──────────────────────────────────────────────────
    try:
        wr = _req.get(
            f"https://stat.ripe.net/data/whois/data.json?resource={ip}",
            timeout=6, headers={"User-Agent": "veritrend-forensics/1.0"}
        ).json()
        org = country = abuse = netname = None
        for record in wr.get("data", {}).get("records", []):
            for field in record:
                k = (field.get("key") or "").lower()
                v = (field.get("value") or "").strip()
                if not v:
                    continue
                if k in ("org-name", "orgname", "owner", "descr") and not org:
                    org = v
                if k == "country" and not country:
                    country = v.upper()
                if k in ("abuse-mailbox", "orgabuseemail", "e-mail") and "@" in v and not abuse:
                    abuse = v
                if k == "netname" and not netname:
                    netname = v
        result["whois"] = {k: v for k, v in {
            "org": org, "country": country, "abuse": abuse, "netname": netname
        }.items() if v}
    except Exception:
        pass

    # ── BGP Prefix-Übersicht via RIPE Stat ───────────────────────────────────
    try:
        br = _req.get(
            f"https://stat.ripe.net/data/prefix-overview/data.json?resource={ip}",
            timeout=6, headers={"User-Agent": "veritrend-forensics/1.0"}
        ).json()
        data = br.get("data", {})
        asns = data.get("asns", [])
        result["bgp"] = {
            "prefix":     data.get("resource", ""),
            "announced":  data.get("announced", False),
            "asns":       [{"asn": a.get("asn"), "holder": a.get("holder", "")} for a in asns[:3]],
        }
    except Exception:
        pass

    return jsonify(result)


def _geojson_to_bbox(geo):
    """Berechnet [lon_min, lat_min, lon_max, lat_max] aus einer GeoJSON-Geometrie."""
    coords = []
    def _extract(obj):
        if isinstance(obj, list):
            if len(obj) >= 2 and isinstance(obj[0], (int, float)):
                coords.append(obj)
            else:
                for item in obj:
                    _extract(item)
    _extract(geo.get("coordinates", []))
    if not coords:
        return None
    lons = [c[0] for c in coords]
    lats = [c[1] for c in coords]
    return [min(lons), min(lats), max(lons), max(lats)]


_GDELT_LANG_ISO = {
    "Afrikaans": "af", "Albanian": "sq", "Arabic": "ar", "Armenian": "hy",
    "Azerbaijani": "az", "Basque": "eu", "Belarusian": "be", "Bengali": "bn",
    "Bosnian": "bs", "Bulgarian": "bg", "Burmese": "my", "Catalan": "ca",
    "Chinese": "zh", "Croatian": "hr", "Czech": "cs", "Danish": "da",
    "Dutch": "nl", "English": "en", "Estonian": "et", "Filipino": "tl",
    "Finnish": "fi", "French": "fr", "Galician": "gl", "Georgian": "ka",
    "German": "de", "Greek": "el", "Gujarati": "gu", "Hebrew": "he",
    "Hindi": "hi", "Hungarian": "hu", "Icelandic": "is", "Indonesian": "id",
    "Irish": "ga", "Italian": "it", "Japanese": "ja", "Kannada": "kn",
    "Kazakh": "kk", "Khmer": "km", "Korean": "ko", "Kurdish": "ku",
    "Latvian": "lv", "Lithuanian": "lt", "Macedonian": "mk", "Malay": "ms",
    "Malayalam": "ml", "Maltese": "mt", "Marathi": "mr", "Mongolian": "mn",
    "Nepali": "ne", "Norwegian": "no", "Pashto": "ps", "Persian": "fa",
    "Polish": "pl", "Portuguese": "pt", "Punjabi": "pa", "Romanian": "ro",
    "Russian": "ru", "Serbian": "sr", "Sinhala": "si", "Slovak": "sk",
    "Slovenian": "sl", "Somali": "so", "Spanish": "es", "Swahili": "sw",
    "Swedish": "sv", "Tamil": "ta", "Telugu": "te", "Thai": "th",
    "Turkish": "tr", "Ukrainian": "uk", "Urdu": "ur", "Uzbek": "uz",
    "Vietnamese": "vi", "Welsh": "cy",
}


def _translate_to_de(title: str, src_lang_name: str) -> str:
    """Übersetzt einen Titel ins Deutsche via deep_translator (GoogleTranslator)."""
    src_iso = _GDELT_LANG_ISO.get(src_lang_name, "")
    if not src_iso or src_iso == "de" or not title:
        return title
    try:
        from deep_translator import GoogleTranslator
        result = GoogleTranslator(source=src_iso, target="de").translate(title[:450])
        return result or title
    except Exception:
        return title


_GNEWS_LANG_MAP = {
    # Frontend-Wert → (hl, gl, ceid, lang_name_for_translation)
    "":        ("de", "DE", "DE:de", None),       # Alle → deutschsprachiger Feed
    "German":  ("de", "DE", "DE:de", None),
    "English": ("en", "US", "US:en", "English"),
    "Arabic":  ("ar", "SA", "SA:ar", "Arabic"),
    "French":  ("fr", "FR", "FR:fr", "French"),
    "Spanish": ("es", "ES", "ES:es", "Spanish"),
}


def _translate_item(args):
    title, lang_name = args
    return _translate_to_de(title, lang_name)


_BING_MKT_MAP = {
    # Frontend-Wert → (mkt, lang_name_for_translation)
    "":        ("de-DE", None),
    "German":  ("de-DE", None),
    "English": ("en-US", "English"),
    "Arabic":  ("ar-SA", "Arabic"),
    "French":  ("fr-FR", "French"),
    "Spanish": ("es-ES", "Spanish"),
}

_NEWSAPI_LANG_MAP = {
    # Frontend-Wert → ISO-639-1-Sprachcode für NewsAPI
    "":        "de",
    "German":  "de",
    "English": "en",
    "Arabic":  "ar",
    "French":  "fr",
    "Spanish": "es",
}


def _rss_parse_items(resp_text: str) -> list:
    """Parst Standard-RSS-XML (Google News, Bing News) → Liste von Ergebnis-Dicts."""
    from xml.etree import ElementTree as ET
    from email.utils import parsedate_to_datetime
    from urllib.parse import urlparse

    try:
        root    = ET.fromstring(resp_text)
        channel = root.find("channel")
        items   = channel.findall("item") if channel is not None else []
    except Exception:
        return []

    results = []
    for item in items:
        title_raw = (item.findtext("title") or "").strip()
        url       = (item.findtext("link")  or "").strip()
        pubdate   = (item.findtext("pubDate") or "").strip()
        source    = (item.findtext("source") or "").strip()

        title = title_raw
        if " - " in title_raw:
            parts = title_raw.rsplit(" - ", 1)
            title = parts[0].strip()
            if not source:
                source = parts[1].strip()

        dt_iso = ""
        if pubdate:
            try:
                dt_iso = parsedate_to_datetime(pubdate).strftime("%Y-%m-%dT%H:%M")
            except Exception:
                pass

        domain = source
        if url and not domain:
            try:
                domain = urlparse(url).netloc
            except Exception:
                pass

        results.append({
            "title":          title,
            "title_original": None,
            "url":            url,
            "domain":         domain,
            "seendate":       dt_iso,
            "language":       None,
            "sourcecountry":  None,
        })
    return results


@app.route("/api/events/find-news")
@login_required
def api_events_find_news():
    """
    Sucht über Nachrichtenquellen nach der ersten Berichterstattung zu einem Thema.
    Fremdsprachige Schlagzeilen werden automatisch ins Deutsche übersetzt (MyMemory).
    Ergebnisse sind chronologisch sortiert (älteste zuerst).

    Query-Parameter:
      q      – Suchbegriff (erforderlich)
      from   – Startdatum YYYY-MM-DD (optional)
      to     – Enddatum YYYY-MM-DD (optional)
      lang   – Sprachregion: "" | "German" | "English" | "Arabic" | "French" | "Spanish"
      source – Nachrichtenquelle: "google" (Standard) | "bing" | "all" | "newsapi"
    """
    import requests as req_lib
    from concurrent.futures import ThreadPoolExecutor

    q = request.args.get("q", "").strip()
    if not q:
        abort(400, "q Parameter fehlt")

    from_date = request.args.get("from",   "").strip()
    to_date   = request.args.get("to",     "").strip()
    lang_key  = request.args.get("lang",   "").strip()
    source    = request.args.get("source", "google").strip()
    limit     = min(int(request.args.get("limit", 25)), 100)

    hl, gl, ceid, translate_as = _GNEWS_LANG_MAP.get(lang_key, _GNEWS_LANG_MAP[""])
    bing_mkt, _                = _BING_MKT_MAP.get(lang_key,   _BING_MKT_MAP[""])

    search_q = q
    if from_date:
        search_q += f" after:{from_date}"
    if to_date:
        search_q += f" before:{to_date}"

    UA = (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/120.0.0.0 Safari/537.36"
    )

    def fetch_google():
        try:
            r = req_lib.get(
                "https://news.google.com/rss/search",
                params={"q": search_q, "hl": hl, "gl": gl, "ceid": ceid},
                headers={"User-Agent": UA},
                timeout=15,
            )
            r.raise_for_status()
            items = _rss_parse_items(r.text)
            for it in items:
                it["language"]      = hl
                it["sourcecountry"] = gl
            return items
        except Exception as exc:
            log.warning("Google News RSS fehlgeschlagen: %s", exc)
            return None

    def fetch_bing():
        try:
            r = req_lib.get(
                "https://www.bing.com/news/search",
                params={"q": search_q, "format": "RSS", "mkt": bing_mkt},
                headers={"User-Agent": UA},
                timeout=15,
            )
            r.raise_for_status()
            items = _rss_parse_items(r.text)
            lang_code = bing_mkt.split("-")[0] if "-" in bing_mkt else bing_mkt
            cc        = bing_mkt.split("-")[-1] if "-" in bing_mkt else bing_mkt
            for it in items:
                it["language"]      = lang_code
                it["sourcecountry"] = cc
            return items
        except Exception as exc:
            log.warning("Bing News RSS fehlgeschlagen: %s", exc)
            return None

    def fetch_newsapi(api_key):
        na_lang = _NEWSAPI_LANG_MAP.get(lang_key, "de")
        params  = {
            "q":        q,
            "apiKey":   api_key,
            "sortBy":   "publishedAt",
            "pageSize": 100,
            "language": na_lang,
        }
        if from_date:
            params["from"] = from_date
        if to_date:
            params["to"] = to_date
        try:
            r = req_lib.get(
                "https://newsapi.org/v2/everything",
                params=params,
                headers={"User-Agent": UA},
                timeout=15,
            )
            r.raise_for_status()
            data = r.json()
        except Exception as exc:
            log.warning("NewsAPI fehlgeschlagen: %s", exc)
            return None

        items = []
        for article in data.get("articles", []):
            title       = (article.get("title") or "").strip()
            url         = (article.get("url")   or "").strip()
            pub         = (article.get("publishedAt") or "").strip()
            source_name = (article.get("source", {}).get("name") or "").strip()

            if not title or title == "[Removed]" or "removed.com" in url:
                continue

            dt_iso = ""
            if pub:
                try:
                    dt_iso = datetime.fromisoformat(
                        pub.replace("Z", "+00:00")
                    ).strftime("%Y-%m-%dT%H:%M")
                except Exception:
                    pass

            items.append({
                "title":          title,
                "title_original": None,
                "url":            url,
                "domain":         source_name,
                "seendate":       dt_iso,
                "language":       na_lang,
                "sourcecountry":  "",
            })
        return items

    def fetch_gdelt(max_records=None):
        """GDELT DOC 2.0 Article Search – gibt Einzelartikel zurück."""
        import time as _time
        _max = max_records or limit
        for attempt in range(3):
            try:
                params = {
                    "query":       q,
                    "mode":        "artlist",
                    "maxrecords":  min(_max, 250),
                    "format":      "json",
                    "sort":        "DateAsc",
                }
                if from_date:
                    params["startdatetime"] = from_date.replace("-", "") + "000000"
                if to_date:
                    params["enddatetime"]   = to_date.replace("-", "") + "235959"
                r = req_lib.get(
                    "https://api.gdeltproject.org/api/v2/doc/doc",
                    params=params,
                    headers={"User-Agent": UA},
                    timeout=20,
                )
                if r.status_code == 429 and attempt < 2:
                    log.info("GDELT rate-limited, retry in %ds", (attempt + 1) * 5)
                    _time.sleep((attempt + 1) * 5)
                    continue
                r.raise_for_status()
                data = r.json()
                items = []
                for art in data.get("articles", []):
                    title   = (art.get("title") or "").strip()
                    url     = (art.get("url")   or "").strip()
                    seendt  = (art.get("seendate") or "").strip()
                    domain  = (art.get("domain")  or "").strip()
                    lang    = (art.get("language") or "").strip()
                    country = (art.get("sourcecountry") or "").strip()
                    if not title:
                        continue
                    # GDELT seendate Format: "YYYYMMDDTHHMMSSz" → "YYYY-MM-DDTHH:MM"
                    dt_iso = ""
                    if seendt and len(seendt) >= 8:
                        try:
                            dt_iso = f"{seendt[:4]}-{seendt[4:6]}-{seendt[6:8]}"
                            if len(seendt) >= 13:
                                dt_iso += f"T{seendt[9:11]}:{seendt[11:13]}"
                        except Exception:
                            pass
                    items.append({
                        "title":          title,
                        "title_original": None,
                        "url":            url,
                        "domain":         domain,
                        "seendate":       dt_iso,
                        "language":       lang,
                        "sourcecountry":  country,
                    })
                log.info("GDELT: %d Artikel abgerufen", len(items))
                return items
            except Exception as exc:
                if attempt < 2:
                    _time.sleep((attempt + 1) * 3)
                    continue
                log.warning("GDELT Article Search fehlgeschlagen: %s", exc)
                return None
        return None

    results = []
    if source == "bing":
        if limit > 25:
            with ThreadPoolExecutor(max_workers=2) as pool:
                bing_future  = pool.submit(fetch_bing)
                gdelt_future = pool.submit(fetch_gdelt, limit)
            bing_items  = bing_future.result()
            gdelt_items = gdelt_future.result() or []
            if bing_items is None and not gdelt_items:
                abort(502, "Bing News nicht erreichbar")
            bing_items = bing_items or []
            log.info("Bing: %d, GDELT: %d (limit=%d)", len(bing_items), len(gdelt_items), limit)
            seen = set()
            for it in bing_items + gdelt_items:
                key = it["title"].lower().strip()
                if key not in seen:
                    seen.add(key)
                    results.append(it)
        else:
            items = fetch_bing()
            if items is None:
                abort(502, "Bing News nicht erreichbar")
            results = list(items)
    elif source == "gdelt":
        items = fetch_gdelt()
        if items is None:
            abort(502, "GDELT nicht erreichbar")
        results = items
    elif source == "newsapi":
        _na_setting = AppSetting.query.filter_by(key="newsapi_key", user_id=None).first()
        api_key = (_na_setting.value if _na_setting and _na_setting.value else os.getenv("NEWSAPI_KEY", "")).strip()
        if not api_key:
            abort(400, "NewsAPI-Key nicht konfiguriert – bitte im Admin-Bereich unter 'API & Konfiguration' eintragen")
        items = fetch_newsapi(api_key)
        if items is None:
            abort(502, "NewsAPI nicht erreichbar")
        results = items
    elif source == "all":
        google_items = fetch_google() or []
        bing_items   = fetch_bing()   or []
        gdelt_items  = fetch_gdelt()  or []
        seen = set()
        for it in google_items + bing_items + gdelt_items:
            key = it["title"].lower().strip()
            if key not in seen:
                seen.add(key)
                results.append(it)
    else:  # "google" (Standard)
        if limit > 25:
            # Google RSS liefert max ~25 Ergebnisse → GDELT parallel dazuholen
            with ThreadPoolExecutor(max_workers=2) as pool:
                google_future = pool.submit(fetch_google)
                gdelt_future  = pool.submit(fetch_gdelt, limit)
            google_items = google_future.result() or []
            gdelt_items  = gdelt_future.result()  or []
            log.info("Google News: %d, GDELT: %d (limit=%d)", len(google_items), len(gdelt_items), limit)
            # Merge: Google zuerst (höhere Relevanz), dann GDELT-Ergänzung
            seen = set()
            for it in google_items + gdelt_items:
                key = it["title"].lower().strip()
                if key not in seen:
                    seen.add(key)
                    results.append(it)
            log.info("Merged: %d Ergebnisse", len(results))
        else:
            items = fetch_google()
            if items is None:
                abort(502, "Google News nicht erreichbar")
            results = list(items)
            log.info("Google News: %d Ergebnisse", len(results))

    # Älteste zuerst
    results.sort(key=lambda r: r["seendate"] or "")

    # Post-fetch Datumsfilter (falls Quelle die after:/before:-Operatoren ignoriert)
    if from_date:
        results = [r for r in results if not r["seendate"] or r["seendate"][:10] >= from_date]
    if to_date:
        results = [r for r in results if not r["seendate"] or r["seendate"][:10] <= to_date]

    # Fremdsprachige Titel parallel ins Deutsche übersetzen
    if translate_as:
        pairs = [(r["title"], translate_as) for r in results if r["title"]]
        if pairs:
            with ThreadPoolExecutor(max_workers=6) as pool:
                translated_list = list(pool.map(_translate_item, pairs))
            for i, translated in enumerate(translated_list):
                if translated and translated != results[i]["title"]:
                    results[i]["title_original"] = results[i]["title"]
                    results[i]["title"]          = translated

    return jsonify(results[:limit])


@app.route("/api/events/trend-check")
@login_required
def api_events_trend_check():
    """
    Prüft Suchnachfrage für einen Begriff zu mehreren Datumsangaben in der lokalen DB.

    Query-Parameter:
      q     – Suchbegriff (case-insensitiver Teilstring-Match gegen Keyword.keyword)
      dates – kommagetrennte YYYY-MM-DD-Datumsangaben
    """
    from datetime import timedelta

    q         = request.args.get("q",     "").strip()
    dates_raw = request.args.get("dates", "").strip()

    if not q or not dates_raw:
        return jsonify({"keywords": [], "by_date": {}})

    dates = list({d[:10] for d in dates_raw.split(",") if d.strip()})

    matching_kws = (Keyword.query
                    .filter(Keyword.keyword.ilike(f"%{q}%"))
                    .order_by(db.func.length(Keyword.keyword))
                    .all())
    kw_ids   = [kw.id for kw in matching_kws]
    kw_names = list(dict.fromkeys(kw.keyword for kw in matching_kws))  # dedupliziert
    kw_map   = {kw.id: kw.keyword for kw in matching_kws}

    by_date = {}
    for date_str in dates:
        try:
            target = datetime.fromisoformat(date_str).date()
        except ValueError:
            by_date[date_str] = None
            continue

        if not kw_ids:
            by_date[date_str] = None
            continue

        lo = datetime(target.year, target.month, target.day) - timedelta(days=7)
        hi = datetime(target.year, target.month, target.day) + timedelta(days=7)

        rows = TrendData.query.filter(
            TrendData.keyword_id.in_(kw_ids),
            TrendData.date >= lo,
            TrendData.date <= hi,
        ).all()

        if not rows:
            by_date[date_str] = None
            continue

        # Nächsten Tag ermitteln, dann Maximum über alle run_tags
        closest_day = min(
            (r.date.date() for r in rows),
            key=lambda d: abs((d - target).days),
        )
        day_rows = [r for r in rows if r.date.date() == closest_day]
        best     = max(day_rows, key=lambda r: r.value)

        by_date[date_str] = {
            "value":    best.value,
            "keyword":  kw_map.get(best.keyword_id, ""),
            "date":     closest_day.isoformat(),
            "days_off": abs((closest_day - target).days),
        }

    return jsonify({"keywords": kw_names, "by_date": by_date})


@app.route("/api/events/trend-series")
@login_required
def api_events_trend_series():
    """
    Zeitreihen für bis zu 5 Keywords (unscharf gegen q gematcht) im gewählten Zeitraum.

    Query-Parameter:
      q    – Suchbegriff (ILIKE-Match gegen Keyword.keyword)
      from – Anfang des Zeitraums (YYYY-MM-DD), optional
      to   – Ende   des Zeitraums (YYYY-MM-DD), optional
    """
    from collections import defaultdict
    from datetime import timedelta

    q       = request.args.get("q",    "").strip()
    from_dt = request.args.get("from", "").strip()
    to_dt   = request.args.get("to",   "").strip()

    if not q:
        return jsonify({"keywords": [], "series": []})

    matching_kws = (Keyword.query
                    .filter(Keyword.keyword.ilike(f"%{q}%"))
                    .order_by(db.func.length(Keyword.keyword))
                    .limit(5).all())
    if not matching_kws:
        return jsonify({"keywords": [], "series": []})

    # Zeitfenster: 30 Tage vor dem ersten Artikel bis 7 Tage nach dem letzten
    try:
        lo = datetime.fromisoformat(from_dt) - timedelta(days=30) if from_dt else None
    except ValueError:
        lo = None
    try:
        hi = datetime.fromisoformat(to_dt) + timedelta(days=7) if to_dt else None
    except ValueError:
        hi = None

    series = []
    for kw in matching_kws:
        q_td = TrendData.query.filter(TrendData.keyword_id == kw.id)
        if lo is not None:
            q_td = q_td.filter(TrendData.date >= lo)
        if hi is not None:
            q_td = q_td.filter(TrendData.date <= hi)
        rows = q_td.order_by(TrendData.date).all()

        # Auflösung erkennen: sub-tägliche Zeitstempel → stündliche Aggregation
        has_subday = any(
            r.date.hour != 0 or r.date.minute != 0 or r.date.second != 0
            for r in rows
        ) if rows else False

        bucket_max: dict = defaultdict(int)
        for r in rows:
            if has_subday:
                # Stündliche Bucket-Schlüssel: "YYYY-MM-DDTHH:MM"
                bucket = r.date.replace(minute=0, second=0, microsecond=0).strftime("%Y-%m-%dT%H:%M")
            else:
                bucket = r.date.date().isoformat()
            if r.value > bucket_max[bucket]:
                bucket_max[bucket] = r.value

        data = [{"date": d, "value": v} for d, v in sorted(bucket_max.items())]
        series.append({
            "id":        kw.id,
            "keyword":   kw.keyword,
            "geo":       kw.geo or "",
            "gprop":     kw.gprop or "",
            "timeframe": kw.timeframe or "",
            "data":      data,
        })

    return jsonify({
        "keywords": [kw.keyword for kw in matching_kws],
        "series":   series,
    })


# ---------------------------------------------------------------------------
# API – Fetch-Log
# ---------------------------------------------------------------------------

@app.route("/api/logs")
@login_required
def api_logs():
    per_page = 50
    page  = max(int(request.args.get("page", 1)), 1)
    total = FetchLog.query.filter_by(user_id=current_user.id).count()
    pages = max((total + per_page - 1) // per_page, 1)
    if page > pages:
        page = pages
    logs  = (FetchLog.query.filter_by(user_id=current_user.id)
             .order_by(FetchLog.started_at.desc())
             .offset((page - 1) * per_page)
             .limit(per_page)
             .all())
    return jsonify({
        "logs":  [l.to_dict() for l in logs],
        "total": total,
        "pages": pages,
        "page":  page,
    })


@app.route("/api/logs/clear", methods=["POST"])
@login_required
def api_logs_clear():
    """Löscht alle Fetch-Log-Einträge des aktuellen Users."""
    try:
        db.session.execute(text("DELETE FROM fetch_log WHERE user_id = :uid"), {"uid": current_user.id})
        db.session.commit()
        log.info("Fetch-Logs gelöscht")
        return jsonify({"ok": True})
    except Exception as e:
        db.session.rollback()
        log.error("logs/clear fehlgeschlagen: %s", e)
        abort(500, str(e))


@app.route("/api/translate", methods=["POST"])
@login_required
def api_translate():
    """
    Übersetzt ein Keyword in EN, ES, FR, ZH-CN, RU und erkennt die Quellsprache.
    Body: { "text": "Klimawandel" }
    """
    data = request.get_json(force=True) or {}
    text_input = (data.get("text") or "").strip()
    if not text_input:
        abort(400, "text fehlt")

    from translator import detect_language, translate_to_targets, LANG_NAMES
    detected = detect_language(text_input)
    translations = translate_to_targets(text_input)

    return jsonify({
        "detected_lang": detected,
        "detected_lang_name": LANG_NAMES.get(detected.lower(), detected),
        "translations": translations,
    })


@app.route("/api/translate-content", methods=["POST"])
@login_required
def api_translate_content():
    """
    Übersetzt beliebigen Text in eine Zielsprache.
    Body: { "text": "...", "target": "de" }
    """
    import re as _re
    data = request.get_json(force=True) or {}
    text_input = (data.get("text") or "").strip()
    target = (data.get("target") or "de").strip()
    if not text_input:
        abort(400, "text fehlt")
    if not _re.match(r'^[a-zA-Z]{2,3}(-[a-zA-Z]{2,4})?$', target):
        abort(400, "ungültiger target-Sprachcode")

    from translator import detect_language, LANG_NAMES
    from deep_translator import GoogleTranslator

    detected = detect_language(text_input)
    try:
        translated = GoogleTranslator(source="auto", target=target).translate(text_input[:5000])
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    return jsonify({
        "detected_lang": detected,
        "detected_lang_name": LANG_NAMES.get(detected.lower(), detected),
        "translated": translated or "",
        "target": target,
    })


@app.route("/api/fetch/trigger", methods=["POST"])
@login_required
def api_fetch_trigger():
    """
    Manuellen Fetch-Lauf anstoßen (läuft im Hintergrund).
    Body (JSON, optional): { "run_tag": "mein-label" }
      run_tag leer oder fehlend → Standard-Abruf (überschreibt bestehende Werte).
      run_tag gesetzt            → neue parallele Abruf-Reihe.
    """
    data    = request.get_json(force=True, silent=True) or {}
    run_tag = (data.get("run_tag") or "").strip()[:100]
    uid     = current_user.id

    def _do_fetch():
        from fetcher import run_fetch
        log.info("Manueller Fetch-Lauf (run_tag='%s')", run_tag)
        try:
            result = run_fetch(app, run_tag=run_tag, user_id=uid)
            log.info("Manueller Fetch abgeschlossen: %s", result)
        except Exception as exc:
            log.error("Manueller Fetch-Fehler: %s", exc)

    scheduler.add_job(_do_fetch, id="manual_fetch", replace_existing=True)
    audit_log("keyword_fetch", "keyword", None,
              f"Manueller Fetch gestartet (run_tag='{run_tag}')")
    return jsonify({"ok": True, "message": f"Fetch-Lauf gestartet (run_tag='{run_tag}')"})


@app.route("/api/fetch/trigger/<int:kw_id>", methods=["POST"])
@login_required
def api_fetch_trigger_single(kw_id):
    """Einzelnen Keyword-Abruf anstoßen (läuft im Hintergrund)."""
    kw = Keyword.query.filter_by(id=kw_id, user_id=current_user.id).first()
    if not kw:
        abort(404, "Keyword nicht gefunden")

    data    = request.get_json(force=True, silent=True) or {}
    run_tag = (data.get("run_tag") or "").strip()[:100]
    uid     = current_user.id

    def _do_fetch():
        from fetcher import run_fetch
        log.info("Einzelabruf für Keyword-ID %d ('%s'), run_tag='%s'", kw_id, kw.keyword, run_tag)
        try:
            result = run_fetch(app, keyword_ids=[kw_id], run_tag=run_tag, user_id=uid)
            log.info("Einzelabruf abgeschlossen: %s", result)
        except Exception as exc:
            log.error("Einzelabruf-Fehler: %s", exc)

    scheduler.add_job(_do_fetch, id=f"fetch_kw_{kw_id}", replace_existing=True)
    audit_log("keyword_fetch", "keyword", kw_id,
              f"Einzelabruf: {kw.keyword} (run_tag='{run_tag}')")
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# API – Live-Check (einmaliger Abruf ohne Speicherung, dann optional speichern)
# ---------------------------------------------------------------------------

@app.route("/api/live-check", methods=["POST"])
@login_required
def api_live_check():
    """
    Ruft Google-Trends-Daten für ein Keyword ab, ohne sie zu speichern.
    Body: { keyword, geo, timeframe, gprop }
    Gibt { ok, data: [{date, value}], backend, keyword, geo, timeframe, gprop } zurück.
    """
    import json as _json
    from fetcher import _fetch_single, NoDataError, DEFAULT_WORKFLOW, SERPAPI_KEY, COOKIES_FILE

    body     = request.get_json(force=True, silent=True) or {}
    keyword  = (body.get("keyword") or "").strip()
    geo      = (body.get("geo") if body.get("geo") is not None else "DE").upper()
    timeframe = (body.get("timeframe") or "now 7-d").strip()
    gprop    = (body.get("gprop") or "").strip()

    if not keyword:
        abort(400, "keyword fehlt")

    # Workflow / API-Key / Cookies aus App-Einstellungen lesen (wie run_fetch)
    workflow = _get_user_workflow(current_user)
    k_setting   = AppSetting.query.filter_by(key="serpapi_key", user_id=None).first()
    serpapi_key = (k_setting.value if k_setting and k_setting.value else None) or SERPAPI_KEY
    c_setting   = AppSetting.query.filter_by(key="cookies_file", user_id=None).first()
    cookies_file = (c_setting.value if c_setting and c_setting.value else None) or COOKIES_FILE
    hl = _get_query_language(current_user)

    try:
        raw, backend = _fetch_single(
            keyword, geo, timeframe, gprop,
            workflow=workflow, serpapi_key=serpapi_key, cookies_file=cookies_file, hl=hl,
        )
    except NoDataError as e:
        return jsonify({"ok": False, "error": str(e)}), 422
    except Exception as e:
        log.error("live-check Fehler: %s", e)
        return jsonify({"ok": False, "error": str(e)}), 500

    def fmt_dt(dt):
        if dt.hour == 0 and dt.minute == 0 and dt.second == 0:
            return dt.strftime("%Y-%m-%d")
        return dt.strftime("%Y-%m-%dT%H:%M")

    data_points = [{"date": fmt_dt(k), "value": v} for k, v in sorted(raw.items())]
    return jsonify({
        "ok": True, "backend": backend,
        "keyword": keyword, "geo": geo, "timeframe": timeframe, "gprop": gprop,
        "data": data_points,
    })


@app.route("/api/live-check/save", methods=["POST"])
@login_required
def api_live_check_save():
    """
    Speichert ein Live-Check-Ergebnis in die Datenbank.
    Body: { keyword, geo, timeframe, gprop, active, data: [{date, value}] }
    Legt ein neues Keyword an (oder nutzt ein bestehendes) und speichert die Datenpunkte.
    """
    body      = request.get_json(force=True, silent=True) or {}
    keyword   = (body.get("keyword") or "").strip()
    geo       = (body.get("geo") if body.get("geo") is not None else "DE").upper()
    timeframe = (body.get("timeframe") or "").strip()
    gprop     = (body.get("gprop") or "").strip()
    active      = bool(body.get("active", False))   # Standard: inaktiv
    points      = body.get("data") or []
    project_ids = body.get("project_ids") or []

    if not keyword:
        abort(400, "keyword fehlt")

    # Keyword anlegen oder vorhandenes verwenden (alle 4 Suchparameter identifizieren ein Keyword)
    kw = Keyword.query.filter_by(keyword=keyword, timeframe=timeframe,
                                 geo=geo, gprop=gprop, user_id=current_user.id).first()
    if kw:
        kw.active = active
    else:
        kw = Keyword(keyword=keyword, geo=geo, timeframe=timeframe, gprop=gprop, active=active, user_id=current_user.id)
        db.session.add(kw)
    db.session.flush()  # kw.id verfügbar machen

    # Projektzuordnung
    for pid in project_ids:
        proj = Project.query.filter_by(id=int(pid), user_id=current_user.id).first()
        if proj and proj not in kw.kw_projects:
            kw.kw_projects.append(proj)

    run_tag = datetime.now(timezone.utc).strftime("%d.%m.%y %H:%M")
    saved = 0
    for p in points:
        try:
            raw_date = p["date"]
            if "T" in raw_date:
                dt = datetime.strptime(raw_date, "%Y-%m-%dT%H:%M")
            else:
                dt = datetime.strptime(raw_date, "%Y-%m-%d")
            val = int(p["value"])
            existing = TrendData.query.filter_by(
                keyword_id=kw.id, date=dt, run_tag=run_tag
            ).first()
            if existing:
                existing.value = val
            else:
                db.session.add(TrendData(keyword_id=kw.id, date=dt, value=val, run_tag=run_tag))
            saved += 1
        except Exception:
            pass

    db.session.commit()
    log.info("live-check/save: Keyword '%s' (id=%d), %d Punkte gespeichert", keyword, kw.id, saved)

    # Verwandte Suchanfragen im Hintergrund abrufen (blockiert die Antwort nicht)
    kw_id_saved   = kw.id
    kw_text_saved = keyword
    geo_saved     = geo
    tf_saved      = timeframe
    gprop_saved   = gprop
    user_saved    = current_user._get_current_object()

    def _fetch_rq_background():
        import json as _json
        from fetcher import _fetch_related, DEFAULT_WORKFLOW, SERPAPI_KEY, COOKIES_FILE
        from models import RelatedQuery
        with app.app_context():
            try:
                workflow = _get_user_workflow(user_saved)
                k_s = AppSetting.query.filter_by(key="serpapi_key", user_id=None).first()
                serpapi_key = (k_s.value if k_s and k_s.value else None) or SERPAPI_KEY
                c_s = AppSetting.query.filter_by(key="cookies_file", user_id=None).first()
                cookies_file = (c_s.value if c_s and c_s.value else None) or COOKIES_FILE
                hl = _get_query_language(user_saved)

                rq_data = _fetch_related(
                    kw_text_saved, geo_saved, tf_saved, gprop_saved,
                    workflow=workflow, serpapi_key=serpapi_key, cookies_file=cookies_file, hl=hl,
                )
                from datetime import timezone as _tz
                now = datetime.now(_tz.utc)
                count = 0
                for qt in ("rising", "top"):
                    for i, item in enumerate(rq_data.get(qt, [])):
                        q = (item.get("query") or "").strip()
                        if not q:
                            continue
                        db.session.add(RelatedQuery(
                            keyword_id=kw_id_saved, query_type=qt,
                            query=q, value=str(item.get("value", "")),
                            rank=i, fetched_at=now,
                        ))
                        count += 1
                db.session.commit()
                log.info("live-check/save: %d Related Queries für '%s' gespeichert",
                         count, kw_text_saved)
            except Exception as rq_err:
                db.session.rollback()
                log.warning("live-check/save: Related Queries für '%s' fehlgeschlagen: %s",
                            kw_text_saved, rq_err)

    import threading
    threading.Thread(target=_fetch_rq_background, daemon=True).start()

    return jsonify({"ok": True, "keyword_id": kw.id, "saved": saved})


# ---------------------------------------------------------------------------
# API – Projekte
# ---------------------------------------------------------------------------

@app.route("/api/projects", methods=["GET"])
@login_required
def api_projects_list():
    return jsonify([p.to_dict() for p in
                    Project.query.filter_by(user_id=current_user.id).order_by(Project.sort_order, Project.id).all()])


@app.route("/api/projects", methods=["POST"])
@login_required
def api_projects_create():
    body  = request.get_json(force=True, silent=True) or {}
    name  = (body.get("name") or "").strip()
    if not name:
        abort(400, "name fehlt")
    if current_user.max_projects and current_user.max_projects > 0:
        count = Project.query.filter_by(user_id=current_user.id).count()
        if count >= current_user.max_projects:
            return jsonify({"error": f"Projektlimit erreicht ({current_user.max_projects})"}), 403
    max_order = db.session.query(db.func.max(Project.sort_order)).filter(Project.user_id == current_user.id).scalar() or 0
    proj = Project(
        name        = name,
        description = (body.get("description") or "").strip(),
        briefing    = (body.get("briefing") or "").strip(),
        color       = (body.get("color") or "#4f8ef7").strip(),
        sort_order  = max_order + 1,
        user_id     = current_user.id,
    )
    db.session.add(proj)
    db.session.commit()
    audit_log("project_create", "project", proj.id, f"Projekt erstellt: {name}",
              project_id=proj.id)
    return jsonify({"ok": True, "project": proj.to_dict()})


@app.route("/api/projects/<int:pid>", methods=["PUT"])
@login_required
def api_projects_update(pid):
    proj = Project.query.filter_by(id=pid, user_id=current_user.id).first_or_404()
    body = request.get_json(force=True, silent=True) or {}
    if "name" in body:
        name = (body["name"] or "").strip()
        if not name:
            abort(400, "name darf nicht leer sein")
        proj.name = name
    if "description" in body:
        proj.description = (body["description"] or "").strip()
    if "color" in body:
        proj.color = (body["color"] or "#4f8ef7").strip()
    if "briefing" in body:
        proj.briefing = (body["briefing"] or "").strip()
    db.session.commit()
    return jsonify({"ok": True, "project": proj.to_dict()})


@app.route("/api/projects/<int:pid>", methods=["DELETE"])
@login_required
def api_projects_delete(pid):
    proj = Project.query.filter_by(id=pid, user_id=current_user.id).first_or_404()
    # Slides gehören fest zum Projekt → löschen
    for slide in list(proj.slides):
        db.session.delete(slide)
    # Snapshots werden nicht gelöscht – Zuweisung wird aufgehoben
    for s in proj.snapshots:
        s.project_id = None
    proj_name = proj.name
    proj_id = proj.id
    db.session.delete(proj)
    db.session.commit()
    audit_log("project_delete", "project", proj_id, f"Projekt gelöscht: {proj_name}")
    return jsonify({"ok": True})


@app.route("/api/projects/reorder", methods=["POST"])
@login_required
def api_projects_reorder():
    """Erwartet { ids: [1, 3, 2, …] } – speichert sort_order entsprechend."""
    body = request.get_json(force=True, silent=True) or {}
    ids  = body.get("ids", [])
    if not isinstance(ids, list):
        abort(400, "ids muss eine Liste sein")
    for i, pid in enumerate(ids):
        Project.query.filter_by(id=int(pid), user_id=current_user.id).update({"sort_order": i})
    db.session.commit()
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# API – Snapshots
# ---------------------------------------------------------------------------

@app.route("/api/snapshots", methods=["GET"])
@login_required
def api_snapshots_list():
    uid = current_user.id
    pid = request.args.get("project_id", type=int)
    q = Snapshot.query.filter_by(user_id=uid)
    if pid is not None:
        q = q.filter_by(project_id=pid)
    import json as _j
    snaps = q.order_by(Snapshot.sort_order, Snapshot.created_at.desc()).all()
    def _markers(s):
        raw = _j.loads(s.markers_json or "[]")
        return raw
    def _has_temporal(s):
        try:
            cj = _j.loads(s.chart_json or "{}")
            t  = cj.get("temporal", {})
            return bool(t and isinstance(t, dict) and t)
        except Exception:
            return False
    def _has_correlations(s):
        try:
            cj = _j.loads(s.chart_json or "{}")
            c  = cj.get("correlations", [])
            return bool(c and isinstance(c, list) and c)
        except Exception:
            return False
    return jsonify([{
        "id":               s.id,
        "title":            s.title or "",
        "comment":          s.comment or "",
        "markers":          _markers(s),
        "markers_count":    len(_markers(s)),
        "has_temporal":     _has_temporal(s),
        "has_correlations": _has_correlations(s),
        "project_id":       s.project_id,
        "project_name":  s.project.name if s.project else None,
        "sort_order":    s.sort_order,
        "created_at":    s.created_at.isoformat() if s.created_at else None,
    } for s in snaps])


@app.route("/api/snapshots/reorder", methods=["POST"])
@login_required
def api_snapshots_reorder():
    """Erwartet { ids: [3, 1, 2, …] } – setzt sort_order für diese Snapshots."""
    body = request.get_json(force=True, silent=True) or {}
    ids  = body.get("ids", [])
    if not isinstance(ids, list):
        abort(400, "ids muss eine Liste sein")
    for i, sid in enumerate(ids):
        Snapshot.query.filter_by(id=int(sid), user_id=current_user.id).update({"sort_order": i})
    db.session.commit()
    return jsonify({"ok": True})


@app.route("/api/snapshots/<int:sid>", methods=["GET"])
@login_required
def api_snapshot_get(sid):
    """Vollständige Snapshot-Daten (inkl. chart_json und markers_json)."""
    snap = Snapshot.query.filter_by(id=sid, user_id=current_user.id).first_or_404()
    return jsonify(snap.to_dict())


@app.route("/api/snapshots/<int:sid>", methods=["PUT"])
@login_required
def api_snapshot_update(sid):
    snap = Snapshot.query.filter_by(id=sid, user_id=current_user.id).first_or_404()
    body = request.get_json(force=True, silent=True) or {}
    if "title" in body:
        snap.title = (body["title"] or "").strip()
    if "comment" in body:
        snap.comment = (body["comment"] or "").strip()
    if "project_id" in body:
        snap.project_id = int(body["project_id"]) if body["project_id"] else None
    if body.get("clear_temporal") or body.get("clear_correlations"):
        import json as _jt
        try:
            cj = _jt.loads(snap.chart_json or "{}")
            if body.get("clear_temporal"):
                cj["temporal"] = {}
            if body.get("clear_correlations"):
                cj["correlations"] = []
            snap.chart_json = _jt.dumps(cj, ensure_ascii=False)
        except Exception:
            pass
    snap.compute_hash()
    db.session.commit()
    return jsonify({"ok": True, "snapshot": {
        "id":           snap.id,
        "title":        snap.title,
        "comment":      snap.comment,
        "project_id":   snap.project_id,
        "project_name": snap.project.name if snap.project else None,
        "content_hash": snap.content_hash,
    }})


@app.route("/api/snapshots", methods=["POST"])
@login_required
def api_snapshots_save():
    """Speichert einen Chart-Snapshot mit Markierungen und Kommentar."""
    import json as _j
    body = request.get_json(force=True, silent=True) or {}
    if body.get("type") == "analysis":
        chart_data = {
            "type":          "analysis",
            "analysis_type": body.get("analysis_type", ""),
            "image":         body.get("image", ""),
            "subtitle":      body.get("subtitle", ""),
            "params":        body.get("params", {}),
        }
    else:
        chart_data = {
            "labels":        body.get("labels", []),
            "datasets":      body.get("datasets", []),
            "keywords_meta": body.get("keywords_meta", []),
            "visible_range": body.get("visible_range", {}),
            "correlations":  body.get("correlations", []),
            "temporal":      body.get("temporal", {}),
        }
    raw_pid    = body.get("project_id")
    project_id = int(raw_pid) if raw_pid else None

    # Snapshot ans Ende der Projektliste anfügen
    max_order = db.session.query(db.func.max(Snapshot.sort_order)).filter_by(project_id=project_id).scalar() or 0

    snap = Snapshot(
        title        = (body.get("title") or "").strip(),
        comment      = (body.get("comment") or "").strip(),
        chart_json   = _j.dumps(chart_data, ensure_ascii=False),
        markers_json = _j.dumps(body.get("markers", []), ensure_ascii=False),
        project_id   = project_id,
        sort_order   = max_order + 1,
        user_id      = current_user.id,
    )
    snap.compute_hash()
    db.session.add(snap)
    db.session.commit()
    log.info("snapshot saved: id=%d title='%s' hash=%s", snap.id, snap.title, snap.content_hash[:12])
    audit_log("snapshot_create", "snapshot", snap.id,
              f"Titel: {snap.title}", content_hash=snap.content_hash,
              project_id=snap.project_id)
    return jsonify({"ok": True, "id": snap.id, "content_hash": snap.content_hash})


@app.route("/api/snapshots/<int:sid>/verify", methods=["GET"])
@login_required
def api_snapshot_verify(sid):
    """Prüft die Integrität eines Snapshots anhand seines SHA-256-Hashes."""
    snap = Snapshot.query.filter_by(id=sid, user_id=current_user.id).first_or_404()
    result = snap.verify_hash()
    return jsonify({
        "id":            snap.id,
        "content_hash":  snap.content_hash or "",
        "verified":      result,  # True=intakt, False=manipuliert, None=kein Hash
        "created_at":    snap.created_at.isoformat() if snap.created_at else None,
    })


@app.route("/api/snapshots/<int:sid>", methods=["DELETE"])
@login_required
def api_snapshot_delete(sid):
    snap = Snapshot.query.filter_by(id=sid, user_id=current_user.id).first_or_404()
    snap_label = snap.keyword if hasattr(snap, 'keyword') else f"Snapshot #{sid}"
    snap_pid = snap.project_id
    db.session.delete(snap)
    db.session.commit()
    audit_log("snapshot_delete", "snapshot", sid, f"Snapshot gelöscht: {snap_label}",
              project_id=snap_pid)
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# API – Audit Trail
# ---------------------------------------------------------------------------

@app.route("/api/audit", methods=["GET"])
@login_required
def api_audit_list():
    """Gibt Audit-Trail-Einträge zurück (mit Filter- und Paging-Optionen)."""
    from models import AuditEntry
    if not _audit_enabled():
        return jsonify([])
    pid    = request.args.get("project_id", type=int)
    action = request.args.get("action")
    limit  = request.args.get("limit", 200, type=int)
    offset = request.args.get("offset", 0, type=int)
    q = AuditEntry.query.filter_by(user_id=current_user.id)
    if pid is not None:
        q = q.filter_by(project_id=pid)
    if action:
        q = q.filter_by(action=action)
    total = q.count()
    entries = q.order_by(AuditEntry.id.desc()).offset(offset).limit(limit).all()
    return jsonify({"total": total, "entries": [e.to_dict() for e in entries]})


@app.route("/api/audit/verify", methods=["GET"])
@login_required
def api_audit_verify():
    """Prüft die Integrität der Hash-Kette des Audit Trails."""
    from models import AuditEntry
    if not _audit_enabled():
        return jsonify({"ok": False, "error": "Audit Trail nicht aktiviert"})
    entries = AuditEntry.query.filter_by(user_id=current_user.id).order_by(AuditEntry.id).all()
    if not entries:
        return jsonify({"ok": True, "count": 0, "msg": "Keine Einträge vorhanden"})
    broken_at = None
    for i, entry in enumerate(entries):
        # Prüfe prev_hash-Verkettung
        expected_prev = entries[i - 1].entry_hash if i > 0 else "GENESIS"
        if entry.prev_hash != expected_prev:
            broken_at = entry.id
            break
        # Prüfe entry_hash-Integrität
        saved_hash = entry.entry_hash
        entry.compute_entry_hash()
        if entry.entry_hash != saved_hash:
            entry.entry_hash = saved_hash  # wiederherstellen
            broken_at = entry.id
            break
    if broken_at:
        return jsonify({"ok": False, "count": len(entries),
                        "broken_at": broken_at,
                        "msg": f"Integritätsverletzung bei Eintrag #{broken_at}"})
    return jsonify({"ok": True, "count": len(entries),
                    "msg": f"Hash-Kette intakt ({len(entries)} Einträge verifiziert)"})


@app.route("/api/audit/actions", methods=["GET"])
@login_required
def api_audit_actions():
    """Gibt alle vorkommenden Aktionstypen zurück (für Filter-Dropdown)."""
    from models import AuditEntry
    if not _audit_enabled():
        return jsonify([])
    actions = db.session.query(db.distinct(AuditEntry.action)).filter_by(
        user_id=current_user.id
    ).all()
    return jsonify(sorted([a[0] for a in actions]))


# ---------------------------------------------------------------------------
# API – Slides (Zwischenseiten)
# ---------------------------------------------------------------------------

@app.route("/api/slides", methods=["GET"])
@login_required
def api_slides_list():
    pid = request.args.get("project_id", type=int)
    own_pids = [p.id for p in Project.query.filter_by(user_id=current_user.id).all()]
    q = Slide.query.filter(Slide.project_id.in_(own_pids))
    if pid is not None:
        q = q.filter_by(project_id=pid)
    slides = q.order_by(Slide.sort_order, Slide.created_at).all()
    return jsonify([s.to_dict() for s in slides])


@app.route("/api/slides", methods=["POST"])
@login_required
def api_slides_create():
    body       = request.get_json(force=True, silent=True) or {}
    pid        = body.get("project_id")
    slide_type = body.get("slide_type", "section")
    if not pid:
        abort(400, "project_id erforderlich")

    pid = int(pid)
    Project.query.filter_by(id=pid, user_id=current_user.id).first_or_404()

    custom_order = body.get("sort_order")
    if custom_order is not None:
        new_order = float(custom_order)
    elif slide_type == "title":
        # Titelseite → ganz oben einfügen (sort_order unterhalb aller vorhandenen Einträge)
        min_snap  = db.session.query(db.func.min(Snapshot.sort_order)).filter_by(project_id=pid).scalar()
        min_slide = db.session.query(db.func.min(Slide.sort_order)).filter_by(project_id=pid).scalar()
        min_vals  = [v for v in [min_snap, min_slide] if v is not None]
        new_order = (min(min_vals) - 1) if min_vals else 0
    else:
        # Alle anderen Typen → ans Ende anhängen
        max_snap  = db.session.query(db.func.max(Snapshot.sort_order)).filter_by(project_id=pid).scalar() or 0
        max_slide = db.session.query(db.func.max(Slide.sort_order)).filter_by(project_id=pid).scalar() or 0
        new_order = max(max_snap, max_slide) + 1

    slide = Slide(
        project_id  = pid,
        slide_type  = slide_type,
        title       = (body.get("title") or "").strip(),
        description = (body.get("description") or "").strip(),
        content     = (body.get("content") or "").strip(),
        sort_order  = new_order,
    )
    db.session.add(slide)
    db.session.commit()
    return jsonify({"ok": True, "slide": slide.to_dict()})


@app.route("/api/slides/<int:sid>", methods=["PUT"])
@login_required
def api_slides_update(sid):
    slide = Slide.query.get_or_404(sid)
    Project.query.filter_by(id=slide.project_id, user_id=current_user.id).first_or_404()
    body  = request.get_json(force=True, silent=True) or {}
    if "slide_type"   in body: slide.slide_type   = body["slide_type"]
    if "title"        in body: slide.title        = (body["title"] or "").strip()
    if "description"  in body: slide.description  = (body["description"] or "").strip()
    if "content"      in body: slide.content      = (body["content"] or "").strip()
    db.session.commit()
    return jsonify({"ok": True, "slide": slide.to_dict()})


@app.route("/api/slides/<int:sid>", methods=["DELETE"])
@login_required
def api_slides_delete(sid):
    slide = Slide.query.get_or_404(sid)
    Project.query.filter_by(id=slide.project_id, user_id=current_user.id).first_or_404()
    db.session.delete(slide)
    db.session.commit()
    return jsonify({"ok": True})


@app.route("/api/slides/<int:sid>/screenshot", methods=["POST"])
@login_required
def api_slide_screenshot(sid):
    """Erstellt einen Screenshot der gespeicherten URL (nur für website-Slides)."""
    import json as _j, base64
    slide = Slide.query.get_or_404(sid)
    if slide.slide_type != "website":
        abort(400, "Nur für Website-Slides verfügbar")

    content = {}
    try:
        content = _j.loads(slide.content or "{}")
    except Exception:
        pass

    url = (content.get("url") or "").strip()
    if not url:
        abort(400, "Keine URL gespeichert – bitte zuerst speichern")
    if not url.startswith(("http://", "https://")):
        url = "https://" + url

    try:
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page    = browser.new_page(viewport={"width": 1280, "height": 800})
            page.goto(url, wait_until="networkidle", timeout=30000)
            screenshot_bytes = page.screenshot(type="jpeg", quality=80, full_page=False)
            browser.close()
        content["screenshot_b64"] = base64.b64encode(screenshot_bytes).decode()
        content["screenshot_at"] = datetime.now(timezone.utc).isoformat()
        slide.content = _j.dumps(content)
        db.session.commit()
        return jsonify({"ok": True, "slide": slide.to_dict()})
    except Exception as e:
        log.error("Screenshot fehlgeschlagen für slide %s: %s", sid, e)
        abort(500, f"Screenshot fehlgeschlagen: {e}")


@app.route("/api/screenshot", methods=["POST"])
@login_required
def api_screenshot():
    """Erstellt einen Screenshot einer URL und gibt die Base64-JPEG-Daten zurück."""
    import base64
    data = request.get_json(force=True) or {}
    url = (data.get("url") or "").strip()
    if not url:
        abort(400, "url fehlt")
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    try:
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page    = browser.new_page(viewport={"width": 1280, "height": 800})
            page.goto(url, wait_until="networkidle", timeout=30000)
            screenshot_bytes = page.screenshot(type="jpeg", quality=80, full_page=False)
            browser.close()
        b64 = base64.b64encode(screenshot_bytes).decode()
        return jsonify({"ok": True, "b64": b64})
    except Exception as e:
        log.error("Screenshot fehlgeschlagen für %s: %s", url, e)
        abort(500, f"Screenshot fehlgeschlagen: {e}")


@app.route("/api/projects/<int:pid>/items/reorder", methods=["POST"])
@login_required
def api_project_items_reorder(pid):
    """Unified reorder for snapshots + slides.
    Body: { items: [{type: 'snapshot'|'slide', id: N}, …] }
    """
    Project.query.filter_by(id=pid, user_id=current_user.id).first_or_404()
    body  = request.get_json(force=True, silent=True) or {}
    items = body.get("items", [])
    if not isinstance(items, list):
        abort(400, "items muss eine Liste sein")
    for i, item in enumerate(items):
        itype = item.get("type")
        iid   = int(item.get("id", 0))
        if itype == "snapshot":
            Snapshot.query.filter_by(id=iid, user_id=current_user.id).update({"sort_order": i})
        elif itype == "slide":
            Slide.query.filter_by(id=iid).update({"sort_order": i})
    db.session.commit()
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# API – Statistik
# ---------------------------------------------------------------------------

@app.route("/api/stats")
@login_required
def api_stats():
    """Abruf-Statistiken für die Statistik-Seite."""
    from sqlalchemy import text as sqtext

    all_sessions = FetchLog.query.filter_by(user_id=current_user.id).order_by(FetchLog.started_at).all()
    sessions = all_sessions[-50:]  # Balkendiagramm: nur letzte 50 Sitzungen

    # TrendData-Zeilen gruppiert nach Tag (nur eigene Keywords)
    own_kw_ids = [k.id for k in Keyword.query.filter_by(user_id=current_user.id).all()]
    if own_kw_ids:
        placeholders = ",".join(str(int(i)) for i in own_kw_ids)
        rows = db.session.execute(sqtext(
            f"SELECT date(fetched_at) AS day, count(*) AS cnt "
            f"FROM trend_data WHERE keyword_id IN ({placeholders}) GROUP BY date(fetched_at) ORDER BY day"
        )).fetchall()
        total_records = db.session.execute(sqtext(
            f"SELECT count(*) FROM trend_data WHERE keyword_id IN ({placeholders})"
        )).scalar() or 0
        total_related = db.session.execute(sqtext(
            f"SELECT count(*) FROM related_queries WHERE keyword_id IN ({placeholders})"
        )).scalar() or 0
    else:
        rows = []
        total_records = 0
        total_related = 0
    data_by_day = [{"day": r[0], "count": r[1]} for r in rows]
    by_backend: dict = {}
    for s in all_sessions:
        b = s.backend or "unbekannt"
        by_backend[b] = by_backend.get(b, 0) + 1

    llm_used, llm_limit = _get_llm_usage(current_user.id)

    return jsonify({
        "sessions":    [s.to_dict() for s in sessions],
        "data_by_day": data_by_day,
        "summary": {
            "total_records":   total_records,
            "total_related":   total_related,
            "fetch_sessions":  len(all_sessions),
            "by_backend":      by_backend,
            "llm_used":        llm_used,
            "llm_limit":       llm_limit,
            "uses_own_llm":    _user_uses_own_llm_key(current_user.id),
        },
    })


@app.route("/api/stats/llm-log")
@login_required
def api_llm_log():
    """Detaillierte Auflistung der LLM-Aufrufe des aktuellen Monats."""
    from models import LlmLog
    month = request.args.get("month", datetime.now().strftime("%Y-%m"))
    try:
        year, mon = month.split("-")
        start = datetime(int(year), int(mon), 1, tzinfo=timezone.utc)
        if int(mon) == 12:
            end = datetime(int(year) + 1, 1, 1, tzinfo=timezone.utc)
        else:
            end = datetime(int(year), int(mon) + 1, 1, tzinfo=timezone.utc)
    except Exception:
        start = datetime(datetime.now().year, datetime.now().month, 1, tzinfo=timezone.utc)
        end = datetime(datetime.now().year, datetime.now().month + 1, 1, tzinfo=timezone.utc) if datetime.now().month < 12 else datetime(datetime.now().year + 1, 1, 1, tzinfo=timezone.utc)

    logs = LlmLog.query.filter(
        LlmLog.user_id == current_user.id,
        LlmLog.created_at >= start,
        LlmLog.created_at < end,
    ).order_by(LlmLog.created_at.desc()).all()

    # Gruppierung nach Source
    by_source = {}
    for l in logs:
        s = l.source or "sonstige"
        by_source[s] = by_source.get(s, 0) + 1

    return jsonify({
        "month": month,
        "total": len(logs),
        "by_source": by_source,
        "logs": [l.to_dict() for l in logs],
    })


@app.route("/api/data/delete-all", methods=["POST"])
@login_required
def api_delete_all():
    """Löscht alle Trend-, Related-Query- und FetchLog-Daten des aktuellen Users."""
    try:
        own_kw_ids = [k.id for k in Keyword.query.filter_by(user_id=current_user.id).all()]
        if own_kw_ids:
            placeholders = ",".join(str(int(i)) for i in own_kw_ids)
            db.session.execute(text(f"DELETE FROM trend_data WHERE keyword_id IN ({placeholders})"))
            db.session.execute(text(f"DELETE FROM related_queries WHERE keyword_id IN ({placeholders})"))
            db.session.execute(text(f"DELETE FROM region_interest WHERE keyword_id IN ({placeholders})"))
        db.session.execute(text("DELETE FROM fetch_log WHERE user_id = :uid"), {"uid": current_user.id})
        db.session.commit()
        log.info("Alle Daten gelöscht (DELETE ALL)")
        return jsonify({"ok": True})
    except Exception as e:
        db.session.rollback()
        log.error("delete-all fehlgeschlagen: %s", e)
        abort(500, str(e))


# ---------------------------------------------------------------------------
# API – Alerts
# ---------------------------------------------------------------------------

@app.route("/api/alerts", methods=["GET"])
@login_required
def api_alerts_list():
    return jsonify([a.to_dict() for a in Alert.query.filter_by(user_id=current_user.id).order_by(Alert.created_at.desc()).all()])


@app.route("/api/alerts", methods=["POST"])
@login_required
def api_alerts_create():
    data = request.get_json(force=True) or {}
    name = (data.get("name") or "").strip()
    atype = (data.get("alert_type") or "").strip()
    if not atype or atype not in ("occurrence", "disappearance", "spike", "volume_rise", "volume_drop"):
        abort(400, "alert_type muss occurrence, disappearance, spike, volume_rise oder volume_drop sein")
    kw_ids = data.get("keyword_ids") or []
    alert = Alert(
        name=name or atype,
        alert_type=atype,
        watch_term=(data.get("watch_term") or "").strip(),
        keyword_ids_json=_json.dumps(kw_ids) if kw_ids else "",
        spike_threshold=int(data.get("spike_threshold") or 20),
        spike_threshold_type=(data.get("spike_threshold_type") or "percent"),
        spike_hours=int(data.get("spike_hours") or 24),
        comment=(data.get("comment") or "").strip(),
        active=bool(data.get("active", True)),
        user_id=current_user.id,
    )
    db.session.add(alert)
    db.session.commit()
    return jsonify(alert.to_dict()), 201


@app.route("/api/alerts/<int:aid>", methods=["PUT"])
@login_required
def api_alerts_update(aid):
    alert = Alert.query.filter_by(id=aid, user_id=current_user.id).first()
    if not alert:
        abort(404)
    data = request.get_json(force=True) or {}
    if "name"                 in data: alert.name                 = (data["name"] or "").strip()
    if "alert_type"           in data: alert.alert_type           = data["alert_type"]
    if "watch_term"           in data: alert.watch_term           = (data["watch_term"] or "").strip()
    if "keyword_ids"          in data:
        kw_ids = data["keyword_ids"] or []
        alert.keyword_ids_json = _json.dumps(kw_ids) if kw_ids else ""
    if "spike_threshold"      in data: alert.spike_threshold      = int(data["spike_threshold"])
    if "spike_threshold_type" in data: alert.spike_threshold_type = data["spike_threshold_type"]
    if "spike_hours"          in data: alert.spike_hours          = int(data["spike_hours"])
    if "comment"              in data: alert.comment              = (data["comment"] or "").strip()
    if "active"               in data: alert.active               = bool(data["active"])
    db.session.commit()
    return jsonify(alert.to_dict())


@app.route("/api/alerts/<int:aid>", methods=["DELETE"])
@login_required
def api_alerts_delete(aid):
    alert = Alert.query.filter_by(id=aid, user_id=current_user.id).first()
    if not alert:
        abort(404)
    db.session.delete(alert)
    db.session.commit()
    return jsonify({"ok": True})


@app.route("/api/alerts/events", methods=["GET"])
@login_required
def api_alert_events():
    limit   = min(int(request.args.get("limit", 100)), 500)
    only_unseen = request.args.get("unseen") == "1"
    own_alert_ids = [a.id for a in Alert.query.filter_by(user_id=current_user.id).all()]
    q = AlertEvent.query.filter(AlertEvent.alert_id.in_(own_alert_ids)).order_by(AlertEvent.triggered_at.desc())
    if only_unseen:
        q = q.filter_by(seen=False)
    return jsonify([e.to_dict() for e in q.limit(limit).all()])


@app.route("/api/alerts/events/seen-all", methods=["POST"])
@login_required
def api_alert_events_seen_all():
    own_alert_ids = [a.id for a in Alert.query.filter_by(user_id=current_user.id).all()]
    AlertEvent.query.filter(AlertEvent.alert_id.in_(own_alert_ids), AlertEvent.seen == False).update({"seen": True}, synchronize_session="fetch")
    db.session.commit()
    return jsonify({"ok": True})


@app.route("/api/alerts/events/<int:eid>/seen", methods=["POST"])
@login_required
def api_alert_event_seen(eid):
    ev = db.session.get(AlertEvent, eid)
    if not ev:
        abort(404)
    ev.seen = True
    db.session.commit()
    return jsonify({"ok": True})


@app.route("/api/dashboard/summary", methods=["GET"])
@login_required
def api_dashboard_summary():
    own_alert_ids = [a.id for a in Alert.query.filter_by(user_id=current_user.id).all()]
    unseen = AlertEvent.query.filter(AlertEvent.alert_id.in_(own_alert_ids), AlertEvent.seen == False).count()
    total  = AlertEvent.query.filter(AlertEvent.alert_id.in_(own_alert_ids)).count()
    return jsonify({"unseen": unseen, "total": total})


@app.route("/api/alerts/evaluate", methods=["POST"])
@login_required
def api_alerts_evaluate():
    from alerts import evaluate_alerts as _eval
    _eval(app)
    return jsonify({"ok": True})


@app.route("/api/proxy-image")
@login_required
def api_proxy_image():
    """Server-seitiger Proxy für externe Bilder (umgeht CORS-Beschränkungen im Browser)."""
    import urllib.request
    from flask import Response as FlaskResponse
    url = request.args.get("url", "").strip()
    if not url.startswith(("http://", "https://")):
        abort(400)
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=8) as r:
            data = r.read()
            ct = r.headers.get_content_type() or "image/png"
        return FlaskResponse(data, content_type=ct)
    except Exception:
        abort(502)


# ---------------------------------------------------------------------------
# Fehlerbehandlung
# ---------------------------------------------------------------------------

@app.errorhandler(400)
def bad_request(e):
    return jsonify({"error": str(e)}), 400


@app.errorhandler(404)
def not_found(e):
    return jsonify({"error": "Nicht gefunden"}), 404


@app.errorhandler(409)
def conflict(e):
    return jsonify({"error": str(e)}), 409


# ---------------------------------------------------------------------------
# Startup
# ---------------------------------------------------------------------------

with app.app_context():
    db.create_all()
    # Schema-Migrationen
    with db.engine.connect() as conn:
        # fetch_log: backend-Spalte
        try:
            conn.execute(text("ALTER TABLE fetch_log ADD COLUMN backend VARCHAR(20) DEFAULT ''"))
            conn.commit()
        except Exception:
            pass

        # Keywords: neue Spalten falls nicht vorhanden
        for col, definition in [
            ("detected_lang",  "VARCHAR(10)  DEFAULT ''"),
            ("translation_de", "VARCHAR(255) DEFAULT ''"),
            ("timeframe",      "VARCHAR(30)  DEFAULT ''"),
            ("gprop",          "VARCHAR(20)  DEFAULT ''"),
        ]:
            try:
                conn.execute(text(f"ALTER TABLE keywords ADD COLUMN {col} {definition}"))
                conn.commit()
            except Exception:
                pass  # Spalte existiert bereits

        # trend_data.date: YYYY-MM-DD → YYYY-MM-DD 00:00:00 (für DateTime-Kompatibilität)
        try:
            conn.execute(text(
                'UPDATE trend_data SET "date" = "date" || \' 00:00:00\' '
                'WHERE "date" NOT LIKE \'% %\' AND "date" NOT LIKE \'%T%\''
            ))
            conn.commit()
        except Exception:
            pass

        # trend_data: run_tag-Spalte + neue Unique-Constraint (Tabellenneubau für SQLite)
        try:
            cols = [r[1] for r in conn.execute(text("PRAGMA table_info(trend_data)")).fetchall()]
            if "run_tag" not in cols:
                log.info("Migriere trend_data → run_tag + neue Unique-Constraint …")
                conn.execute(text("""
                    CREATE TABLE trend_data_v2 (
                        id         INTEGER PRIMARY KEY AUTOINCREMENT,
                        keyword_id INTEGER NOT NULL REFERENCES keywords(id),
                        date       DATETIME NOT NULL,
                        value      INTEGER  NOT NULL,
                        fetched_at DATETIME,
                        run_tag    VARCHAR(100) NOT NULL DEFAULT '',
                        UNIQUE(keyword_id, date, run_tag)
                    )
                """))
                conn.execute(text("""
                    INSERT INTO trend_data_v2 (id, keyword_id, date, value, fetched_at, run_tag)
                    SELECT id, keyword_id, date, value, fetched_at, '' FROM trend_data
                """))
                conn.execute(text("DROP TABLE trend_data"))
                conn.execute(text("ALTER TABLE trend_data_v2 RENAME TO trend_data"))
                conn.commit()
                log.info("Migration trend_data abgeschlossen")
        except Exception as e:
            try: conn.rollback()
            except Exception: pass
            log.error("Migration trend_data fehlgeschlagen: %s", e)

        # Keywords: UNIQUE(keyword) → UNIQUE(keyword, timeframe) — Tabellenneubau in SQLite
        try:
            idx_names = [r[1] for r in conn.execute(text("PRAGMA index_list(keywords)")).fetchall()]
            if "uq_keyword_tf" not in idx_names and "uq_keyword_tf_geo_gprop" not in idx_names:
                log.info("Migriere keywords → composite unique (keyword, timeframe) …")
                conn.execute(text("""
                    CREATE TABLE keywords_v2 (
                        id             INTEGER PRIMARY KEY AUTOINCREMENT,
                        keyword        VARCHAR(255) NOT NULL,
                        geo            VARCHAR(10)  NOT NULL DEFAULT 'DE',
                        active         BOOLEAN      NOT NULL DEFAULT 1,
                        created_at     DATETIME,
                        note           TEXT         DEFAULT '',
                        detected_lang  VARCHAR(10)  DEFAULT '',
                        translation_de VARCHAR(255) DEFAULT '',
                        timeframe      VARCHAR(30)  DEFAULT '',
                        gprop          VARCHAR(20)  DEFAULT ''
                    )
                """))
                conn.execute(text("INSERT INTO keywords_v2 SELECT * FROM keywords"))
                conn.execute(text("DROP TABLE keywords"))
                conn.execute(text("ALTER TABLE keywords_v2 RENAME TO keywords"))
                conn.execute(text(
                    "CREATE UNIQUE INDEX uq_keyword_tf ON keywords(keyword, timeframe)"
                ))
                conn.commit()
                log.info("Migration keywords (uq_keyword_tf) abgeschlossen")
        except Exception as ex:
            try: conn.rollback()
            except Exception: pass
            log.warning("Migration keywords (unique): %s", ex)

        # Keywords: UNIQUE(keyword, timeframe) → UNIQUE(keyword, timeframe, geo, gprop)
        try:
            idx_names = [r[1] for r in conn.execute(text("PRAGMA index_list(keywords)")).fetchall()]
            if "uq_keyword_tf_geo_gprop" not in idx_names:
                log.info("Migriere keywords → composite unique (keyword, timeframe, geo, gprop) …")
                conn.execute(text("""
                    CREATE TABLE keywords_v3 (
                        id             INTEGER PRIMARY KEY AUTOINCREMENT,
                        keyword        VARCHAR(255) NOT NULL,
                        geo            VARCHAR(10)  NOT NULL DEFAULT 'DE',
                        active         BOOLEAN      NOT NULL DEFAULT 1,
                        created_at     DATETIME,
                        note           TEXT         DEFAULT '',
                        detected_lang  VARCHAR(10)  DEFAULT '',
                        translation_de VARCHAR(255) DEFAULT '',
                        timeframe      VARCHAR(30)  DEFAULT '',
                        gprop          VARCHAR(20)  DEFAULT ''
                    )
                """))
                conn.execute(text("INSERT INTO keywords_v3 SELECT * FROM keywords"))
                conn.execute(text("DROP TABLE keywords"))
                conn.execute(text("ALTER TABLE keywords_v3 RENAME TO keywords"))
                conn.execute(text(
                    "CREATE UNIQUE INDEX uq_keyword_tf_geo_gprop ON keywords(keyword, timeframe, geo, gprop)"
                ))
                conn.commit()
                log.info("Migration keywords (uq_keyword_tf_geo_gprop) abgeschlossen")
        except Exception as ex:
            try: conn.rollback()
            except Exception: pass
            log.warning("Migration keywords (uq_keyword_tf_geo_gprop): %s", ex)

        # alerts: comment-Spalte
        try:
            conn.execute(text("ALTER TABLE alerts ADD COLUMN comment TEXT DEFAULT ''"))
            conn.commit()
        except Exception:
            pass  # Spalte existiert bereits

        # snapshots: project_id-Spalte
        try:
            conn.execute(text("ALTER TABLE snapshots ADD COLUMN project_id INTEGER REFERENCES projects(id)"))
            conn.commit()
        except Exception:
            pass  # Spalte existiert bereits oder Tabelle noch nicht vorhanden

        # projects: sort_order-Spalte
        try:
            conn.execute(text("ALTER TABLE projects ADD COLUMN sort_order INTEGER NOT NULL DEFAULT 0"))
            conn.commit()
        except Exception:
            pass  # Spalte existiert bereits

        # projects: briefing-Spalte
        try:
            conn.execute(text("ALTER TABLE projects ADD COLUMN briefing TEXT DEFAULT ''"))
            conn.commit()
        except Exception:
            pass  # Spalte existiert bereits

        # snapshots: sort_order-Spalte
        try:
            conn.execute(text("ALTER TABLE snapshots ADD COLUMN sort_order INTEGER NOT NULL DEFAULT 0"))
            conn.commit()
        except Exception:
            pass  # Spalte existiert bereits

        # slides: Tabelle anlegen (falls noch nicht vorhanden)
        try:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS slides (
                    id          INTEGER PRIMARY KEY AUTOINCREMENT,
                    project_id  INTEGER NOT NULL REFERENCES projects(id) ON DELETE CASCADE,
                    slide_type  VARCHAR(20) NOT NULL DEFAULT 'section',
                    title       VARCHAR(300) DEFAULT '',
                    description TEXT DEFAULT '',
                    content     TEXT DEFAULT '',
                    sort_order  INTEGER NOT NULL DEFAULT 0,
                    created_at  DATETIME DEFAULT CURRENT_TIMESTAMP
                )
            """))
            conn.commit()
        except Exception as ex:
            log.warning("Migration slides: %s", ex)

        # slides: content-Spalte (falls Tabelle bereits ohne sie existiert)
        try:
            conn.execute(text("ALTER TABLE slides ADD COLUMN content TEXT DEFAULT ''"))
            conn.commit()
        except Exception:
            pass  # Spalte existiert bereits

        # region_interest: run_tag-Spalte
        try:
            conn.execute(text(
                "ALTER TABLE region_interest ADD COLUMN run_tag VARCHAR(100) NOT NULL DEFAULT ''"
            ))
            conn.commit()
        except Exception:
            pass  # Spalte existiert bereits

        # keyword_projects: Many-to-many Verknüpfungstabelle
        try:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS keyword_projects (
                    keyword_id INTEGER NOT NULL REFERENCES keywords(id) ON DELETE CASCADE,
                    project_id INTEGER NOT NULL REFERENCES projects(id) ON DELETE CASCADE,
                    PRIMARY KEY (keyword_id, project_id)
                )
            """))
            conn.commit()
        except Exception as ex:
            log.warning("Migration keyword_projects: %s", ex)

        # events: project_id-Spalte
        try:
            conn.execute(text("ALTER TABLE events ADD COLUMN project_id INTEGER REFERENCES projects(id) ON DELETE SET NULL"))
            conn.commit()
        except Exception:
            pass  # Spalte existiert bereits

        # ── Multi-User Migration ──
        # user_id Spalten zu bestehenden Tabellen hinzufuegen
        for tbl, col in [
            ("keywords", "user_id"), ("projects", "user_id"), ("events", "user_id"),
            ("alerts", "user_id"), ("fetch_log", "user_id"), ("snapshots", "user_id"),
        ]:
            try:
                conn.execute(text(f"ALTER TABLE {tbl} ADD COLUMN {col} INTEGER REFERENCES users(id)"))
                conn.commit()
            except Exception:
                pass

        # users: max_keywords Spalte
        try:
            conn.execute(text("ALTER TABLE users ADD COLUMN max_keywords INTEGER DEFAULT 20"))
            conn.commit()
        except Exception:
            pass

        # users: allowed_backends Spalte
        try:
            conn.execute(text("ALTER TABLE users ADD COLUMN allowed_backends TEXT DEFAULT ''"))
            conn.commit()
        except Exception:
            pass

        # snapshots: content_hash Spalte (SHA-256 Integritäts-Hash)
        try:
            conn.execute(text("ALTER TABLE snapshots ADD COLUMN content_hash VARCHAR(64) DEFAULT ''"))
            conn.commit()
            # Hashes für bestehende Snapshots nachberechnen
            for snap in Snapshot.query.filter(
                db.or_(Snapshot.content_hash.is_(None), Snapshot.content_hash == "")
            ).all():
                snap.compute_hash()
            db.session.commit()
            log.info("SHA-256 Hashes für bestehende Snapshots berechnet")
        except Exception:
            pass

        # app_settings: Migration von key-only PK zu id PK
        try:
            cols = [r[1] for r in conn.execute(text("PRAGMA table_info(app_settings)")).fetchall()]
            if "id" not in cols:
                log.info("Migriere app_settings → id + user_id …")
                conn.execute(text("""
                    CREATE TABLE app_settings_v2 (
                        id       INTEGER PRIMARY KEY AUTOINCREMENT,
                        key      VARCHAR(50) NOT NULL,
                        value    TEXT DEFAULT '',
                        user_id  INTEGER REFERENCES users(id),
                        UNIQUE(key, user_id)
                    )
                """))
                conn.execute(text(
                    "INSERT INTO app_settings_v2 (key, value) SELECT key, value FROM app_settings"
                ))
                conn.execute(text("DROP TABLE app_settings"))
                conn.execute(text("ALTER TABLE app_settings_v2 RENAME TO app_settings"))
                conn.commit()
                log.info("Migration app_settings abgeschlossen")
        except Exception as ex:
            try: conn.rollback()
            except Exception: pass
            log.warning("Migration app_settings: %s", ex)

    # Initialen Superadmin anlegen (falls keine User vorhanden)
    if User.query.count() == 0:
        admin = User(
            username="admin",
            display_name="Administrator",
            role="superadmin",
            active=True,
            max_serpapi_calls=0,
            max_llm_calls=0,
            max_projects=0,
        )
        admin.set_password("admin")
        db.session.add(admin)
        db.session.commit()
        log.info("Initialer Superadmin 'admin' angelegt (Passwort: admin)")

        # Bestehende Daten dem Admin zuweisen
        admin_id = admin.id
        for model in [Keyword, Project, Event, Alert, FetchLog, Snapshot]:
            model.query.filter_by(user_id=None).update({"user_id": admin_id})
        AppSetting.query.filter_by(user_id=None).update({"user_id": None})  # globale Settings bleiben global
        db.session.commit()
        log.info("Bestehende Daten dem Admin zugewiesen")

    log.info("Datenbank initialisiert")


def _run_analysis(method, body):
    """Führt eine Analyse intern aus (ohne HTTP-Request). Wird von APA genutzt.

    Delegiert an die compute()-Methode des zuständigen Analysis-Plugins.
    """
    # Methoden-Name → Plugin-ID Mapping
    _METHOD_TO_PLUGIN = {
        "spike_coincidence": "spike_coin",
        "changepoint": "cpd",
        "rolling_correlation": "rc",
        "periodicity": "period_filter",
        "outliers": "outlier",
        "decompose": "decomp",
        "self_similarity": "ssim",
    }

    plugin_id = _METHOD_TO_PLUGIN.get(method)
    if plugin_id:
        plugin = PluginManager.get("analysis", plugin_id)
        if plugin:
            result = plugin.compute(body)
            result.pop("_status", None)
            return result
        return {"error": f"Plugin nicht gefunden: {plugin_id}", "summary": f"Plugin {plugin_id} nicht verfügbar"}

    return {"error": f"Unbekannte Methode: {method}", "summary": f"Unbekannte Methode: {method}"}


@app.route("/api/projects/<int:pid>/export-docx", methods=["GET"])
@login_required
def api_project_export_docx(pid):
    """Exportiert alle Projekt-Inhalte als Word-Dokument (.docx)."""
    import json as _json, io, base64, re
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn

    proj = Project.query.filter_by(id=pid, user_id=current_user.id).first()
    if not proj:
        abort(404)

    snaps  = Snapshot.query.filter_by(project_id=pid).all()
    slides = Slide.query.filter_by(project_id=pid).all()

    # Gleiche Sortierung wie im Frontend: sort_order, dann created_at desc
    items = []
    for s in snaps:
        d = s.to_dict()
        d["item_type"]   = "snapshot"
        d["_sort_order"] = s.sort_order or 0
        d["_created_at"] = s.created_at
        items.append(d)
    for sl in slides:
        d = sl.to_dict()
        d["item_type"]   = "slide"
        d["_sort_order"] = sl.sort_order or 0
        d["_created_at"] = sl.created_at
        items.append(d)
    items.sort(key=lambda x: (x["_sort_order"],
                               -(x["_created_at"].timestamp() if x["_created_at"] else 0)))

    GP = {"": "Web", "news": "News", "images": "Bilder",
          "youtube": "YouTube", "froogle": "Shopping"}
    TF = {"now 1-H": "1h", "now 4-H": "4h", "now 1-d": "24h", "now 7-d": "7 Tage",
          "today 1-m": "1 Monat", "today 3-m": "3 Monate",
          "today 12-m": "12 Monate", "today 5-y": "5 Jahre"}
    ATYPE = {"ssim": "Self-Similarity-Matrix",
             "outlier": "Ausreißer-Erkennung",
             "decomp": "Zeitreihen-Zerlegung"}

    doc = Document()

    # Seitenränder
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Cm(2.5)
        sec.top_margin  = sec.bottom_margin = Cm(2.5)

    # Deckblatt
    doc.add_heading(proj.name, level=0)
    if proj.description:
        doc.add_paragraph(proj.description)
    doc.add_paragraph(
        f"Exportiert: {datetime.now().strftime('%d.%m.%Y %H:%M')}  |  "
        f"Snapshots: {len(snaps)}  |  Seiten: {len(slides)}"
    ).runs[0].font.size = Pt(10)
    doc.add_page_break()

    # ── Markdown → Docx ──────────────────────────────────────────────────────
    def _add_inline(para, text):
        """Inline-Markdown (* ** ) in Runs umwandeln."""
        for part in re.split(r'(\*\*[^*\n]+\*\*|\*[^*\n]+\*)', text):
            if part.startswith("**") and part.endswith("**"):
                para.add_run(part[2:-2]).bold = True
            elif part.startswith("*") and part.endswith("*"):
                para.add_run(part[1:-1]).italic = True
            else:
                para.add_run(part)

    def _add_md(text):
        """Markdown-Text als Word-Absätze einfügen."""
        for line in text.split("\n"):
            hm = re.match(r"^(#{1,4}) (.+)", line)
            if hm:
                doc.add_heading(hm.group(2), level=min(len(hm.group(1)), 4))
                continue
            if re.match(r"^---+$", line.strip()):
                doc.add_paragraph("─" * 50)
                continue
            if re.match(r"^[-*] ", line):
                p = doc.add_paragraph(style="List Bullet")
                _add_inline(p, line[2:])
                continue
            if re.match(r"^\d+\. ", line):
                p = doc.add_paragraph(style="List Number")
                _add_inline(p, re.sub(r"^\d+\. ", "", line))
                continue
            if not line.strip():
                doc.add_paragraph()
                continue
            p = doc.add_paragraph()
            _add_inline(p, line)

    def _add_img_b64(data_url, width_cm=15):
        """Base64-Bild in Dokument einbetten."""
        try:
            raw = base64.b64decode(data_url.split(",")[-1])
            doc.add_picture(io.BytesIO(raw), width=Cm(width_cm))
        except Exception:
            pass

    def _render_trend_chart(chart, width_cm=15):
        """Trend-Liniendiagramm mit matplotlib rendern und ins Dokument einbetten."""
        datasets = chart.get("datasets", [])
        labels   = chart.get("labels", [])
        if not datasets:
            return
        try:
            w_in = width_cm / 2.54
            fig, ax = plt.subplots(figsize=(w_in, w_in / 2.5))
            ax.set_facecolor("#1a1a2e")
            fig.patch.set_facecolor("#16213e")
            for ds in datasets:
                data   = ds.get("data", [])
                color  = ds.get("borderColor", "#aaaaff")
                dlabel = ds.get("label", "")
                if not data:
                    continue
                if isinstance(data[0], dict):
                    xs = [p.get("x") for p in data]
                    ys = [p.get("y") for p in data]
                else:
                    xs = labels[:len(data)]
                    ys = data
                # parse x-axis as dates when possible
                try:
                    from datetime import datetime as _dt
                    xs_d = [_dt.fromisoformat(str(x)[:19]) for x in xs]
                    ax.plot(xs_d, ys, color=color, linewidth=1.2, label=dlabel)
                    ax.xaxis.set_major_formatter(mdates.DateFormatter("%d.%m.%y"))
                    ax.xaxis.set_major_locator(mdates.AutoDateLocator())
                except Exception:
                    ax.plot(range(len(ys)), ys, color=color, linewidth=1.2, label=dlabel)

            ax.set_ylim(0, 100)
            ax.tick_params(colors="#cccccc", labelsize=7)
            ax.spines[:].set_color("#444466")
            ax.yaxis.set_major_locator(MaxNLocator(integer=True, nbins=5))
            fig.autofmt_xdate(rotation=35, ha="right")
            if len(datasets) > 1:
                leg = ax.legend(fontsize=7, facecolor="#1a1a2e", labelcolor="#cccccc",
                                framealpha=0.8)
            fig.tight_layout(pad=0.5)
            buf = io.BytesIO()
            fig.savefig(buf, format="png", dpi=130, bbox_inches="tight",
                        facecolor=fig.get_facecolor())
            plt.close(fig)
            buf.seek(0)
            doc.add_picture(buf, width=Cm(width_cm))
        except Exception as e:
            log.warning("Chart render failed: %s", e)

    def _stats_table(datasets, labels):
        stats = []
        for ds in datasets:
            data = ds.get("data", [])
            if not data:
                continue
            if isinstance(data[0], dict):
                vals = [p["y"] for p in data if p.get("y") is not None]
                d0 = str(data[0].get("x", "?"))[:10]
                d1 = str(data[-1].get("x", "?"))[:10]
                pts = [(str(p.get("x","?"))[:10], p["y"]) for p in data if p.get("y") is not None]
            else:
                vals = [v for v in data if v is not None]
                d0 = str(labels[0])[:10] if labels else "?"
                d1 = str(labels[-1])[:10] if labels else "?"
                pts = [(str(labels[j])[:10] if j < len(labels) else str(j), v)
                       for j, v in enumerate(data) if v is not None]
            if not vals:
                continue
            mn, mx = min(vals), max(vals)
            avg    = sum(vals) / len(vals)
            third  = max(1, len(vals) // 3)
            diff   = (sum(vals[-third:]) / third) - (sum(vals[:third]) / third)
            trend  = "steigend" if diff > 5 else "fallend" if diff < -5 else "stabil"
            peak = max(pts, key=lambda p: p[1]) if pts else None
            stats.append({"label": ds.get("label","?"), "d0": d0, "d1": d1,
                           "n": len(vals), "min": mn, "max": mx, "avg": avg,
                           "trend": trend, "peak": peak})
        if not stats:
            return
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        for i, h in enumerate(["Reihe", "Zeitraum", "Min", "Max", "Ø", "Trend"]):
            tbl.rows[0].cells[i].text = h
        for s in stats:
            row = tbl.add_row().cells
            row[0].text = s["label"]
            row[1].text = f"{s['d0']} – {s['d1']}"
            row[2].text = f"{s['min']:.0f}"
            row[3].text = f"{s['max']:.0f}"
            row[4].text = f"{s['avg']:.1f}"
            row[5].text = s["trend"]
        doc.add_paragraph()

    # ── Inhalte ───────────────────────────────────────────────────────────────
    for item in items:
        if item["item_type"] == "snapshot":
            doc.add_heading(item.get("title") or f"Snapshot #{item['id']}", level=1)
            if item.get("_created_at"):
                p = doc.add_paragraph(
                    item["_created_at"].strftime("%d.%m.%Y %H:%M"))
                p.runs[0].font.size = Pt(9)
            if item.get("content_hash"):
                p = doc.add_paragraph(f"SHA-256: {item['content_hash']}")
                p.runs[0].font.size = Pt(7)
                p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            chart = item.get("chart", {})
            if chart.get("type") == "analysis":
                atype = ATYPE.get(chart.get("analysis_type", ""),
                                  chart.get("analysis_type", "Analyse"))
                doc.add_paragraph(f"Analyse-Typ: {atype}")
                if chart.get("subtitle"):
                    doc.add_paragraph(f"Keyword: {chart['subtitle']}")
                if chart.get("image"):
                    _add_img_b64(chart["image"])
            else:
                kw_meta = chart.get("keywords_meta", [])
                if kw_meta:
                    tbl = doc.add_table(rows=1, cols=4)
                    tbl.style = "Table Grid"
                    for i, h in enumerate(["Keyword", "Geo", "Kategorie", "Zeitraum"]):
                        tbl.rows[0].cells[i].text = h
                    for m in kw_meta:
                        row = tbl.add_row().cells
                        row[0].text = m.get("keyword", "")
                        row[1].text = m.get("geo", "") or "Weltweit"
                        row[2].text = GP.get(m.get("gprop", ""), "Web")
                        row[3].text = TF.get(m.get("timeframe", ""),
                                             m.get("timeframe", ""))
                    doc.add_paragraph()

                _render_trend_chart(chart)
                _stats_table(chart.get("datasets", []), chart.get("labels", []))

                corrs = chart.get("correlations", [])
                if corrs:
                    doc.add_heading("Korrelationen", level=3)
                    for c in corrs:
                        r = c.get("r")
                        if r is not None:
                            s = ("stark" if abs(r) > 0.7
                                 else "mittel" if abs(r) > 0.3 else "schwach")
                            p = doc.add_paragraph(style="List Bullet")
                            p.add_run(
                                f"{c.get('labelA','?')} ↔ {c.get('labelB','?')}: "
                                f"r={r:+.3f} ({s})")

            markers = item.get("markers", [])
            if markers:
                doc.add_heading("Markierungen", level=3)
                for m in markers:
                    lbl = m.get("label", "")
                    cmt = m.get("comment", "")
                    if lbl or cmt:
                        p = doc.add_paragraph(style="List Bullet")
                        run = p.add_run(f"M{m.get('num','')}: ")
                        run.bold = True
                        p.add_run(f"{lbl}{' – ' + cmt if cmt else ''}")

            if item.get("comment", "").strip():
                doc.add_heading("Kommentar", level=3)
                _add_md(item["comment"])

        else:  # slide
            stype = item.get("slide_type", "")
            title = item.get("title") or ""
            doc.add_heading(title or f"[{stype}]", level=1)

            desc = item.get("description", "")
            if desc:
                # Inline-Snapshot-Marker durch Verweistext ersetzen
                desc = re.sub(
                    r'\{\{SNAPSHOT:(\d+)\}\}',
                    lambda m: f"[→ Siehe Snapshot #{m.group(1)}]",
                    desc)
                if stype in ("textbild", "title", "section"):
                    _add_md(desc)
                else:
                    doc.add_paragraph(desc)

            if stype == "textbild":
                content = _json.loads(item.get("content") or "{}")
                if content.get("image_data_url"):
                    _add_img_b64(content["image_data_url"])

        doc.add_paragraph()  # Abstand zwischen Einträgen

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = re.sub(r'[^\w\s\-äöüÄÖÜß]', '_', proj.name).strip()
    filename  = f"{safe_name}.docx"
    import hashlib as _hl
    export_hash = _hl.sha256(buf.getvalue()).hexdigest()
    audit_log("export_docx", "project", pid,
              f"Datei: {filename}, SHA-256: {export_hash[:16]}…",
              content_hash=export_hash, project_id=pid)
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )




# (Scientific Paper Export → plugins/ai/scientific_paper/)
_REMOVED_SCI_PAPER = True  # Marker: alter Code entfernt, jetzt Plugin

# APA-Tools (ausgelagert nach apa_tools.py)
from apa_tools import _WZ_CORE_TOOLS, _get_wz_tools, _get_wz_tools_oai, _execute_wz_tool  # noqa: F401

# APA-Stream (ausgelagert nach apa_stream.py)
from apa_stream import api_ai_project_assist_stream
app.add_url_rule("/api/ai-project-assist/stream",
                 endpoint="api_ai_project_assist_stream",
                 view_func=login_required(api_ai_project_assist_stream),
                 methods=["POST"])


# ---------------------------------------------------------------------------
scheduler.start()
log.info(
    "Scheduler gestartet – täglicher Fetch um %02d:%02d Uhr.",
    FETCH_HOUR, FETCH_MINUTE,
)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=False, use_reloader=False)
